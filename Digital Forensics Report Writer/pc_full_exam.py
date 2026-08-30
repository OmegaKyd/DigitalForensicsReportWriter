import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt
import os
import sys
import time
import ctypes
import re
import PyPDF2
from tkinter.scrolledtext import ScrolledText
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from settings_manager import SettingsManager
from ui_theme import apply_theme, add_header_bar, add_action_bar, pack_right_actions, size_window, build_extracted_info_pane, parse_mdy_date, add_date_entry, COLORS, FormTabs, close_and_return
from app_menu import attach_app_menu
from paragraphs_manager import fill_paragraph, load_paragraphs
from report_common import (
    ask_open_file,
    ask_directory,
    default_export_dir,
    remember_folder,
    suggested_report_filename,
    unique_output_path,
    apply_suggested_filename,
    show_placeholder_preview,
    pc_preview_rows,
    prefer_gui_over_parsed,
    looks_like_digital_collector_log,
    parse_digital_collector_log,
    apply_log_device_fields_to_form,
    merge_log_device_into_report_data,
    load_request_titles,
    setup_title_combobox,
    remember_titles_from_form,
    setup_agency_combobox,
    remember_agencies_from_form,
    title_agency_words,
    saved_examiner_agency,
    bind_prefix_typeahead,
    refresh_request_title_values,
    current_dfr_prefix,
    is_complete_dfr_number,
    add_template_picker,
    sync_template_choice,
)
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

#ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)

class PCFullExam(TkinterDnD.Tk):
    def __init__(self, master=None):
        super().__init__()
        
        self.role_type = None # Initialize to avoid AttributeError

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.master = master
        apply_theme(self)
        attach_app_menu(self)
        add_header_bar(self, "Full Computer Exam", "TX1, FTK, X-Ways, and Digital Collector image reports")
        self.action_bar = add_action_bar(self)
        self.preview_button = ttk.Button(self.action_bar, text="Preview", command=self.preview_placeholders)
        self.generate_button = ttk.Button(self.action_bar, text="Generate Report", command=self.generate_report)
        self.exit_button = ttk.Button(self.action_bar, text="Exit", command=self.on_closing)
        pack_right_actions(self.preview_button, self.generate_button, self.exit_button)
        
        # Initialize settings manager
        self.settings_manager = SettingsManager()
        self.current_settings = self.settings_manager.load_settings()
        self._save_timer = None 
        
        # Initialize PC exam specific paragraphs
        self.paragraphs = {}
        self._paragraph_kind = "pc_full"
        self.initialize_pc_exam_paragraphs()
        
        self.title("Ω Digital Forensics Report Writer - Full Computer Device Exam Ω")
        
        size_window(self, 1260, 700)
        
        self.main_frame = ttk.Frame(self)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.paned_window = ttk.PanedWindow(self.main_frame, orient=tk.HORIZONTAL)
        self.paned_window.pack(fill=tk.BOTH, expand=True)
        
        self.left_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.left_frame, weight=3)

        self.form_tabs = FormTabs(self.left_frame)
        self.tab_request = self.form_tabs.add_tab("request", "Request Info")
        self.tab_examiner = self.form_tabs.add_tab("examiner", "Examiner Info")
        self.tab_device = self.form_tabs.add_tab("device", "Device Info")
        self.tab_output = self.form_tabs.add_tab("output", "Output Info")
        self.scrollable_frame = self.tab_request
        self.middle_frame = self.tab_device

        self.right_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.right_frame, weight=2)
        
        self.create_widgets()
        self.create_info_display()
        
        self.bind_all("<MouseWheel>", self._on_mousewheel)
        self.bind_tab_status_events()
        self.refresh_tab_status()

    def _on_mousewheel(self, event):
        if hasattr(self, "form_tabs"):
            self.form_tabs.on_mousewheel(event)

    def _field_filled(self, widget):
        if widget is None:
            return False
        try:
            return bool(widget.get().strip())
        except Exception:
            return False

    def bind_tab_status_events(self):
        for name in (
            "role_type", "request_date", "request_agency", "request_officer", "case_type",
            "offense_type", "examiner_name", "DFR_Num", "device_owner", "save_location",
        ):
            widget = getattr(self, name, None)
            if widget is None:
                continue
            try:
                widget.bind("<KeyRelease>", lambda e: self.refresh_tab_status(), add="+")
                widget.bind("<<ComboboxSelected>>", lambda e: self.refresh_tab_status(), add="+")
            except Exception:
                pass

    def refresh_tab_status(self):
        if not hasattr(self, "form_tabs"):
            return
        request_ok = self._field_filled(getattr(self, "role_type", None))
        if getattr(self, "role_type", None) and self.role_type.get() == "Agency Assist":
            request_ok = all(self._field_filled(getattr(self, name, None)) for name in (
                "request_date", "request_agency", "request_officer", "case_type"
            ))
        elif getattr(self, "role_type", None) and self.role_type.get() == "Case Agent":
            request_ok = self._field_filled(getattr(self, "offense_type", None))
        examiner_ok = self._field_filled(getattr(self, "examiner_name", None)) and is_complete_dfr_number(self.DFR_Num.get() if hasattr(self, "DFR_Num") else "")
        device_ok = self._field_filled(getattr(self, "device_owner", None))
        try:
            device_ok = device_ok and bool(self.get_selected_forensic_software())
        except Exception:
            pass
        output_ok = self._field_filled(getattr(self, "save_location", None))
        if output_ok:
            try:
                output_ok = os.path.isdir(self.save_location.get().strip())
            except Exception:
                output_ok = False
        self.form_tabs.set_complete("request", request_ok)
        self.form_tabs.set_complete("examiner", examiner_ok)
        self.form_tabs.set_complete("device", device_ok)
        self.form_tabs.set_complete("output", output_ok)

    def preview_placeholders(self):
        self.generate_report(preview_only=True)

    def create_widgets(self):
        """Main method to create all UI widgets - calls individual frame creation methods"""
        style = ttk.Style()
        style.configure('Header.TLabel', font=('Arial', 12, 'bold'))
        style.configure('Section.TFrame', padding='10')

        # Create each section in separate methods
        self.create_left_column_widgets()
        self.create_middle_column_widgets() 
        self.create_right_column_widgets()

    def create_left_column_widgets(self):
        """Create all widgets for the left column"""
        self.create_role_selection_frame()
        self.create_request_information_frame()
        self.create_transfer_information_frame()
        self.create_examiner_information_frame()

    def create_role_selection_frame(self):
        """Create the examination type/role selection frame"""
        role_frame = ttk.LabelFrame(self.scrollable_frame, text="Examination Type", padding="10")
        role_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(role_frame, text="Select Role:").grid(row=0, column=0, sticky="w", pady=2)
        self.role_type = ttk.Combobox(role_frame, 
            values=["Agency Assist", "Case Agent"], state="readonly")
        self.role_type.grid(row=0, column=1, sticky="ew", pady=2)
        self.role_type.set("Agency Assist")  
        self.role_type.bind('<<ComboboxSelected>>', self.toggle_request_content)
        self.setup_auto_complete_dropdown(self.role_type)

        role_frame.columnconfigure(1, weight=1)

    def create_request_information_frame(self):
        """Create the request information frame with dynamic content"""
        request_frame = ttk.LabelFrame(self.scrollable_frame, text="Request Information", padding="10")
        request_frame.pack(fill=tk.X, padx=5, pady=5)

        # Create two frames for different sets of fields
        self.agency_assist_frame = ttk.Frame(request_frame)
        self.agency_assist_frame.pack(fill=tk.X, padx=5, pady=5)

        self.case_agent_frame = ttk.Frame(request_frame)
        self.case_agent_frame.pack(fill=tk.X, padx=5, pady=5)
        self.case_agent_frame.pack_forget()  # Hide initially

        self.create_agency_assist_fields()
        self.create_case_agent_fields()

    def create_agency_assist_fields(self):
        """Create agency assist specific fields"""
        # Request Date
        ttk.Label(self.agency_assist_frame, text="Exam Request Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.request_date = add_date_entry(self.agency_assist_frame, row=0, column=1)

        # Request Agency
        ttk.Label(self.agency_assist_frame, text="Requesting Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.request_agency = ttk.Combobox(self.agency_assist_frame)
        self.request_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.request_agency)

        # Request Title (with dropdown and custom entry)
        ttk.Label(self.agency_assist_frame, text="Requesting Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        request_title_frame = ttk.Frame(self.agency_assist_frame)
        request_title_frame.grid(row=2, column=1, sticky="ew", pady=2)

        self.request_title_type = ttk.Combobox(request_title_frame, values=load_request_titles())
        self.request_title_type.pack(fill=tk.X, expand=True)
        self.request_title_type.set("")
        bind_prefix_typeahead(self.request_title_type)

        self.request_title_entry = ttk.Entry(request_title_frame)
        self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        self.request_title_entry.pack_forget()

        # Request Officer
        ttk.Label(self.agency_assist_frame, text="Requesting Officer:").grid(row=3, column=0, sticky="w", pady=2)
        self.request_officer = ttk.Entry(self.agency_assist_frame)
        self.request_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Offense Type
        ttk.Label(self.agency_assist_frame, text="Primary Case Offense:").grid(row=4, column=0, sticky="w", pady=2)
        self.case_type = ttk.Entry(self.agency_assist_frame)
        self.case_type.grid(row=4, column=1, sticky="ew", pady=2)

        # Legal Authority
        ttk.Label(self.agency_assist_frame, text="Legal Authority:").grid(row=5, column=0, sticky="w", pady=2)
        self.legal_authority = ttk.Combobox(self.agency_assist_frame, 
            values=['Search Warrant', 'Consent', 'Parole', 'Implied Consent'], state="readonly")
        self.legal_authority.grid(row=5, column=1, sticky="ew", pady=2)
        self.legal_authority.set('Search Warrant')
        self.legal_authority.bind('<<ComboboxSelected>>', self.toggle_time_frame_section)
        self.setup_auto_complete_dropdown(self.legal_authority)

        # Time Frame Section (initially visible since default is Search Warrant)
        self.time_frame_label = ttk.Label(self.agency_assist_frame, text="Time Frame Limited?:")
        self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)

        self.time_frame_checkbox_frame = ttk.Frame(self.agency_assist_frame)
        self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)

        self.time_frame_var = tk.IntVar(self.agency_assist_frame)
        self.time_frame_checkbox = ttk.Checkbutton(
            self.time_frame_checkbox_frame,
            text="Yes",
            variable=self.time_frame_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_time_frame_fields
        )
        self.time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        self.time_frame_var.set(0)  # Default to unchecked

        # Time Frame Date Fields (initially hidden)
        self.time_frame_fields = ttk.Frame(self.agency_assist_frame)
        self.time_frame_fields.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        self.time_frame_fields.grid_remove()

        ttk.Label(self.time_frame_fields, text="Time Frame Start Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.time_frame_start = add_date_entry(self.time_frame_fields, row=0, column=1)

        ttk.Label(self.time_frame_fields, text="Time Frame End Date (M/D/Y):").grid(row=1, column=0, sticky="w", pady=2)
        self.time_frame_end = add_date_entry(self.time_frame_fields, row=1, column=1)

        # Configure grid weights
        self.time_frame_fields.columnconfigure(1, weight=1)
        self.agency_assist_frame.columnconfigure(1, weight=1)

    def create_case_agent_fields(self):
        """Create case agent specific fields"""
        # Offense Type
        ttk.Label(self.case_agent_frame, text="Primary Case Offense:").grid(row=0, column=0, sticky="w", pady=2)
        self.offense_type = ttk.Entry(self.case_agent_frame)
        self.offense_type.grid(row=0, column=1, sticky="ew", pady=2)

        # Legal Authority
        ttk.Label(self.case_agent_frame, text="Legal Authority:").grid(row=1, column=0, sticky="w", pady=2)
        self.legal_self = ttk.Combobox(self.case_agent_frame, 
            values=['Search Warrant', 'Consent', 'Parole', 'Implied Consent'], state="readonly")
        self.legal_self.grid(row=1, column=1, sticky="ew", pady=2)
        self.legal_self.set('Search Warrant')
        self.legal_self.bind('<<ComboboxSelected>>', self.toggle_sw_date_and_time_frame)
        self.setup_auto_complete_dropdown(self.legal_self)

        # Search Warrant Date (initially visible since default is Search Warrant)
        self.sw_date_frame = ttk.Frame(self.case_agent_frame)
        self.sw_date_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=2)

        ttk.Label(self.sw_date_frame, text="Search Warrant Service Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.sw_service_date = add_date_entry(self.sw_date_frame, row=0, column=1)

        # Time Frame Section for Case Agent (initially visible since default is Search Warrant)
        self.case_agent_time_frame_label = ttk.Label(self.case_agent_frame, text="Time Frame Limited?:")
        self.case_agent_time_frame_label.grid(row=3, column=0, sticky="w", pady=2)

        self.case_agent_time_frame_checkbox_frame = ttk.Frame(self.case_agent_frame)
        self.case_agent_time_frame_checkbox_frame.grid(row=3, column=1, sticky="ew", pady=2)

        self.case_agent_time_frame_var = tk.IntVar(self.case_agent_frame)
        self.case_agent_time_frame_checkbox = ttk.Checkbutton(
            self.case_agent_time_frame_checkbox_frame,
            text="Yes",
            variable=self.case_agent_time_frame_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_case_agent_time_frame_fields
        )
        self.case_agent_time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        self.case_agent_time_frame_var.set(0)  # Default to unchecked

        # Time Frame Date Fields for Case Agent (initially hidden)
        self.case_agent_time_frame_fields = ttk.Frame(self.case_agent_frame)
        self.case_agent_time_frame_fields.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        self.case_agent_time_frame_fields.grid_remove()

        ttk.Label(self.case_agent_time_frame_fields, text="Time Frame Start Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.case_agent_time_frame_start = add_date_entry(self.case_agent_time_frame_fields, row=0, column=1)

        ttk.Label(self.case_agent_time_frame_fields, text="Time Frame End Date (M/D/Y):").grid(row=1, column=0, sticky="w", pady=2)
        self.case_agent_time_frame_end = add_date_entry(self.case_agent_time_frame_fields, row=1, column=1)

        # Configure grid weights
        self.sw_date_frame.columnconfigure(1, weight=1)
        self.case_agent_time_frame_fields.columnconfigure(1, weight=1)
        self.case_agent_frame.columnconfigure(1, weight=1)

    def toggle_time_frame_section(self, event=None):
        """Show/hide time frame section based on legal authority selection for Agency Assist"""
        if self.legal_authority.get() == 'Search Warrant':
            # Show time frame section
            self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)
            self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)
            # If checkbox is checked, also show the fields
            self.toggle_time_frame_fields()
        else:
            # Hide time frame section and clear values
            self.time_frame_label.grid_remove()
            self.time_frame_checkbox_frame.grid_remove()
            self.time_frame_fields.grid_remove()
            
            # Clear values
            self.time_frame_var.set(0)
            if hasattr(self, 'time_frame_start'):
                self.time_frame_start.delete(0, tk.END)
            if hasattr(self, 'time_frame_end'):
                self.time_frame_end.delete(0, tk.END)

    def toggle_time_frame_fields(self, event=None):
        """Show/hide time frame date fields based on checkbox state for Agency Assist"""
        if self.time_frame_var.get() == 1:
            self.time_frame_fields.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            self.time_frame_fields.grid_remove()
            # Clear values when hiding
            if hasattr(self, 'time_frame_start'):
                self.time_frame_start.delete(0, tk.END)
            if hasattr(self, 'time_frame_end'):
                self.time_frame_end.delete(0, tk.END)

    def toggle_sw_date_and_time_frame(self, event=None):
        """Handle both SW date and time frame toggling for Case Agent"""
        if hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant':
            # Show the SW date frame
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=2)
            
            # Show time frame section
            self.case_agent_time_frame_label.grid(row=3, column=0, sticky="w", pady=2)
            self.case_agent_time_frame_checkbox_frame.grid(row=3, column=1, sticky="ew", pady=2)
            # If checkbox is checked, also show the fields
            self.toggle_case_agent_time_frame_fields()
        else:
            # Hide the SW date frame and clear its value
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid_remove()
                
                # Clear the SW service date field when hiding it
                if hasattr(self, 'sw_service_date'):
                    self.sw_service_date.delete(0, tk.END)
            
            # Hide time frame section and clear values
            self.case_agent_time_frame_label.grid_remove()
            self.case_agent_time_frame_checkbox_frame.grid_remove()
            self.case_agent_time_frame_fields.grid_remove()
            
            # Clear values
            self.case_agent_time_frame_var.set(0)
            if hasattr(self, 'case_agent_time_frame_start'):
                self.case_agent_time_frame_start.delete(0, tk.END)
            if hasattr(self, 'case_agent_time_frame_end'):
                self.case_agent_time_frame_end.delete(0, tk.END)

    def toggle_case_agent_time_frame_fields(self, event=None):
        """Show/hide time frame date fields based on checkbox state for Case Agent"""
        if self.case_agent_time_frame_var.get() == 1:
            self.case_agent_time_frame_fields.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            self.case_agent_time_frame_fields.grid_remove()
            # Clear values when hiding
            if hasattr(self, 'case_agent_time_frame_start'):
                self.case_agent_time_frame_start.delete(0, tk.END)
            if hasattr(self, 'case_agent_time_frame_end'):
                self.case_agent_time_frame_end.delete(0, tk.END)

    def create_transfer_information_frame(self):
        """Create the transfer information frame"""
        transfer_frame = ttk.LabelFrame(self.scrollable_frame, text="Transfer Information", padding="10")
        transfer_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Transfer - Changed from Combobox to Single Yes Checkbox
        ttk.Label(transfer_frame, text="Device Transferred by Another Officer:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(transfer_frame)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variable
        self.device_transfer_var = tk.IntVar(transfer_frame)
        
        # Create single Yes checkbox
        self.transfer_yes_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Yes",
            variable=self.device_transfer_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_transfer_fields
        )
        self.transfer_yes_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Set default to unchecked (0 = No)
        self.device_transfer_var.set(0)

        # Transfer Fields (initially hidden)
        self.transfer_fields = ttk.Frame(transfer_frame)
        self.transfer_fields.grid(row=1, column=0, columnspan=2, sticky="ew", pady=2)
        self.transfer_fields.grid_remove()

        ttk.Label(self.transfer_fields, text="Transfer Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.transfer_date = add_date_entry(self.transfer_fields, row=0, column=1)

        ttk.Label(self.transfer_fields, text="Transfer Officer Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.transfer_agency = ttk.Combobox(self.transfer_fields)
        self.transfer_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.transfer_agency)

        ttk.Label(self.transfer_fields, text="Transfer Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        self.transfer_title = ttk.Combobox(self.transfer_fields, values=load_request_titles())
        self.transfer_title.grid(row=2, column=1, sticky="ew", pady=2)
        setup_title_combobox(self.transfer_title)

        ttk.Label(self.transfer_fields, text="Transfer Officer Name:").grid(row=3, column=0, sticky="w", pady=2)
        self.transfer_officer = ttk.Entry(self.transfer_fields)
        self.transfer_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Configure grid weights
        transfer_frame.columnconfigure(1, weight=1)
        self.transfer_fields.columnconfigure(1, weight=1)

    def create_examiner_information_frame(self):
        examiner_frame = ttk.LabelFrame(self.tab_examiner, text="Examiner Information", padding="10")
        examiner_frame.pack(fill=tk.X, padx=5, pady=5)

        # Examiner Agency
        ttk.Label(examiner_frame, text="Examiner Agency:").grid(row=0, column=0, sticky="w", pady=2)
        examiner_agency_frame = ttk.Frame(examiner_frame)
        examiner_agency_frame.grid(row=0, column=1, sticky="ew", pady=2)

        self.examiner_agency_type = ttk.Combobox(examiner_agency_frame)
        self.examiner_agency_type.pack(fill=tk.X, expand=True)
        setup_agency_combobox(
            self.examiner_agency_type,
            saved=saved_examiner_agency(self.current_settings),
            on_change=self.auto_save_settings,
        )

        self.examiner_agency_entry = ttk.Entry(examiner_agency_frame)
        self.examiner_agency_entry.pack_forget()

        # Examiner Title
        ttk.Label(examiner_frame, text="Examiner Title:").grid(row=1, column=0, sticky="w", pady=2)
        examiner_title_frame = ttk.Frame(examiner_frame)
        examiner_title_frame.grid(row=1, column=1, sticky="ew", pady=2)

        self.examiner_title_type = ttk.Combobox(examiner_title_frame, values=load_request_titles())
        self.examiner_title_type.pack(fill=tk.X, expand=True)
        setup_title_combobox(
            self.examiner_title_type,
            saved=self.current_settings.get("examiner_title", ""),
            on_change=self.auto_save_settings,
        )

        self.examiner_title_entry = ttk.Entry(examiner_title_frame)
        self.examiner_title_entry.pack_forget()

        # Examiner Name
        ttk.Label(examiner_frame, text="Examiner Name:").grid(row=2, column=0, sticky="w", pady=2)
        self.examiner_name = ttk.Entry(examiner_frame)
        self.examiner_name.grid(row=2, column=1, sticky="ew", pady=2)
        # Load saved setting
        if self.current_settings.get("examiner_name"):
            self.examiner_name.insert(0, self.current_settings["examiner_name"])
        
        # Bind auto-save events
        self.examiner_name.bind('<KeyRelease>', self.auto_save_settings)
        self.examiner_name.bind('<FocusOut>', self.auto_save_settings)

        # Case Number
        ttk.Label(examiner_frame, text="Case Number:").grid(row=3, column=0, sticky="w", pady=2)
        self.case_number = ttk.Entry(examiner_frame)
        self.case_number.grid(row=3, column=1, sticky="ew", pady=2)

        # Evidence Number
        ttk.Label(examiner_frame, text="Evidence Number:").grid(row=4, column=0, sticky="w", pady=2)
        self.evidence_number = ttk.Entry(examiner_frame)
        self.evidence_number.grid(row=4, column=1, sticky="ew", pady=2)

        # DFR Report Number
        ttk.Label(examiner_frame, text="DFR Report #:").grid(row=5, column=0, sticky="w", pady=2)
        self.DFR_Num = ttk.Entry(examiner_frame)
        self.DFR_Num.grid(row=5, column=1, sticky="ew", pady=2)
        
        # Load saved prefix or use default
        self.DFR_Num.insert(0, current_dfr_prefix())

        examiner_frame.columnconfigure(1, weight=1)


    def toggle_examiner_title_entry(self, event=None):
        if self.examiner_title_type.get() == "Other (specify)":
            self.examiner_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.examiner_title_entry.pack_forget()
            self.examiner_title_entry.delete(0, tk.END)
        
        # Auto-save when selection changes
        self.auto_save_settings()

    def on_agency_type_changed(self, event=None):
        # Call the original toggle method
        self.toggle_agency_entry(event)
        # Auto-save the change
        self.auto_save_settings()

    def auto_save_settings(self, event=None):
        # Use after_idle to prevent excessive saving during rapid typing
        if hasattr(self, '_save_timer') and self._save_timer:
            self.after_cancel(self._save_timer)
        
        # Schedule save for after user stops typing (500ms delay)
        self._save_timer = self.after(500, self._perform_auto_save)

    def _perform_auto_save(self):
        try:
            # Get examiner title value based on dropdown selection
            examiner_title_value = self.examiner_title_type.get()
                
            settings_to_save = {
                "examiner_agency_type": self.examiner_agency_type.get(),
                "examiner_agency_custom": "",
                "examiner_title": examiner_title_value,
                "examiner_name": self.examiner_name.get(),
                "dfr_number_prefix": self.get_dfr_prefix(),
                "version": "1.0.1"
            }
            
            self.settings_manager.save_settings(settings_to_save)
            # Update current_settings to reflect the save
            self.current_settings.update(settings_to_save)
            remember_agencies_from_form(self)
            
        except Exception as e:
            print(f"Error auto-saving examiner settings: {e}")

    def get_dfr_prefix(self):
        dfr_value = self.DFR_Num.get()
        # Find the last occurrence of "DFR" and include everything up to and including any year/dash pattern
        import re
        match = re.match(r'(DFR\d{4}-)', dfr_value)
        if match:
            return match.group(1)
        else:
            # Fallback to current year
            from datetime import datetime
            current_year = datetime.now().year
            return f"DFR{current_year}-"

    def create_middle_column_widgets(self):
        self.create_forensic_software_frame()
        self.create_device_info_frame()
        self.create_output_file_frame()

    def create_forensic_software_frame(self):
        forensic_software_frame = ttk.LabelFrame(self.tab_device, text="Forensic Software", padding="10")
        forensic_software_frame.pack(fill=tk.X, padx=5, pady=5)

        forensic_software_content = ttk.Frame(forensic_software_frame)
        forensic_software_content.pack(fill=tk.X, expand=True)

        # Forensic Processing Software - Using checkboxes all on same line
        ttk.Label(forensic_software_content, text="Forensic Processing Software:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(forensic_software_content)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variables - matching artifacts checkbox style exactly
        self.axiom_var = tk.IntVar(forensic_software_frame)
        self.xways_var = tk.IntVar(forensic_software_frame)
        self.griffeye_var = tk.IntVar(forensic_software_frame)
        
        # Create checkboxes exactly like artifacts checkboxes
        self.axiom_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Axiom",
            variable=self.axiom_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.axiom_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        self.xways_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="X-Ways",
            variable=self.xways_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.xways_checkbox.grid(row=0, column=1, sticky="w", padx=(0, 10))
        
        self.griffeye_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Griffeye",
            variable=self.griffeye_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.griffeye_checkbox.grid(row=0, column=2, sticky="w", padx=(0, 10))
        
        # Create artifacts button frame
        self.artifacts_button_frame = ttk.Frame(forensic_software_content)
        self.artifacts_button_frame.grid(row=1, column=0, columnspan=2, sticky="w", pady=5)

        self.selected_artifacts = []
        self.select_artifacts_button = ttk.Button(
            self.artifacts_button_frame, 
            text="Select Artifacts", 
            command=self.open_artifacts_popup
        )

        # Add the missing artifacts count label
        self.artifacts_count_label = ttk.Label(
            self.artifacts_button_frame,
            text="(0 selected)"
        )

        # Don't show artifacts button initially since nothing is selected by default
        # Both button and label will be shown/hidden together

        forensic_software_content.columnconfigure(1, weight=1)

    def on_forensic_software_change(self):
        # Get selected software
        selected_software = []
        if self.axiom_var.get():
            selected_software.append("Axiom")
        if self.xways_var.get():
            selected_software.append("X-Ways")
        if self.griffeye_var.get():
            selected_software.append("Griffeye")
        
        # Set the main forensic software variable (for backwards compatibility)
        if selected_software:
            self.forensic_software = type('obj', (object,), {'get': lambda: ', '.join(selected_software)})()
        else:
            self.forensic_software = type('obj', (object,), {'get': lambda: ''})()
        
        # Show/hide artifacts button based on Axiom selection
        if self.axiom_var.get():
            self.select_artifacts_button.pack(side=tk.LEFT, padx=5)
            if hasattr(self, 'artifacts_count_label'):
                self.artifacts_count_label.pack(side=tk.LEFT, padx=5)
                self.update_artifacts_count_label()
        else:
            self.select_artifacts_button.pack_forget()
            if hasattr(self, 'artifacts_count_label'):
                self.artifacts_count_label.pack_forget()
            # Clear artifacts when Axiom is deselected
            self.selected_artifacts = []


    def toggle_artifacts_button(self, event=None):
        # First check if the forensic software selected is Axiom
        if self.forensic_software.get() == 'Axiom':
            # Make sure the button and label are shown
            self.select_artifacts_button.pack(side=tk.LEFT, padx=5)
            self.artifacts_count_label.pack(side=tk.LEFT, padx=5)
            self.update_artifacts_count_label()
        else:
            # Hide the button and label
            self.select_artifacts_button.pack_forget()
            self.artifacts_count_label.pack_forget()
            # Clear any selected artifacts when switching away from Axiom
            self.selected_artifacts = []

    def update_artifacts_count_label(self):
        count = len(self.selected_artifacts) if hasattr(self, 'selected_artifacts') else 0
        self.artifacts_count_label.config(text=f"({count} selected)")

    def open_artifacts_popup(self):
        if hasattr(self, 'extraction_file') and self.extraction_file:
            # Always use PC for device type
            device_type = "PC" 
            self.selected_artifacts = self.select_artifacts_popup(device_type)
            # Update the count label instead of showing a popup
            self.update_artifacts_count_label()
        else:
            messagebox.showwarning("No Extraction File", "Please upload an extraction file first.")

    def create_device_info_frame(self):
        device_frame = ttk.LabelFrame(self.tab_device, text="Device Information", padding="10")
        device_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Owner
        ttk.Label(device_frame, text="Device Owner:").grid(row=0, column=0, sticky="w", pady=2)
        self.device_owner = ttk.Entry(device_frame)
        self.device_owner.grid(row=0, column=1, sticky="ew", pady=2)

        # Device Type with dropdown
        ttk.Label(device_frame, text="Device Type:").grid(row=1, column=0, sticky="w", pady=2)
        
        # Create a frame to contain the combobox and optional entry field
        device_type_frame = ttk.Frame(device_frame)
        device_type_frame.grid(row=1, column=1, sticky="ew", pady=2)
        device_type_frame.columnconfigure(0, weight=1)  # Make sure this frame expands properly
        
        self.device_type = ttk.Combobox(device_type_frame, 
            values=["Computer", "Loose Hard Drive", "USB Drive", "Memory Card", "Other Storage Device"], 
            state="readonly")
        self.device_type.pack(fill=tk.X, expand=True)
        self.device_type.set("Computer")
        self.device_type.bind("<<ComboboxSelected>>", self.toggle_device_type_fields)
        self.setup_auto_complete_dropdown(self.device_type)
        
        # Custom device type entry field (initially hidden)
        self.device_type_entry = ttk.Entry(device_type_frame)
        self.device_type_entry.pack(fill=tk.X, expand=True, pady=2)
        self.device_type_entry.pack_forget()  # Hide initially

        # Device Make
        ttk.Label(device_frame, text="Device Make:").grid(row=2, column=0, sticky="w", pady=2)
        self.device_PCMan = ttk.Entry(device_frame)
        self.device_PCMan.grid(row=2, column=1, sticky="ew", pady=2)

        # Device Model
        ttk.Label(device_frame, text="Device Model:").grid(row=3, column=0, sticky="w", pady=2)
        self.device_PCMod = ttk.Entry(device_frame)
        self.device_PCMod.grid(row=3, column=1, sticky="ew", pady=2)

        # Serial Number
        ttk.Label(device_frame, text="Device Serial Number:").grid(row=4, column=0, sticky="w", pady=2)
        self.device_PCSerial = ttk.Entry(device_frame)
        self.device_PCSerial.grid(row=4, column=1, sticky="ew", pady=2)

        # Device Color
        ttk.Label(device_frame, text="Device Color:").grid(row=5, column=0, sticky="w", pady=2)
        self.device_color = ttk.Entry(device_frame)
        self.device_color.grid(row=5, column=1, sticky="ew", pady=2)

        # Device Password
        ttk.Label(device_frame, text="Device Password:").grid(row=6, column=0, sticky="w", pady=2)
        self.device_password = ttk.Entry(device_frame)
        self.device_password.grid(row=6, column=1, sticky="ew", pady=2)

        # Computer-specific hard drive fields
        self.computer_fields_frame = ttk.Frame(device_frame)
        self.computer_fields_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)

        # Hard Drive Make
        ttk.Label(self.computer_fields_frame, text="Hard Drive Make:").grid(row=0, column=0, sticky="w", pady=2)
        self.hd_make = ttk.Entry(self.computer_fields_frame)
        self.hd_make.grid(row=0, column=1, sticky="ew", pady=2)

        # Hard Drive Model
        ttk.Label(self.computer_fields_frame, text="Hard Drive Model:").grid(row=1, column=0, sticky="w", pady=2)
        self.hd_model = ttk.Entry(self.computer_fields_frame)
        self.hd_model.grid(row=1, column=1, sticky="ew", pady=2)

        # Hard Drive Serial Number
        ttk.Label(self.computer_fields_frame, text="Hard Drive Serial:").grid(row=2, column=0, sticky="w", pady=2)
        self.hd_serial = ttk.Entry(self.computer_fields_frame)
        self.hd_serial.grid(row=2, column=1, sticky="ew", pady=2)

        # Configure grid weights for computer fields frame - THIS IS THE KEY FIX
        # Make the label column the same width as the main device_frame labels
        self.computer_fields_frame.columnconfigure(0, weight=0, minsize=120)  # Fixed width for labels
        self.computer_fields_frame.columnconfigure(1, weight=1)  # Entry fields expand

        # Storage Capacity - SINGLE FIELD that moves position based on device type
        self.capacity_label = ttk.Label(device_frame, text="Storage Capacity:")
        self.device_capacity = ttk.Entry(device_frame)

        # Configure grid weights for main device frame
        device_frame.columnconfigure(0, weight=0, minsize=120)  # Fixed width for labels to match
        device_frame.columnconfigure(1, weight=1)  # Entry fields expand
        
        # Show appropriate fields initially since "Computer" is the default
        self.toggle_device_type_fields()
        
        return device_frame
    
    def create_output_file_frame(self):
        output_frame = ttk.LabelFrame(self.tab_output, text="Output File", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        output_content = ttk.Frame(output_frame)
        output_content.pack(fill=tk.X, expand=True)
        
        # Output filename
        ttk.Label(output_content, text="Output File Name:").grid(row=0, column=0, sticky="w", pady=2)
        self.output_filename = ttk.Entry(output_content)
        self.output_filename.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Hint for filename (directly under the filename field)
        hint_style = ttk.Style()
        hint_style.configure("Hint.TLabel", font=('Arial', 7))  
        hint_label = ttk.Label(output_content, 
                              text="*Default File Name: \"(DFR #) - (Owner Name) (Device Model)\"", 
                              style="Hint.TLabel",
                              wraplength=350)
        hint_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 5))
        
        # Save location
        ttk.Label(output_content, text="Save Location:").grid(row=2, column=0, sticky="w", pady=2)
        
        # Frame for save location entry and browse button
        location_frame = ttk.Frame(output_content)
        location_frame.grid(row=2, column=1, sticky="ew", pady=2)
        location_frame.columnconfigure(0, weight=1)
        
        # Save location entry (initialize with desktop path)
        self.save_location = ttk.Entry(location_frame)
        self.save_location.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        # Set default to desktop
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        self.save_location.insert(0, desktop_path)
        
        # Browse button
        browse_button = ttk.Button(location_frame, text="Browse", command=self.browse_save_location)
        browse_button.grid(row=0, column=1)
        
        output_content.columnconfigure(1, weight=1)

    def browse_save_location(self):
        """Allow user to browse and select a save location"""
        try:
            # Get current path from the entry, default to desktop if empty
            current_path = self.save_location.get().strip()
            if not current_path or not os.path.exists(current_path):
                current_path = os.path.join(os.path.expanduser("~"), "Desktop")
            
            # Open folder dialog
            selected_path = filedialog.askdirectory(
                title="Select Save Location",
                initialdir=current_path
            )
            
            if selected_path:
                # Update the entry with selected path
                self.save_location.delete(0, tk.END)
                self.save_location.insert(0, selected_path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not browse for save location: {str(e)}")

    def create_right_column_widgets(self):
        self.create_file_upload_frame()

    def create_file_upload_frame(self):
        file_frame = ttk.LabelFrame(self.right_frame, text="File Upload", padding="10")
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        add_template_picker(self, file_frame, preferred="DFR Computer (2026).docx", keywords=("pc", "storage", "computer"))

        # Extraction File Drag-and-Drop
        self.create_extraction_file_ui(file_frame)    

    def toggle_device_type_fields(self, event=None):
        selected_type = self.device_type.get()
        
        if selected_type == "Other Storage Device":
            # Show custom entry field for other storage devices
            self.device_type_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            # Hide custom entry field
            self.device_type_entry.pack_forget()
        
        if selected_type == "Computer":
            # Show computer-specific hard drive fields
            self.computer_fields_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
            # Position capacity field after hard drive fields (row 8)
            self.capacity_label.grid(row=8, column=0, sticky="w", pady=2)
            self.device_capacity.grid(row=8, column=1, sticky="ew", pady=2)
        else:
            # Hide computer-specific fields and clear their values
            self.computer_fields_frame.grid_remove()
            if hasattr(self, 'hd_make'):
                self.hd_make.delete(0, tk.END)
            if hasattr(self, 'hd_model'):
                self.hd_model.delete(0, tk.END)
            if hasattr(self, 'hd_serial'):
                self.hd_serial.delete(0, tk.END)
            
            # Position capacity field after device password (row 7)
            self.capacity_label.grid(row=7, column=0, sticky="w", pady=2)
            self.device_capacity.grid(row=7, column=1, sticky="ew", pady=2)
    
    def create_extraction_file_ui(self, file_frame):
        # Extraction File Drag-and-Drop
        self.extraction_drop_label = ttk.Label(file_frame, text="Drag & Drop Log File or Click to Browse",
                                      relief="solid", borderwidth=2, padding=10)
        self.extraction_drop_label.pack(fill=tk.X, pady=10, padx=5)
        self.extraction_drop_label.bind('<Button-1>', self.browse_extraction_file)
        self.extraction_drop_label.drop_target_register(DND_FILES)
        self.extraction_drop_label.dnd_bind('<<Drop>>', self.drop_extraction_file)

    def create_info_display(self):
        self.info_frame, self.info_display = build_extracted_info_pane(
            self.right_frame,
            "Add a log file (.txt) to see the extracted information.",
        )
        self.no_evidence_var = tk.IntVar(self)
        self.no_evidence_checkbox = ttk.Checkbutton(
            self.info_frame,
            text="No Evidence Found",
            variable=self.no_evidence_var,
            onvalue=1,
            offvalue=0,
        )
        self.no_evidence_checkbox.pack(anchor="w", pady=(6, 0))

    def update_info_display(self, message=None, data=None):
        self.info_display.config(state=tk.NORMAL)
        self.info_display.delete(1.0, tk.END)
        
        if message:
            self.info_display.insert(tk.END, message)
        
        if data:
            # Add file type header based on extraction type
            if hasattr(self, "extraction_type"):
                if self.extraction_type == "TX1":
                    self.info_display.insert(tk.END, "TX1 Log File Detected\n", "header")
                elif self.extraction_type == "FTK":
                    self.info_display.insert(tk.END, "FTK Imager Log File Detected\n", "header")
                elif self.extraction_type == "XWAYS":
                    self.info_display.insert(tk.END, "X-Ways Forensics Log File Detected\n", "header")
                elif self.extraction_type == "DC":
                    self.info_display.insert(tk.END, "Cellebrite Digital Collector Log Detected\n", "header")
            
            self.info_display.insert(tk.END, "\nExtracted Information:\n\n", "section")
            
            display_data = data.copy()
            
            exclude_keys = ['article', 'extraction_date', 'examiner']
            
            for key in exclude_keys:
                if key in display_data:
                    del display_data[key]
            
            # Determine extraction type for display formatting
            is_tx1 = self.extraction_type == "TX1" if hasattr(self, "extraction_type") else False
            is_ftk = self.extraction_type == "FTK" if hasattr(self, "extraction_type") else False
            is_xways = self.extraction_type == "XWAYS" if hasattr(self, "extraction_type") else False
            is_dc = self.extraction_type == "DC" if hasattr(self, "extraction_type") else False
            
            # Define the display order based on extraction type
            if is_tx1:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('TX1_OS', 'TX1 Version'),
                    ('extraction_serial', 'TX1 Serial Number'),
                    ('case_number', 'Case Number'),
                    ('case_id', 'Case ID'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Case Notes'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_ftk:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('FTK_OS', 'FTK Imager Version'),
                    ('case_number', 'Case Number'),
                    ('evidence_number', 'Evidence Number'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Description'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_xways:
                display_order = [   
                    ('formatted_date', 'Extraction Date/Time'),
                    ('xways_OS', 'X-Ways Version'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Internal Description'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_dc:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('DC_OS', 'Digital Collector Version'),
                    ('case_number', 'Case Number'),
                    ('case_name', 'Case Name'),
                    ('evidence_number', 'Exhibit ID / Evidence #'),
                    ('device_model', 'Hard Drive Model'),
                    ('device_serial', 'Hard Drive Serial'),
                    ('device_capacity', 'Capacity'),
                    ('case_notes', 'Description'),
                    ('image_format', 'Image Format'),
                    ('md5_hash', 'MD5 Hash'),
                    ('sha1_hash', 'SHA1 Hash'),
                ]
            else:
                # Generic fallback display order
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('device_model', 'Source Device Model'),
                ]
            
            displayed_keys = set()
            
            for key_pair in display_order:
                key, display_name = key_pair
                if key in display_data and display_data[key] and key not in displayed_keys:
                    value = display_data[key]
                    self.info_display.insert(tk.END, f"{display_name}: ", "key")
                    self.info_display.insert(tk.END, f"{value}\n", "value")
                    displayed_keys.add(key)
             
            # Configure text tags for styling
            self.info_display.tag_configure("header", font=("Arial", 12, "bold"), foreground="green", justify="center")
            self.info_display.tag_configure("key", font=("Arial", 10, "bold"))
            self.info_display.tag_configure("value", font=("Arial", 10))
            self.info_display.tag_configure("section", font=("Arial", 11, "bold"))
        
        self.info_display.config(state=tk.DISABLED)

    def parse_extraction_file(self):
        if hasattr(self, 'extraction_type'):
            return self.parse_log_file()
        else:
            messagebox.showerror("Error", "No extraction file selected")
            return {}

    def toggle_request_content(self, event=None):
        # Clear all request information fields when role is changed
        if hasattr(self, 'request_date'):
            self.request_date.delete(0, tk.END)
        if hasattr(self, 'request_agency'):
            self.request_agency.delete(0, tk.END)
        if hasattr(self, 'request_title_type'):
            self.request_title_type.set("")  # Reset to default
        if hasattr(self, 'request_title_entry'):
            self.request_title_entry.delete(0, tk.END)
            self.request_title_entry.pack_forget()  # Hide custom entry
        if hasattr(self, 'request_officer'):
            self.request_officer.delete(0, tk.END)
        if hasattr(self, 'case_type'):
            self.case_type.delete(0, tk.END)
        if hasattr(self, 'legal_authority'):
            self.legal_authority.set('Search Warrant')  # Reset to default
        if hasattr(self, 'offense_type'):
            self.offense_type.delete(0, tk.END)
        if hasattr(self, 'legal_self'):
            self.legal_self.set('Search Warrant')  # Reset to default
        if hasattr(self, 'sw_service_date'):
            self.sw_service_date.delete(0, tk.END)
        
        # Clear time frame fields for both roles
        if hasattr(self, 'time_frame_var'):
            self.time_frame_var.set(0)
        if hasattr(self, 'time_frame_start'):
            self.time_frame_start.delete(0, tk.END)
        if hasattr(self, 'time_frame_end'):
            self.time_frame_end.delete(0, tk.END)
        if hasattr(self, 'case_agent_time_frame_var'):
            self.case_agent_time_frame_var.set(0)
        if hasattr(self, 'case_agent_time_frame_start'):
            self.case_agent_time_frame_start.delete(0, tk.END)
        if hasattr(self, 'case_agent_time_frame_end'):
            self.case_agent_time_frame_end.delete(0, tk.END)
        
        # Hide search warrant date fields when switching roles
        if hasattr(self, 'sw_date_frame'):
            for widget in self.sw_date_frame.winfo_children():
                widget.grid_remove()
        
        # Now show/hide the appropriate frame
        if self.role_type.get() == "Agency Assist":
            self.agency_assist_frame.pack(fill=tk.X, padx=5, pady=5)
            self.case_agent_frame.pack_forget()
            # Show time frame section since default is Search Warrant
            self.toggle_time_frame_section()
        else:  # Case Agent
            self.agency_assist_frame.pack_forget()
            self.case_agent_frame.pack(fill=tk.X, padx=5, pady=5)
            # Show the SW date field and time frame section since default is Search Warrant
            self.toggle_sw_date_and_time_frame()

    def toggle_sw_date(self, event=None):
        if hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant':
            # Show the SW date frame
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            # Hide the SW date frame and clear its value
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid_remove()
                
                # Clear the SW service date field when hiding it
                if hasattr(self, 'sw_service_date'):
                    self.sw_service_date.delete(0, tk.END)

    def toggle_title_entry(self, event=None):
        if self.request_title_type.get() == "Other (specify)":
            self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.request_title_entry.pack_forget()

    def toggle_agency_entry(self, event=None):
        return

    def toggle_transfer_fields(self, event=None):
        if self.device_transfer_var.get() == 1:
            self.transfer_fields.grid()
        else:
            self.transfer_fields.grid_remove()
            if hasattr(self, 'transfer_title'):
                self.transfer_title.delete(0, tk.END)
            if hasattr(self, 'transfer_officer'):
                self.transfer_officer.delete(0, tk.END)
            if hasattr(self, 'transfer_agency'):
                self.transfer_agency.delete(0, tk.END)
            if hasattr(self, 'transfer_date'):
                self.transfer_date.delete(0, tk.END)

    def setup_auto_complete_dropdown(self, combobox):
        combobox.bind('<KeyRelease>', lambda event, cb=combobox: self.auto_complete(event, cb))

    def auto_complete(self, event, combobox):
        if event.keysym in ('Up', 'Down', 'Left', 'Right', 'Return', 'Tab', 'BackSpace', 'Delete'):
            return
        
        value = combobox.get()
        
        if not value:
            return
        
        all_values = combobox['values']
        
        previous_value = combobox.get()
        
        for option in all_values:
            if option.lower().startswith(value.lower()):
                combobox.set(option)
                combobox.icursor(len(option))
                combobox.selection_range(len(value), len(option))
                           
                break

    def drop_extraction_file(self, event):
        file_path = self.tk.splitlist(event.data)[0]
        
        if file_path.lower().endswith('.txt'):
            # Clear previously populated fields from log files
            self.clear_log_populated_fields()
            
            # Set the new extraction file
            self.extraction_file = file_path
            self.extraction_drop_label.configure(text=f"Selected: {os.path.basename(file_path)} (Forensic Log)")
            
            # Parse the file and update the display
            extraction_data = self.parse_log_file()
            apply_log_device_fields_to_form(self, extraction_data)
            self.update_info_display(data=extraction_data)
            
            # No popup messages - the file type will be shown in the display window header
        else:
            messagebox.showerror("Error", "Please select a TXT log file (TX1, FTK, X-Ways, or Digital Collector)")
            self.update_info_display("Error: Please select a valid TXT log file.")

    def browse_extraction_file(self, event):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("Log files", "*.txt")
            ])
        
        if file_path:
            if file_path.lower().endswith('.txt'):
                # Clear previously populated fields from log files
                self.clear_log_populated_fields()
                
                # Set the new extraction file
                self.extraction_file = file_path
                self.extraction_drop_label.configure(text=f"Selected: {os.path.basename(file_path)} (Forensic Log)")
                
                # Parse the file and update the display
                extraction_data = self.parse_log_file()
                apply_log_device_fields_to_form(self, extraction_data)
                self.update_info_display(data=extraction_data)
                
                # No popup messages - the file type will be shown in the display window header
            else:
                messagebox.showerror("Error", "Please select a TXT log file")

    def clear_log_populated_fields(self):        
        # Clear any previously stored extraction data
        if hasattr(self, 'extraction_data'):
            delattr(self, 'extraction_data')
        
        # Clear extraction type
        if hasattr(self, 'extraction_type'):
            delattr(self, 'extraction_type')
        
        # Clear the set of populated fields
        if hasattr(self, 'log_populated_fields'):
            self.log_populated_fields.clear()
        
        # Reset extraction file reference if it exists
        if hasattr(self, 'extraction_file'):
            # Don't delete the file path, just clear any cached data
            pass
     
    def drop_template_file(self, event):
        file_path = self.tk.splitlist(event.data)[0]
        if file_path.endswith('.docx'):
            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        else:
            messagebox.showerror("Error", "Please select a Word document template")
     
    def browse_template_file(self, event):
        file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
        if file_path:
            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")

    def add_bold_underline_paragraph(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.underline = True
        run.font.name = 'Arial'  # This line already exists
        run.font.size = Pt(11)   # This line already exists
        return p

    def parse_log_file(self):
        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            messagebox.showerror("Error", "No log file selected")
            return {}
            
        try:
            with open(self.extraction_file, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Initialize/reset the set of fields that were populated from the log
                if not hasattr(self, 'log_populated_fields'):
                    self.log_populated_fields = set()
                else:
                    self.log_populated_fields.clear()
                
                # Determine the log type
                if "TX1 Log Entry" in content:
                    self.extraction_type = "TX1"
                    return self.parse_tx1_log(content)
                elif "FTK® Imager" in content or "FTK Imager" in content:
                    self.extraction_type = "FTK"
                    return self.parse_ftk_log(content)
                elif "X-Ways Forensics" in content:
                    self.extraction_type = "XWAYS"
                    return self.parse_xways_log(content)
                elif looks_like_digital_collector_log(content):
                    self.extraction_type = "DC"
                    data, populated = parse_digital_collector_log(content)
                    self.log_populated_fields.update(populated)
                    return data
                else:
                    messagebox.showwarning("Unknown Log Format", 
                                         "The selected file doesn't appear to be a TX1, FTK Imager, X-Ways, or Digital Collector log.")
                    return {}
        except Exception as e:
            messagebox.showerror("Error", f"Failed to parse log file: {str(e)}")
            return {}

    def parse_xways_log(self, content):
        extraction_data = {
            'extraction_tool': 'X-Ways Forensics',
            'extraction_type': 'Disk Imaging',  # Default value
        }
        
        lines = content.split('\n')
        
        # Extract acquisition start date from the first line
        if lines:
            first_line = lines[0].strip()
            # Remove BOM character if present
            first_line = first_line.lstrip('\ufeff')
            
            # Format: 12/13/2024, 10:19:06
            extraction_data['extraction_date'] = first_line  # Store original (cleaned)
            extraction_data['formatted_date'] = self.parse_xways_date(first_line)
            self.log_populated_fields.add('extraction_date')
            self.log_populated_fields.add('formatted_date')
        
        # Extract X-Ways version from second line
        if len(lines) > 1:
            second_line = lines[1].strip()
            # Format: X-Ways Forensics 20.9 SR-5 x64
            version_match = re.search(r'X-Ways Forensics\s+(\d+\.\d+)', second_line)
            if version_match:
                extraction_data['xways_OS'] = version_match.group(1)
                self.log_populated_fields.add('xways_OS')
        
        # Extract other fields line by line
        for line in lines:
            line = line.strip()
            
            # Extract Model
            if line.startswith('Model:'):
                model = line[len('Model:'):].strip()
                extraction_data['device_model'] = model
                self.log_populated_fields.add('device_model')
            
            # Extract Serial Number
            elif line.startswith('Serial No.:'):
                serial = line[len('Serial No.:'):].strip()
                # X-Ways shows serial like: 0000000000000000 / 0xecc4800670043
                # Take the part before the slash if there is one
                if '/' in serial:
                    serial = serial.split('/')[0].strip()
                extraction_data['device_serial'] = serial
                self.log_populated_fields.add('device_serial')
            
            # Extract Internal description (this could be used as case notes or additional device info)
            elif line.startswith('Internal description:'):
                description = line[len('Internal description:'):].strip()
                extraction_data['case_notes'] = description
                self.log_populated_fields.add('case_notes')
            
            # Extract MD5 hash
            elif line.startswith('Hash of source data:') and 'MD5' in line:
                # Format: Hash of source data: 640605156B461F053675C093173C7648 (MD5)
                hash_match = re.search(r'Hash of source data:\s*([a-fA-F0-9]+)\s*\(MD5\)', line)
                if hash_match:
                    extraction_data['md5_hash'] = hash_match.group(1)
                    self.log_populated_fields.add('md5_hash')
            
            # Extract Examiner(s) field
            elif line.startswith('Examiner(s):'):
                examiner = line[len('Examiner(s):'):].strip()
                extraction_data['examiner'] = examiner
                self.log_populated_fields.add('examiner')
        
        # If no formatted date was set, set a default
        if 'formatted_date' not in extraction_data:
            extraction_data['formatted_date'] = "Unknown Date"
            self.log_populated_fields.add('formatted_date')
        
        return extraction_data
    
    def parse_xways_date(self, date_str):
        try:
            # Remove BOM character if present
            date_str = date_str.lstrip('\ufeff')
            
            # Split by comma to get date and time parts
            if ', ' in date_str:
                date_part, time_part = date_str.split(', ')
            else:
                # If no comma, assume the whole string is just the date
                date_part = date_str.strip()
                time_part = "00:00:00"  # Default time
            
            # Parse the date and time using datetime
            full_datetime_str = f"{date_part.strip()} {time_part.strip()}"
            date_obj = datetime.strptime(full_datetime_str, '%m/%d/%Y %H:%M:%S')
            
            # Format the date for output 
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing X-Ways date: {e}")
            return date_str
    
    def parse_tx1_log(self, content):
        extraction_data = {
            'extraction_tool': 'TX1',
            'extraction_type': 'Disk Duplication',  # Default value
        }
        
        # Handle special parsing for Case ID and Case Notes (which may be indented on next line)
        # First, split the content into sections by looking for lines that start with hyphens
        sections = []
        current_section = []
        
        for line in content.split('\n'):
            if line.startswith('----'):  # Section separator
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
            current_section.append(line)
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        # Look for the section with Case ID and Case Notes
        for section in sections:
            if 'Case ID:' in section:
                # Extract Case ID - store it as case_id, not case_number
                case_id_match = re.search(r'Case ID:\s*\n?\s*(.*?)(?:\n\s*[A-Za-z]|\n\n)', section, re.DOTALL)
                if case_id_match:
                    case_id = case_id_match.group(1).strip()
                    extraction_data['case_id'] = case_id
                    # Track that this field was populated from the log
                    self.log_populated_fields.add('case_id')
                
                # Extract Case Notes (which typically follow Case ID and may be indented)
                case_notes_match = re.search(r'Case Notes:\s*\n?\s*(.*?)(?:\n\s*[A-Za-z]|\n\n|\n-----)', section, re.DOTALL)
                if case_notes_match:
                    # Clean up newlines and excessive whitespace in the notes
                    notes = case_notes_match.group(1).strip()
                    notes = re.sub(r'\n\s+', ' ', notes)  # Replace newline+whitespace with a single space
                    extraction_data['case_notes'] = notes
                    # Track that this field was populated from the log
                    self.log_populated_fields.add('case_notes')
                break
        
        # Find the Source Disk section specifically for device information
        source_disk_section = None
        for section in sections:
            if 'Source Disk' in section and '-----' in section:
                source_disk_section = section
                break
        
        # Process general fields (not device-specific) from the entire content
        general_patterns = {
            'Task:': ('extraction_type', lambda x: x.strip()),
            'Imager Ver:': ('TX1_OS', lambda x: x.strip()),
            'TX1 S/N:': ('extraction_serial', lambda x: x.strip()),
            'Created:': ('extraction_date', self.parse_tx1_date),
            'Examiner:': ('examiner', lambda x: x.strip()),
            'Case Number:': ('case_number', lambda x: x.strip()),
            'Acquisition Md5:': ('md5_hash', lambda x: x.replace(' ', '').strip()), 
        }
        
        # Parse general fields from entire content - DO NOT UPDATE ANY UI FIELDS
        for line in content.split('\n'):
            line = line.strip()
            for key, (field_name, processor) in general_patterns.items():
                if line.startswith(key):
                    value = line[len(key):].strip()
                    extraction_data[field_name] = processor(value)
                    
                    # Track that this field was populated from the log
                    self.log_populated_fields.add(field_name)
                    
                    # DO NOT UPDATE ANY UI FIELDS - just store in extraction_data
                    break
        
        # Process device-specific fields ONLY from Source Disk section
        if source_disk_section:
            device_patterns = {
                'Model:': ('device_model', lambda x: x.strip()),
                'Serial number:': ('device_serial', lambda x: x.strip()),
            }
            
            for line in source_disk_section.split('\n'):
                line = line.strip()
                for key, (field_name, processor) in device_patterns.items():
                    if line.startswith(key):
                        value = line[len(key):].strip()
                        extraction_data[field_name] = processor(value)
                        
                        # Track that this field was populated from the log
                        self.log_populated_fields.add(field_name)
                        
                        # DO NOT update device model or serial fields in the UI - only store in extraction_data
                        break
        
        # Generate formatted date in a readable format
        if 'extraction_date' in extraction_data:
            extraction_data['formatted_date'] = extraction_data['extraction_date']
        else:
            extraction_data['formatted_date'] = "Unknown Date"
                
        return extraction_data
    
    def parse_ftk_log(self, content):
        extraction_data = {
            'extraction_tool': 'AccessData FTK Imager',
            'extraction_type': 'Disk Imaging',  # Default value
        }
        
        # Extract FTK Imager version from first line
        version_match = re.search(r'FTK®?\s+Imager\s+(\d+\.\d+\.\d+\.\d+)', content)
        
        if version_match:
            extraction_data['FTK_OS'] = version_match.group(1)
            self.log_populated_fields.add('FTK_OS')
        
        # Extract fields based on patterns - all separate searches to handle different formats
        case_number_match = re.search(r'Case Number:\s*(.+?)(?:\r?\n)', content)
        if case_number_match:
            case_number = case_number_match.group(1).strip()
            extraction_data['case_number'] = case_number
            self.log_populated_fields.add('case_number')
            # DO NOT UPDATE UI FIELD - just store in extraction_data
        
        evidence_number_match = re.search(r'Evidence Number:\s*(.+?)(?:\r?\n)', content)
        if evidence_number_match:
            extraction_data['evidence_number'] = evidence_number_match.group(1).strip()
            self.log_populated_fields.add('evidence_number')
        
        description_match = re.search(r'Unique [Dd]escription:\s*(.+?)(?:\r?\n)', content)

        if description_match:
            extraction_data['case_notes'] = description_match.group(1).strip()
            self.log_populated_fields.add('case_notes')
        
        examiner_match = re.search(r'Examiner:\s*(.+?)(?:\r?\n)', content)
        if examiner_match:
            examiner = examiner_match.group(1).strip()
            extraction_data['examiner'] = examiner
            self.log_populated_fields.add('examiner')
        
        acquisition_date_match = re.search(r'Acquisition started:\s*(.+?)(?:\r?\n)', content)
        if acquisition_date_match:
            date_str = acquisition_date_match.group(1).strip()
            extraction_data['extraction_date'] = self.parse_ftk_date(date_str)
            extraction_data['formatted_date'] = extraction_data['extraction_date']
            self.log_populated_fields.add('extraction_date')
            self.log_populated_fields.add('formatted_date')
        
        drive_model_match = re.search(r'Drive Model:\s*(.+?)(?:\r?\n)', content)
        if drive_model_match:
            model = drive_model_match.group(1).strip()
            extraction_data['device_model'] = model
            self.log_populated_fields.add('device_model')
        
        drive_serial_match = re.search(r'Drive Serial Number:\s*(.+?)(?:\r?\n)', content)
        if drive_serial_match:
            serial = drive_serial_match.group(1).strip()
            extraction_data['device_serial'] = serial
            self.log_populated_fields.add('device_serial')
        
        md5_match = re.search(r'MD5 checksum:\s*([a-fA-F0-9]+)', content)
        if md5_match:
            extraction_data['md5_hash'] = md5_match.group(1).strip()
            self.log_populated_fields.add('md5_hash')

        # If no formatted date was set
        if 'formatted_date' not in extraction_data:
            extraction_data['formatted_date'] = "Unknown Date"
            self.log_populated_fields.add('formatted_date')
        
        return extraction_data
   
    def parse_tx1_date(self, date_str):
        try:
            # Format: Mon Apr  7 10:05:07 2025 (UTC-0500)
            # Remove timezone info in parentheses for parsing
            cleaned_date = re.sub(r'\s*\([^)]*\)', '', date_str).strip()
            
            # Parse the date using datetime
            date_obj = datetime.strptime(cleaned_date, '%a %b %d %H:%M:%S %Y')
            
            # Format the date for output
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing TX1 date: {e}")
            return date_str

    def parse_ftk_date(self, date_str):
        try:
            # Format: Thu Mar 27 10:10:41 2025
            date_obj = datetime.strptime(date_str, '%a %b %d %H:%M:%S %Y')
            
            # Format the date for output
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)  # Using 0 as offset since we're working with local time
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing FTK date: {e}")
            return date_str

    def get_local_timezone_name(self, utc_offset, is_dst):
        import time
        
        # Get the local timezone name directly from the system
        timezone_name = time.tzname[1] if is_dst else time.tzname[0]
        
        # Dictionary of timezone abbreviations to full names
        timezone_names = {
            "PST": "Pacific Standard Time",
            "PDT": "Pacific Daylight Time",
            "MST": "Mountain Standard Time",
            "MDT": "Mountain Daylight Time", 
            "CST": "Central Standard Time",
            "CDT": "Central Daylight Time",
            "EST": "Eastern Standard Time",
            "EDT": "Eastern Daylight Time",
            "HST": "Hawaii Standard Time",
            "AKST": "Alaska Standard Time",
            "AKDT": "Alaska Daylight Time",
            "AST": "Atlantic Standard Time",
            "ADT": "Atlantic Daylight Time"
        }
        
        return timezone_names.get(timezone_name, timezone_name)
        
    def browse_template_file(self, event):
        try:
            file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
            if file_path:
                with open(file_path, 'rb') as test_file:
                    pass
                self.template_file = file_path
                sync_template_choice(self, file_path)
                self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        except IOError as e:
            messagebox.showerror("File Error", f"Could not read the selected file: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

### PARAGRAPHS ###      
    def initialize_pc_exam_paragraphs(self):
        self.paragraphs = load_paragraphs("pc_full")

    def select_artifacts_popup(self, device_type):
        artifacts = []
        previously_selected = self.selected_artifacts if hasattr(self, 'selected_artifacts') else []
        
        left_artifacts = [
            "Operating System Information",
            "File System Information", 
            "Jump Lists",
            "Keyword Searches",
            "LNK Files",
            "MRU Opened/Saved Files",
            "MRU Recent Files And Folders",
            "Recycle Bin",
            "Installed Programs",
            "Microsoft Installed Programs",
            "USB Devices",
            "Encrypted Files",
            "Encryption / Anti-forensic Tools",
        ]
        
        middle_artifacts = [
            "Dropbox",
            "Google Drive",
            "OneDrive",
            "Torrent File Fragments",
            "EML(X) Files",
            "Gmail Fragments",
            "Gmail Webmail",
            "MBOX Emails",
            "Outlook Emails",
            "Windows Mail",
            "Google Maps",
        ]
        
        browser_artifacts = [
            "Chrome",
            "Edge",
            "Firefox",
            "Opera",
        ]
        
        right_artifacts = [
            "Google Docs",
            "Microsoft Excel Documents",
            "Microsoft Word Documents",
            "PDF Documents",
            "Text Documents",
            "Audio",
            "VLC Recently Played Files",
            "Pictures",
            "Videos",
        ]
        
        selection_window = tk.Toplevel(self)
        selection_window.title("Select Artifacts to Include")
        
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        popup_width = 1050
        popup_height = 600
        
        x_position = (screen_width - popup_width) // 2
        y_position = (screen_height - popup_height) // 2
        
        selection_window.geometry(f"{popup_width}x{popup_height}+{x_position}+{y_position}")
        
        selection_window.transient(self)
        selection_window.grab_set()
        
        main_frame = ttk.Frame(selection_window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        instruction_frame = ttk.Frame(main_frame)
        instruction_frame.pack(fill="x", pady=5)

        instruction_text = ttk.Label(
            instruction_frame,
            text="Select common artifacts tagged during examination:",
            wraplength=900,
            justify="left",
            font=("Arial", 10)
        )
        instruction_text.pack(anchor="w")

        separator = ttk.Separator(main_frame, orient="horizontal")
        separator.pack(fill="x", pady=5)
        
        canvas_frame = ttk.Frame(main_frame)
        canvas_frame.pack(fill="both", expand=True)
        
        canvas = tk.Canvas(canvas_frame)
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        columns_container = ttk.Frame(scrollable_frame)
        columns_container.pack(fill="both", expand=True, padx=10, pady=10)
        
        left_column = ttk.LabelFrame(columns_container, text="System Artifacts", padding=10)
        left_column.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
        
        middle_column = ttk.LabelFrame(columns_container, text="Cloud & Communication Artifacts", padding=10)
        middle_column.grid(row=0, column=1, sticky="nsew", padx=5, pady=5)
        
        browser_column = ttk.LabelFrame(columns_container, text="Browser Artifacts", padding=10)
        browser_column.grid(row=0, column=2, sticky="nsew", padx=5, pady=5)
        
        right_column = ttk.LabelFrame(columns_container, text="Media Artifacts", padding=10)
        right_column.grid(row=0, column=3, sticky="nsew", padx=5, pady=5)
        
        for i in range(4):
            columns_container.columnconfigure(i, weight=1, uniform="column")
        columns_container.rowconfigure(0, weight=1)
        
        artifact_vars = {}
        checkbox_widgets = {}
        subtag_frames = {}
        
        def create_artifact_checkboxes(column_frame, artifact_list):
            for artifact in artifact_list:
                artifact_frame = ttk.Frame(column_frame)
                artifact_frame.pack(anchor="w", fill="x", pady=2)
                
                var = tk.IntVar(selection_window)
                
                if artifact in previously_selected:
                    var.set(1)
                else:
                    var.set(0)
                
                cb = ttk.Checkbutton(
                    artifact_frame,
                    text=artifact,
                    variable=var,
                    onvalue=1,
                    offvalue=0
                )
                cb.pack(anchor="w", padx=5, pady=2)
                
                artifact_vars[artifact] = var
                checkbox_widgets[artifact] = cb
                
                if artifact in ["Pictures", "Videos", "Chrome", "Edge", "Firefox", "Opera"]:
                    subtag_container = ttk.Frame(artifact_frame)
                    subtag_frames[artifact] = subtag_container
                    
                    if artifact == "Pictures" or artifact == "Videos":
                        subtags = ["Child Pornography", "Child Erotica", "Age Difficult"]
                    elif artifact == "Chrome":
                        subtags = ["Autofill", "Bookmarks", "Current Tabs", "Downloads", "Keyword Search Terms", "Web History", "Web Visits"]
                    elif artifact == "Edge":
                        subtags = ["Autofill", "Bookmarks", "Current Tabs", "Downloads", "Keyword Search Terms", "Web History", "Web Visits"]
                    elif artifact == "Firefox":
                        subtags = ["Bookmarks", "Downloads", "Private Browsing History", "Web History", "Web Visits"]
                    elif artifact == "Opera":
                        subtags = ["Autofill", "Bookmarks", "Current Tabs", "Downloads", "Keyword Search Terms", "Web History", "Web Visits"]
                    
                    for subtag in subtags:
                        subtag_name = f"{artifact} -- {subtag}"
                        subvar = tk.IntVar(selection_window)
                        
                        if subtag_name in previously_selected:
                            subvar.set(1)
                        else:
                            subvar.set(0)
                        
                        sub_cb = ttk.Checkbutton(
                            subtag_container,
                            text=f"{subtag}",
                            variable=subvar,
                            onvalue=1,
                            offvalue=0
                        )
                        sub_cb.pack(anchor="w", padx=20, pady=1)
                        
                        artifact_vars[subtag_name] = subvar
                    
                    if artifact in previously_selected:
                        subtag_container.pack(anchor="w", fill="x")
                    else:
                        subtag_container.pack_forget()
                    
                    def make_toggle_function(artifact_name, container):
                        def toggle_func(*args):
                            if artifact_vars[artifact_name].get() == 1:
                                container.pack(anchor="w", fill="x")
                            else:
                                container.pack_forget()
                                for key, var in artifact_vars.items():
                                    if key.startswith(f"{artifact_name} -- "):
                                        var.set(0)
                        return toggle_func
                    
                    toggle_func = make_toggle_function(artifact, subtag_container)
                    var.trace("w", toggle_func)
        
        create_artifact_checkboxes(left_column, left_artifacts)
        create_artifact_checkboxes(middle_column, middle_artifacts)
        create_artifact_checkboxes(browser_column, browser_artifacts)
        create_artifact_checkboxes(right_column, right_artifacts)
        
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=10)
        
        def select_all():
            for key, var in artifact_vars.items():
                if not key.startswith("Pictures -- ") and not key.startswith("Videos -- ") and not key.startswith("Chrome -- ") and not key.startswith("Edge -- ") and not key.startswith("Firefox -- ") and not key.startswith("Opera -- "):
                    var.set(1)
            for artifact in ["Pictures", "Videos", "Chrome", "Edge", "Firefox", "Opera"]:
                if artifact in subtag_frames:
                    subtag_frames[artifact].pack(anchor="w", fill="x")
        
        def deselect_all():
            for key, var in artifact_vars.items():
                var.set(0)
            for artifact in ["Pictures", "Videos", "Chrome", "Edge", "Firefox", "Opera"]:
                if artifact in subtag_frames:
                    subtag_frames[artifact].pack_forget()
        
        select_all_button = ttk.Button(button_frame, text="Select All", command=select_all)
        select_all_button.pack(side="left", padx=5)
        
        deselect_all_button = ttk.Button(button_frame, text="Deselect All", command=deselect_all)
        deselect_all_button.pack(side="left", padx=5)
        
        def on_done():
            nonlocal artifacts
            artifacts = []  
            
            for artifact_name, var in artifact_vars.items():
                if var.get() == 1:
                    artifacts.append(artifact_name)
            
            selection_window.destroy()
        
        done_button = ttk.Button(button_frame, text="Done", command=on_done)
        done_button.pack(side="right", padx=5)
        
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        self.wait_window(selection_window)
        
        try:
            canvas.unbind_all("<MouseWheel>")
        except:
            pass
        
        return artifacts
    
    def generate_artifact_paragraphs(self, doc, selected_artifacts):
        # Mapping PC artifact names to their descriptions and tag names
        pc_artifacts = {
            "Operating System Information": {
                "description": "Operating System Information contains details about the Windows operating system installed on the computer, including version, build number, installation date, and system configuration settings.",
                "tag": "Operating System Information"
            },
            "File System Information": {
                "description": "Information pertaining to the File System that was searched.",
                "tag": "File System Information"
            },
            "Jump Lists": {
                "description": "Jump lists are quick lists of recent applications or files that a user launched.",
                "tag": "Jump Lists"
            },
            "Keyword Searches": {
                "description": "A list of keywords that were searched for on the system.",
                "tag": "Keyword Searches"
            },
            "LNK Files": {
                "description": "LNK files are Windows shortcut files that point to other files on the system.",
                "tag": "LNK Files"
            },
            "MRU Opened/Saved Files": {
                "description": "MRU Opened/Saved Files contains information about last files accessed by any application through 'Open File' or 'Save File' dialog window. Windows versions above XP use PIDL to store file path. PIDL paths might contain GUIDs instead of relative path strings.",
                "tag": "MRU Opened/Saved Files"
            },
            "MRU Recent Files And Folders": {
                "description": "The MRU Recent Files And Folders artifact contains information about files that were recently opened or saved and folders that were opened. This data is often related to items found in the Recent folder in the Users directory.",
                "tag": "MRU Recent Files And Folders"
            },
            "Recycle Bin": {
                "description": "Recycle Bin displays all items that were moved to the Recycle Bin.",
                "tag": "Recycle Bin"
            },
            "Microsoft Installed Programs": {
                "description": "Installed Microsoft Programs contains applications installed on the machine which are published by Microsoft.",
                "tag": "Microsoft Installed Programs"
            },
            "Installed Programs": {
                "description": "Installed Programs contain applications installed on the machine which are not published by Microsoft.",
                "tag": "Installed Programs"
            },
            "USB Devices": {
                "description": "USB Devices contains a history of all USB devices that have been connected to the system.",
                "tag": "USB Devices"
            },
            "Encrypted Files": {
                "description": "Encrypted Files contains information about any files that have been recovered on the system that are encrypted",
                "tag": "Encrypted Files"
            },
            "Encryption / Anti-forensic Tools": {
                "description": "Encryption/Anti-forensics Tools contains the encryption or anti-forensics tool(s) that have been found in the searched evidence.",
                "tag": "Encryption / Anti-forensic Tools"
            },
            "Dropbox": {
                "description": "Dropbox contains information about files that users uploaded and synced to Dropbox.",
                "tag": "Dropbox"
            },
            "Google Drive": {
                "description": "Google Drive is a file hosting service that allows users to upload and sync files to a cloud service.",
                "tag": "Google Drive"
            },
            "OneDrive": {
                "description": "These are artifacts left behind using OneDrive to upload and view files via the web or through the OneDrive desktop application. Data recovered can include file names, dates and times, user IDs, file sizes, sharing settings, and more.",
                "tag": "OneDrive"
            },
            "Torrent File Fragments": {
                "description": "Torrent File Fragments contains data that is carved or parsed from .torrent files that are used to download torrents from various networks on the Internet.",
                "tag": "Torrent File Fragments"
            },
            "EML(X) Files": {
                "description": "EML(X) Files contains the emails in .eml and .emlx formats, that have been found on the device.",
                "tag": "EML(X) Files"
            },
            "Gmail Fragments": {
                "description": "Gmail Email Fragments contains the Gmail email fragments that were recovered from a Windows or OS X computer.",
                "tag": "Gmail Fragments"
            },
            "Gmail Webmail": {
                "description": "Gmail is a webmail website that allows users to send and receive emails.",
                "tag": "Gmail Webmail"
            },
            "MBOX Emails": {
                "description": "MBOX is the default format used in Linux mail clients such as Thunderbird.",
                "tag": "MBOX Emails"
            },
            "Outlook Emails": {
                "description": "Microsft Outlook is a personal information manager and email client. This table captures information related to emails sent and received in Outlook.",
                "tag": "Outlook Emails"
            },
            "Windows Mail": {
                "description": "Windows Mail contains email messages sent or received using Windows Mail.",
                "tag": "Windows Mail"
            },
            "Google Maps": {
                "description": "Google Maps is a free web service that allows users to get directions.",
                "tag": "Google Maps"
            },
            "Chrome": {
                "description": "Google Chrome web browser artifacts containing various browsing data.",
                "tag": "Chrome"
            },
            "Chrome -- Autofill": {
                "description": "Chrome Autofill contains records of the autofill values that Chrome saves for different types of text fields.",
                "tag": "Chrome Autofill"
            },
            "Chrome -- Bookmarks": {
                "description": "Chrome Bookmarks contains browser bookmarks that reference saved webpages.",
                "tag": "Chrome Bookmarks"
            },
            "Chrome -- Current Tabs": {
                "description": "Chrome Current Tabs contains information about the tabs that are open in the current browser session.",
                "tag": "Chrome Current Tabs"
            },
            "Chrome -- Downloads": {
                "description": "Chrome Downloads contains information about the files that a user downloads from the Internet.",
                "tag": "Chrome Downloads"
            },
            "Chrome -- Keyword Search Terms": {
                "description": "Chrome Keyword Search Terms contains information about the keyword search terms that a user enters.",
                "tag": "Chrome Keyword Search Terms"
            },
            "Chrome -- Web History": {
                "description": "Chrome Web History contains a history of the websites that the user visits (includes unique visits only).",
                "tag": "Chrome Web History"
            },
            "Chrome -- Web Visits": {
                "description": "Chrome Web Visits contains a history of the websites that the user visits (includes all visits).",
                "tag": "Chrome Web Visits"
            },
            "Edge": {
                "description": "Microsoft Edge web browser artifacts containing various browsing data.",
                "tag": "Edge"
            },
            "Edge -- Autofill": {
                "description": "Edge Autofill contains records of the autofill values that Edge saves for different types of text fields.",
                "tag": "Edge Autofill"
            },
            "Edge -- Bookmarks": {
                "description": "Edge Bookmarks contains browser bookmarks that reference saved webpages.",
                "tag": "Edge Bookmarks"
            },
            "Edge -- Current Tabs": {
                "description": "Edge Current Tabs contains information about the tabs that are open in the current browser session.",
                "tag": "Edge Current Tabs"
            },
            "Edge -- Downloads": {
                "description": "Edge Downloads contains information about the files that a user downloads from the Internet.",
                "tag": "Edge Downloads"
            },
            "Edge -- Keyword Search Terms": {
                "description": "Edge Keyword Search Terms contains information about the keyword search terms that a user enters.",
                "tag": "Edge Keyword Search Terms"
            },
            "Edge -- Web History": {
                "description": "Edge Web History contains a history of the websites that the user visits (includes unique visits only).",
                "tag": "Edge Web History"
            },
            "Edge -- Web Visits": {
                "description": "Edge Web Visits contains a history of the websites that the user visits (includes all visits).",
                "tag": "Edge Web Visits"
            },
            "Firefox": {
                "description": "Mozilla Firefox web browser artifacts containing various browsing data.",
                "tag": "Firefox"
            },
            "Firefox -- Bookmarks": {
                "description": "Firefox Bookmarks contains the bookmarks from the Firefox web browser on a device.",
                "tag": "Firefox Bookmarks"
            },
            "Firefox -- Downloads": {
                "description": "Firefox Downloads contains the downloads from the Firefox web browser on a device.",
                "tag": "Firefox Downloads"
            },
            "Firefox -- Private Browsing History": {
                "description": "Firefox Private Browsing History contains the URLs that were loaded during a Private Browsing session from the Firefox web browser on a device.",
                "tag": "Firefox Private Browsing History"
            },
            "Firefox -- Web History": {
                "description": "Firefox Web History contains the webpages from the last active session from the Firefox web browser on a device.",
                "tag": "Firefox Web History"
            },
            "Firefox -- Web Visits": {
                "description": "Firefox Web Visits contains all of the non-archived URL visits for Firefox.",
                "tag": "Firefox Web Visits"
            },
            "Opera": {
                "description": "Opera web browser artifacts containing various browsing data.",
                "tag": "Opera"
            },
            "Opera -- Autofill": {
                "description": "Opera Autofill contains records of the autofill values that Opera saves for different types of text fields.",
                "tag": "Opera Autofill"
            },
            "Opera -- Bookmarks": {
                "description": "Opera is a web browser developed by Opera Software, and uses the Blink layout engine. Opera runs on Microsoft Windows and OS X operating systems.",
                "tag": "Opera Bookmarks"
            },
            "Opera -- Current Tabs": {
                "description": "Opera is a web browser developed by Opera Software, and Opera uses the Blink layout engine. Opera runs on Microsoft Windows and OS X operating systems.",
                "tag": "Opera Current Tabs"
            },
            "Opera -- Downloads": {
                "description": "Opera is a web browser developed by Opera Software, and uses the Blink layout engine. Opera runs on Microsoft Windows and OS X operating systems.",
                "tag": "Opera Downloads"
            },
            "Opera -- Keyword Search Terms": {
                "description": "Opera Keyword Search Terms contains information about the keyword search terms that a user entered.",
                "tag": "Opera Keyword Search Terms"
            },
            "Opera -- Web History": {
                "description": "Opera is a web browser developed by Opera Software. Web history are recently visited webpages. Opera stores a user's browsing history so that he or she can view it later. This search carves and parses web history from the Opera web browser, including the typed history (i.e. URLs or search terms entered by the user).",
                "tag": "Opera Web History"
            },
            "Opera -- Web Visits": {
                "description": "Opera Web Visits contains a history of the websites that the user visits (includes all visits).",
                "tag": "Opera Web Visits"
            },  
            "Google Docs": {
                "description": "Google Docs is a word processing suite available to all Google account holders.",
                "tag": "Google Docs"
            },
            "Microsoft Excel Documents": {
                "description": "Microsoft Excel is a spreadsheet processor developed by Microsoft.",
                "tag": "Microsoft Excel Documents"
            },
            "Microsoft Word Documents": {
                "description": "Microsoft Word is a word processor developed by Microsoft.",
                "tag": "Microsoft Word Documents"
            },
            "PDF Documents": {
                "description": "Portable Document Format (PDF) is a file format used to present documents in a manner independent of application software, hardware, and operating systems. This table captures documents in this file format, extracted from the filesystem and carved from unallocated space.",
                "tag": "PDF Documents"
            },
            "Text Documents": {
                "description": "Text documents (.txt) that are located on the system",
                "tag": "Text Documents"
            },
            "Audio": {
                "description": "Audio contains audio files that are recovered that use the .mp3 or .wav formats.",
                "tag": "Audio"
            },
            "VLC Recently Played Files": {
                "description": "VLC Recently Played Files contains information about the media files that are played using the VLC Media Player. This artifact can reveal information on the user's interaction with the application.",
                "tag": "VLC Recently Played Files"
            },
            "Pictures": {
                "description": "Pictures contains image files recovered from the computer storage, including photographs, graphics, and other visual media files in various formats.",
                "tag": "Pictures"
            },
            "Pictures -- Child Pornography": {
                "description": "The image files contained in this tag appear to depict child pornography. Child pornography consists of any visual depiction, including photographs, film, videos, or pictures depicting sexually explicit conduct. \"Sexually explicit conduct\" means material depicting any person under the age of 18 years engaged in graphic sexual intercourse, including genital-genital, oral-genital, anal-genital, or oral-anal whether between persons of the same or opposite sex, or lascivious simulated sexual intercourse where the genitals, breast or pubic area of any person is exhibited. In addition, any material depicting a minor involved in bestiality; masturbation; sadistic or masochistic abuse; or lascivious exhibition of the genitals or pubic area.",
                "tag": "Pictures -- Child Pornography"
            },
            "Pictures -- Child Erotica": {
                "description": "The image files contained in this tag do not meet the statutory requirement to be considered child pornography. These image files depict juvenile subjects who are shown wearing sexually suggestive clothing, posing in sexually suggestive positions, or that appear to be possessed for a sexual purpose. Image files of this nature are often referred to as \"child erotica\" by forensic examiners and investigators.",
                "tag": "Pictures -- Child Erotica"
            },
            "Pictures -- Age Difficult": {
                "description": "The image files in this tag are pornographic in nature and depict younger looking subjects who may be juveniles under the age of 18, or may be young adults who are 18 years of age or older. Therefore, without positive identification of the subjects shown in the image files in this tag, I cannot make an accurate determination regarding the legal or illegal nature of the image files. Pornographic image files and video files of this nature are often referred to as \"age difficult\" pornography by forensic examiners and investigators.",
                "tag": "Pictures -- Age Difficult"
            },
            "Videos": {
                "description": "Videos contains video files recovered from the computer storage, including movies, recordings, and other multimedia content in various video formats.",
                "tag": "Videos"
            },
            "Videos -- Child Pornography": {
                "description": "The video files contained in this tag appear to depict child pornography. Child pornography consists of any visual depiction, including photographs, film, videos, or pictures depicting sexually explicit conduct. \"Sexually explicit conduct\" means material depicting any person under the age of 18 years engaged in graphic sexual intercourse, including genital-genital, oral-genital, anal-genital, or oral-anal whether between persons of the same or opposite sex, or lascivious simulated sexual intercourse where the genitals, breast or pubic area of any person is exhibited. In addition, any material depicting a minor involved in bestiality; masturbation; sadistic or masochistic abuse; or lascivious exhibition of the genitals or pubic area.",
                "tag": "Videos -- Child Pornography"
            },
            "Videos -- Child Erotica": {
                "description": "The video files contained in this tag do not meet the statutory requirement to be considered child pornography. These video files depict juvenile subjects who are shown wearing sexually suggestive clothing, posing in sexually suggestive positions, or that appear to be possessed for a sexual purpose. Video files of this nature are often referred to as \"child erotica\" by forensic examiners and investigators.",
                "tag": "Videos -- Child Erotica"
            },
            "Videos -- Age Difficult": {
                "description": "The video files in this tag are pornographic in nature and depict younger looking subjects who may be juveniles under the age of 18, or may be young adults who are 18 years of age or older. Therefore, without positive identification of the subjects shown in the video files in this tag, I cannot make an accurate determination regarding the legal or illegal nature of the video files. Video files of this nature are often referred to as \"age difficult\" pornography by forensic examiners and investigators.",
                "tag": "Videos -- Age Difficult"
            },
        }
        
        predefined_order = [
            "Operating System Information",
            "File System Information", 
            "Jump Lists",
            "Keyword Searches",
            "LNK Files",
            "MRU Opened/Saved Files",
            "MRU Recent Files And Folders",
            "Recycle Bin",
            "Installed Programs",
            "Microsoft Installed Programs",
            "USB Devices",
            "Encrypted Files",
            "Encryption / Anti-forensic Tools",
            
            "Dropbox",
            "Google Drive",
            "OneDrive",
            "Torrent File Fragments",
            "EML(X) Files",
            "Gmail Fragments",
            "Gmail Webmail",
            "MBOX Emails",
            "Outlook Emails",
            "Windows Mail",
            "Google Maps",
            
            "Chrome",
            "Chrome -- Autofill",
            "Chrome -- Bookmarks",
            "Chrome -- Current Tabs",
            "Chrome -- Downloads",
            "Chrome -- Keyword Search Terms",
            "Chrome -- Web History",
            "Chrome -- Web Visits",
            "Edge",
            "Edge -- Autofill",
            "Edge -- Bookmarks",
            "Edge -- Current Tabs",
            "Edge -- Downloads",
            "Edge -- Keyword Search Terms",
            "Edge -- Web History",
            "Edge -- Web Visits",
            "Firefox",
            "Firefox -- Bookmarks",
            "Firefox -- Downloads",
            "Firefox -- Private Browsing History",
            "Firefox -- Web History",
            "Firefox -- Web Visits",
            "Opera",
            "Opera -- Autofill",
            "Opera -- Bookmarks",
            "Opera -- Current Tabs",
            "Opera -- Downloads",
            "Opera -- Keyword Search Terms",
            "Opera -- Web History",
            "Opera -- Web Visits",
            
            "Google Docs",
            "Microsoft Excel Documents",
            "Microsoft Word Documents",
            "PDF Documents",
            "Text Documents",
            "Audio",
            "VLC Recently Played Files",
            "Pictures",
            "Pictures -- Child Pornography",
            "Pictures -- Child Erotica",
            "Pictures -- Age Difficult",
            "Videos",
            "Videos -- Child Pornography",
            "Videos -- Child Erotica",
            "Videos -- Age Difficult",
        ]
        
        organized_artifacts = []
        for artifact in predefined_order:
            if artifact in selected_artifacts:
                organized_artifacts.append(artifact)
        
        from docx.oxml.shared import qn
        from docx.shared import Inches
        
        for artifact in organized_artifacts:
            if artifact in pc_artifacts:
                is_pictures_or_videos_subtag = (artifact.startswith("Pictures -- ") or 
                                               artifact.startswith("Videos -- "))
                
                is_other_subtag = (" -- " in artifact and not is_pictures_or_videos_subtag)
                
                is_main_category = " -- " not in artifact
                
                if is_main_category or is_other_subtag:
                    p = doc.add_paragraph()
                    
                    pPr = p._element.get_or_add_pPr()
                    numPr = pPr.get_or_add_numPr()
                    numPr.get_or_add_ilvl().val = 0
                    numPr.get_or_add_numId().val = 1
                    
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    
                    # Use the full artifact name for browser artifacts
                    if is_other_subtag and any(browser in artifact for browser in ["Chrome", "Edge", "Firefox", "Opera"]):
                        # For browser artifacts, use the full name but format it properly
                        display_name = artifact.replace(" -- ", " ").upper()
                    elif is_other_subtag:
                        # For non-browser subtags, use just the part after --
                        display_name = artifact.split(" -- ")[1].upper()
                    else:
                        # For main categories, use the full name
                        display_name = artifact.upper()
                    
                    run = p.add_run(display_name)
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    doc.add_paragraph()
                    
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(pc_artifacts[artifact]["description"])
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    doc.add_paragraph()
                    
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(f"Tag: {pc_artifacts[artifact]['tag']}")
                    run.font.underline = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                else:
                    # For Pictures/Videos subtags, keep existing logic
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(f"Tag: {pc_artifacts[artifact]['tag']}")
                    run.font.underline = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    doc.add_paragraph()
                    
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(pc_artifacts[artifact]["description"])
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                
                doc.add_paragraph()
                
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run("(DEVICE ARTIFACTS)\n")
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                
                doc.add_paragraph()

    def validate_fields(self):
        missing_fields = []

        examiner_agency = (self.examiner_agency_type.get() or "").strip()
        
        # Get examiner title value based on dropdown selection
        examiner_title = self.examiner_title_type.get()
        if not examiner_title.strip():
            missing_fields.append("Examiner Title")

        # Get device type (either from dropdown or custom entry)
        device_type = self.device_type_entry.get() if self.device_type.get() == "Other Storage Device" else self.device_type.get()

        # Base fields that are always required (removed optional fields)
        base_fields = {
            "Device Owner": self.device_owner.get(),
            "Device Type": device_type,
            "Device Make": self.device_PCMan.get(),
            # Removed: "Device Model", "Device Serial Number", "Storage Capacity"
            "Examiner Agency": examiner_agency,
            "Examiner Title": examiner_title,
            "Examiner Name": self.examiner_name.get(),
            "Case Number": self.case_number.get(),
            # Removed: "Evidence Number"
            "DFR Number": self.DFR_Num.get() if is_complete_dfr_number(self.DFR_Num.get()) else "",
        }

        # Check role-specific fields
        if self.role_type.get() == "Agency Assist":
            role_fields = {
                "Exam Request Date": self.request_date.get(),
                "Requesting Agency": self.request_agency.get(),
                "Requesting Officer": self.request_officer.get(),
                "Case Type": self.case_type.get(),  # Agency Assist uses case_type
            }
            
            if not self.get_request_title():
                missing_fields.append("Requesting Officer Title")
            
            # Check time frame fields for Agency Assist
            if (self.legal_authority.get() == 'Search Warrant' and 
                hasattr(self, 'time_frame_var') and self.time_frame_var.get() == 1):
                if not self.time_frame_start.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.time_frame_end.get().strip():
                    missing_fields.append("Time Frame End Date")
                    
        else:  # Case Agent
            role_fields = {
                "Case Type": self.offense_type.get(),  # Case Agent uses offense_type
            }
            
            if self.legal_self.get() == 'Search Warrant' and not self.sw_service_date.get().strip():
                missing_fields.append("Search Warrant Service Date")
            
            # Check time frame fields for Case Agent
            if (self.legal_self.get() == 'Search Warrant' and 
                hasattr(self, 'case_agent_time_frame_var') and self.case_agent_time_frame_var.get() == 1):
                if not self.case_agent_time_frame_start.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.case_agent_time_frame_end.get().strip():
                    missing_fields.append("Time Frame End Date")

        # Combine all fields
        fields = {**base_fields, **role_fields}

        # Check for custom entry fields
        
        if self.device_type.get() == "Other Storage Device" and not self.device_type_entry.get().strip():
            missing_fields.append("Custom Device Type")

        # Computer-specific fields are now optional - removed this section entirely
        # The hard drive fields (make, model, serial) are now optional

        # Check all fields
        for field_name, field_value in fields.items():
            if not field_value.strip():
                missing_fields.append(field_name)

        # Check transfer fields if transfer is Yes
        if self.device_transfer_var.get() == 1:
            if not self.transfer_title.get().strip():
                missing_fields.append("Transfer Officer Title")
            if not self.transfer_officer.get().strip():
                missing_fields.append("Transfer Officer Name")
            if not self.transfer_agency.get().strip():
                missing_fields.append("Transfer Officer Agency")
            if not self.transfer_date.get().strip():
                missing_fields.append("Transfer Date")

        # Check that at least one forensic software is selected
        selected_software = []
        if hasattr(self, 'axiom_var') and self.axiom_var.get():
            selected_software.append("Axiom")
        if hasattr(self, 'xways_var') and self.xways_var.get():
            selected_software.append("X-Ways")
        if hasattr(self, 'griffeye_var') and self.griffeye_var.get():
            selected_software.append("Griffeye")
        
        if not selected_software:
            missing_fields.append("At least one Forensic Software must be selected")

        # Check that files are selected
        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            missing_fields.append("Extraction File (TX1, FTK, X-Ways, or Digital Collector Log)")
        if not hasattr(self, 'template_file') or not self.template_file:
            missing_fields.append("Template File")
        if not hasattr(self, 'extraction_type') or not self.extraction_type:
            missing_fields.append("Extraction Type (TX1, FTK, X-Ways, or Digital Collector)")

        return missing_fields

    def set_document_default_font(self, doc):
        # Set the default character style to Arial
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def search_and_replace_content_controls_simple(self, doc, search_docs):
        """
        Enhanced version that handles multi-paragraph replacements for PY_TEXT and PY_ACQUIRE
        with proper newline handling
        """
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        # Get the document XML as bytes
        doc_xml_str = etree.tostring(doc._element, encoding='unicode')
        
        # Perform replacements directly in the XML string
        for search_string, replacement_doc in search_docs.items():
            if search_string in doc_xml_str:
                # Special handling for PY_TEXT and PY_ACQUIRE - multi-paragraph content
                if search_string in ["PY_TEXT", "PY_ACQUIRE"]:
                    # Build the complete replacement XML for all paragraphs
                    replacement_xml_parts = []
                    
                    for para in replacement_doc.paragraphs:
                        # Get the paragraph text
                        para_text = para.text if para.text else ""
                        
                        # For PY_ACQUIRE, don't split by newlines since we already handled that in generate_replacements
                        # Each paragraph already represents a line from the log file
                        text_parts = [para_text] if search_string == "PY_ACQUIRE" else para_text.split('\n')
                        
                        for text_part in text_parts:
                            # Create paragraph XML with proper formatting
                            para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            
                            # Add paragraph properties if needed (for alignment, spacing, etc.)
                            para_xml += '<w:pPr></w:pPr>'
                            
                            # Check if this entire paragraph should be bold/underline
                            # (for headers like "FORENSIC EXTRACTION")
                            is_header = False
                            for run in para.runs:
                                if run.font.bold and run.font.underline:
                                    is_header = True
                                    break
                            
                            if text_part or not text_parts or len(text_parts) == 1:  # Always add run for single empty paragraph
                                para_xml += '<w:r>'
                                para_xml += '<w:rPr>'
                                
                                # Apply formatting based on the original run formatting
                                if is_header or (para.runs and para.runs[0].font.bold):
                                    para_xml += '<w:b/>'
                                if is_header or (para.runs and para.runs[0].font.underline):
                                    para_xml += '<w:u w:val="single"/>'
                                if para.runs and para.runs[0].font.italic:
                                    para_xml += '<w:i/>'
                                
                                # Font settings
                                font_name = 'Arial'
                                font_size = 22  # 11pt * 2 (half-points)
                                if para.runs and para.runs[0].font.name:
                                    font_name = para.runs[0].font.name
                                if para.runs and para.runs[0].font.size:
                                    font_size = para.runs[0].font.size.pt * 2 if hasattr(para.runs[0].font.size, 'pt') else 22
                                
                                para_xml += f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
                                para_xml += f'<w:sz w:val="{int(font_size)}"/>'
                                
                                para_xml += '</w:rPr>'
                                
                                if text_part:  # Only add text if there's content
                                    # Add the text content (escape XML special characters)
                                    text_content = text_part
                                    text_content = text_content.replace('&', '&amp;')
                                    text_content = text_content.replace('<', '&lt;')
                                    text_content = text_content.replace('>', '&gt;')
                                    text_content = text_content.replace('"', '&quot;')
                                    text_content = text_content.replace("'", '&apos;')
                                    
                                    para_xml += f'<w:t xml:space="preserve">{text_content}</w:t>'
                                
                                para_xml += '</w:r>'
                            
                            para_xml += '</w:p>'
                            replacement_xml_parts.append(para_xml)
                    
                    # Join all paragraphs
                    replacement_text = ''.join(replacement_xml_parts)
                    
                else:
                    # Handle single paragraph replacements (existing logic)
                    if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                        replacement_text = replacement_doc.paragraphs[0].text
                    else:
                        replacement_text = ""
                    
                    # Escape XML special characters in replacement text
                    replacement_text = replacement_text.replace('&', '&amp;')
                    replacement_text = replacement_text.replace('<', '&lt;')
                    replacement_text = replacement_text.replace('>', '&gt;')
                    replacement_text = replacement_text.replace('"', '&quot;')
                    replacement_text = replacement_text.replace("'", '&apos;')
                
                # Replace in the XML string
                doc_xml_str = doc_xml_str.replace(search_string, replacement_text)
                replaced_strings[search_string] = True
                
                if search_string in ["PY_TEXT", "PY_ACQUIRE"]:
                    paragraph_count = len(replacement_doc.paragraphs) if replacement_doc.paragraphs else 0
                    print(f"Replaced '{search_string}' with formatted content ({paragraph_count} paragraphs)")
                else:
                    print(f"Replaced '{search_string}' with '{replacement_text[:50]}...'")
        
        # Parse the modified XML back into the document
        new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
        doc._element.clear()
        for child in new_element:
            doc._element.append(child)
        
        return replaced_strings
           
    def search_and_replace_split_placeholders(self, doc, search_docs):
        """
        Handle placeholders that are split across XML elements
        """
        import re  # Move the import to the top
        
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        try:
            # Get the document XML as string
            doc_xml_str = etree.tostring(doc._element, encoding='unicode')
            original_xml_str = doc_xml_str
            
            print("\nHandling split placeholders...")
            
            # Look for placeholders that might be split by XML tags
            target_placeholders = ['PY_DFR', 'PY_OWNER', 'PY_EXAMINER', 'PY_CASENUMBER', 'PY_EVIDENCE', 'PY_REQOFF', 'PY_PCSERIAL', 'PY_DEVMAKE', 'PY_TX1VER', 'PY_XWVER', 'PY_DCVER'
            ]
            
            for search_string in target_placeholders:
                if replaced_strings[search_string]:
                    continue
                    
                if search_string not in search_docs:
                    continue
                    
                replacement_doc = search_docs[search_string]
                if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                    replacement_text = replacement_doc.paragraphs[0].text
                else:
                    replacement_text = ""
                
                # Escape XML special characters
                replacement_text = replacement_text.replace('&', '&amp;')
                replacement_text = replacement_text.replace('<', '&lt;')
                replacement_text = replacement_text.replace('>', '&gt;')
                replacement_text = replacement_text.replace('"', '&quot;')
                replacement_text = replacement_text.replace("'", '&apos;')
                
                # Create pattern allowing XML tags between each character
                pattern_parts = []
                for char in search_string:
                    pattern_parts.append(re.escape(char))
                
                # Join with pattern that allows XML tags between characters
                pattern = r'(?:<[^>]*>)*'.join(pattern_parts)
                
                # Find and replace the split placeholder
                matches = list(re.finditer(pattern, doc_xml_str))
                
                if matches:
                    # Replace from last match to first to avoid position shifts
                    for match in reversed(matches):
                        start_pos = match.start()
                        end_pos = match.end()
                        doc_xml_str = doc_xml_str[:start_pos] + replacement_text + doc_xml_str[end_pos:]
                    
                    replaced_strings[search_string] = True
                    print(f"  ✓ Replaced split placeholder '{search_string}' ({len(matches)} occurrences)")
                else:
                    print(f"  ✗ Could not find split placeholder '{search_string}'")
            
            # Parse back the document if changes were made
            if doc_xml_str != original_xml_str:
                new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
                doc._element.clear()
                for child in new_element:
                    doc._element.append(child)
            
        except Exception as e:
            print(f"Split placeholder replacement failed: {e}")
        
        return replaced_strings

    def generate_report(self, preview_only=False):
        try:
            missing_fields = self.validate_fields()
            if missing_fields and not preview_only:
                messagebox.showerror(
                    "Missing Fields",
                    f"The following fields are missing:\n\n{', '.join(missing_fields)}\n\nPlease complete them before proceeding."
                )
                return

            # Handle officer name based on role
            examiner_title_value = self.examiner_title_type.get()

            if self.role_type.get() == "Agency Assist":
                officer_name = self.request_officer.get().strip()
                officer_last_name = officer_name.split()[-1] if officer_name and ' ' in officer_name else officer_name
            else:
                officer_name = self.examiner_name.get().strip()
                officer_last_name = officer_name.split()[-1] if officer_name and ' ' in officer_name else officer_name

            examiner_agency = self.get_examiner_agency()
            
            if self.role_type.get() == "Agency Assist":
                request_date = self.parse_request_date(self.request_date.get())
            elif self.role_type.get() == "Case Agent" and self.legal_self.get() == 'Search Warrant':
                request_date = self.parse_request_date(self.sw_service_date.get())
            else:
                request_date = ""
            
            # Handle agency formatting based on role
            examiner_agency_formatted, examiner_agency_abbr = self.format_agency(examiner_agency, return_abbreviation=True)
            if self.role_type.get() == "Agency Assist":
                request_agency_formatted, request_agency_abbr = self.format_agency(self.request_agency.get(), return_abbreviation=True)
            else:
                request_agency_formatted, request_agency_abbr = examiner_agency_formatted, examiner_agency_abbr

            examiner_agency_formatted, examiner_agency_abbr = self.format_agency(examiner_agency, return_abbreviation=True)

            # Get the case type based on role
            if self.role_type.get() == "Agency Assist":
                case_type = self.case_type.get()
            else:  # Case Agent
                case_type = self.offense_type.get()

            # Get the device type from the UI (either selected from dropdown or custom entry)
            device_type = self.device_type_entry.get() if self.device_type.get() == "Other Storage Device" else self.device_type.get()

            # Determine source_device value based on device type
            if device_type.lower() in ['computer', 'loose hard drive']:
                source_device = "source hard drive"
            else:  # USB Drive, Memory Card, Other Storage Device
                source_device = "source device"

            # Get examiner title value based on dropdown selection
            examiner_title = self.examiner_title_type.get()

            data = {
                'Request_Date': request_date,
                'Request_Agency': request_agency_formatted,
                'Request_Agency_Abbr': request_agency_abbr,
                'Request_Title': self.format_title(self.get_request_title()) if self.role_type.get() == "Agency Assist" else "",
                'Request_Officer': self.request_officer.get().title() if self.role_type.get() == "Agency Assist" else "",
                'Request_Officer_LastName': officer_last_name.title() if officer_last_name else '',
                'Request_Case': case_type,
                'Device_Owner': self.device_owner.get().title(),
                'Device_Type': device_type,
                'device_type': device_type,
                'Device_Model': self.device_PCMod.get().title(),
                'Device_Serial': self.device_PCSerial.get(),
                'Device_Capacity': self.device_capacity.get(),
                'device_capacity': self.device_capacity.get(),
                'Device_Color': self.device_color.get().title() if hasattr(self, 'device_color') else '',
                'Examiner_Agency': examiner_agency_formatted,
                'Examiner_Agency_Abbr': examiner_agency_abbr,
                'Examiner_Title': self.format_title(examiner_title),
                'Examiner_Name': self.examiner_name.get().title(),
                'Forensic_Software': self.get_selected_forensic_software(),
                'Case_Number': self.case_number.get(),
                'evidence_ID': self.evidence_number.get(),
                'DFR_Num': self.DFR_Num.get(),
                'source_device': source_device,
                'device_PCMan': self.device_PCMan.get(),
                'device_PCMod': self.device_PCMod.get(),
                'device_PCSerial': self.device_PCSerial.get(),
            }

            if self.device_type.get() == "Computer":
                data.update({
                    'device_password': self.device_password.get().strip() if hasattr(self, 'device_password') else '',
                    'hd_make': self.hd_make.get().title(),
                    'hd_model': self.hd_model.get().title(),
                    'hd_serial': self.hd_serial.get(),
                    'article_hd': 'an' if self.hd_make.get().lower().startswith(('a','e','i','o','u')) else 'a',
                })

            if self.role_type.get() == "Case Agent" and self.legal_self.get() == 'Search Warrant':
                data['SW_Date'] = self.parse_request_date(self.sw_service_date.get())

            if self.device_transfer_var.get() == 1:
                transfer_agency_formatted, transfer_agency_abbr = self.format_agency(self.transfer_agency.get(), return_abbreviation=True)
                data.update({
                    'Transfer_Date': self.parse_request_date(self.transfer_date.get()),
                    'Transfer_Title': self.format_title(self.transfer_title.get()),
                    'Transfer_Officer': self.transfer_officer.get().title(),
                    'Transfer_Agency': transfer_agency_formatted,
                    'Transfer_Agency_Abbr': transfer_agency_abbr,
                })

            # Get extraction data for forensic extraction section only
            try:
                extraction_data = self.parse_extraction_file()
            except Exception:
                extraction_data = {}
            
            # Only copy specific fields from extraction_data - NOT device info fields
            extraction_fields_to_copy = [
                'case_id',
                'extraction_type',
                'extraction_tool',
                'TX1_OS',
                'FTK_OS',
                'xways_OS',
                'DC_OS',
                'case_name',
                'sha1_hash',
                'image_format',
                'extraction_software',  
                'extraction_serial',
                'extraction_date',
                'formatted_date',
                'md5_hash',
                'evidence_number',
            ]
            
            # Copy only the allowed fields
            for field in extraction_fields_to_copy:
                if field in extraction_data:
                    data[field] = extraction_data[field]
            data = merge_log_device_into_report_data(data, extraction_data, device_type)
                    
            # Set article based on device type
            data['article'] = 'an' if device_type.lower().startswith(('a','e','i','o','u')) else 'a'
            if preview_only:
                model = data.get("device_PCMod") or data.get("hd_model") or ""
                suggested = apply_suggested_filename(self, "PC", model)
                show_placeholder_preview(
                    self,
                    pc_preview_rows(
                        data,
                        officer_text=self.format_case_officer(data) if hasattr(self, "format_case_officer") else data.get("Request_Officer", ""),
                        image_date=data.get("formatted_date", ""),
                    ),
                    suggested,
                )
                return
            
            # Load the template document
            doc = Document(self.template_file)
            self.set_document_default_font(doc)
            
            # METHOD 1: Try direct paragraph replacement first
            py_text_found = False
            for para_index, para in enumerate(doc.paragraphs):
                if 'PY_TEXT' in para.text:
                    py_text_found = True
                    print(f"Found PY_TEXT in paragraph {para_index}")
                    
                    # Clear the paragraph
                    para.clear()
                    
                    # Generate all the content paragraphs
                    temp_doc = Document()
                    self.generate_paragraphs(data, temp_doc)
                    
                    # Get the parent element
                    parent = para._element.getparent()
                    insert_index = parent.index(para._element)
                    
                    # Insert all generated paragraphs at this position
                    for new_para in reversed(temp_doc.paragraphs):
                        # Create a copy of the paragraph in the main document
                        copied_para = doc.add_paragraph()
                        
                        # Copy all runs with formatting
                        for run in new_para.runs:
                            new_run = copied_para.add_run(run.text)
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                            new_run.font.underline = run.font.underline
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                        
                        # Move to correct position
                        parent.insert(insert_index, copied_para._element)
                    
                    # Remove the original PY_TEXT paragraph
                    parent.remove(para._element)
                    break
            
            # Handle other replacements
            search_docs = {
                "PY_DFR": Document(),
                "PY_CASENUMBER": Document(),
                "PY_EVIDENCE": Document(),
                "PY_REQDATE": Document(),
                "PY_OWNER": Document(),
                "PY_REQAGENCY": Document(),
                "PY_REQOFF": Document(),
                "PY_LIMITSTART": Document(),
                "PY_LIMITEND": Document(),
                "PY_EXAMINER": Document(),
                "PY_IMAGEDATE": Document(),

                "PY_DEVMAKE": Document(),
                "PY_DEVMODEL": Document(),
                "PY_PCSERIAL": Document(),
                "PY_COLOR": Document(),
                "PY_PASSCODE": Document(),
                
                "PY_HDMAKE": Document(),
                "PY_HDMODEL": Document(),
                "PY_HDSERIAL": Document(),
                "PY_CAPACITY": Document(),        
                
                "PY_FTKVER": Document(),
                "PY_TX1VER": Document(),
                "PY_XWVER": Document(),
                "PY_DCVER": Document(),
                
                "PY_ACQUIRE": Document(),   
            }
            
            # Generate replacement content for other fields
            self.generate_replacements(data, search_docs)
            
            # METHOD 2: If PY_TEXT wasn't found in paragraphs, it might be in a content control
            if not py_text_found:
                print("PY_TEXT not found in paragraphs, trying XML method...")
                # Create a document with all the generated paragraphs for PY_TEXT
                new_doc = Document()
                self.generate_paragraphs(data, new_doc)
                search_docs["PY_TEXT"] = new_doc
            
            # Replace remaining placeholders using XML method
            try:
                replaced_strings = self.search_and_replace_content_controls_simple(doc, search_docs)
                
                # Check for any remaining missing placeholders
                missing_strings = [s for s, replaced in replaced_strings.items() if not replaced]
                
                if missing_strings:
                    print(f"Trying split placeholder method for: {missing_strings}")
                    split_replaced = self.search_and_replace_split_placeholders(doc, search_docs)
                    
                    # Update the replaced_strings with any newly found ones
                    for search_string, was_replaced in split_replaced.items():
                        if was_replaced:
                            replaced_strings[search_string] = True
                
                # Check for missing strings and show appropriate warnings
                final_missing = [s for s, replaced in replaced_strings.items() if not replaced]
                
                if final_missing:
                    if "PY_TEXT" in final_missing:
                        messagebox.showwarning("Warning", "The string 'PY_TEXT' was not found in the document.")
                    else:
                        messagebox.showwarning("Warning", f"The following strings were not found in the document: {', '.join(final_missing)}")
                    
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"Replacement method failed: {e}")
                print(f"Error details: {error_details}")

            self.save_document(doc)

            # Reset selected artifacts after generating report
            self.selected_artifacts = []
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            messagebox.showerror("Error", f"Failed to generate report:\n\n{str(e)}\n\nDetails:\n{error_details[:500]}")
    
    def get_selected_forensic_software(self):
        selected = []
        if hasattr(self, 'axiom_var') and self.axiom_var.get():
            selected.append("Axiom")
        if hasattr(self, 'xways_var') and self.xways_var.get():
            selected.append("X-Ways")
        if hasattr(self, 'griffeye_var') and self.griffeye_var.get():
            selected.append("Griffeye")

        return ', '.join(selected) if selected else ''

    def format_request_agency_full(self, data):
        """Returns the full agency name (not abbreviated) for PY_REQAGENCY"""
        # For Case Agent role, use examiner agency (full name)
        if self.role_type.get() == "Case Agent":
            return data.get('Examiner_Agency', '')
        else:  # Agency Assist role
            return data.get('Request_Agency', '')

    def generate_replacements(self, data, search_docs):
        def add_text_to_doc(doc, text):
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p
        
        def add_formatted_log_to_doc(doc, log_content):
            """Special handler for log content that preserves formatting"""
            if not log_content:
                return add_text_to_doc(doc, "")
            
            # Split the log content into lines
            lines = log_content.split('\n')
            
            # Add each line as a separate paragraph to preserve line breaks
            for i, line in enumerate(lines):
                # For empty lines, add an empty paragraph to preserve spacing
                if not line.strip():
                    p = doc.add_paragraph()
                else:
                    p = doc.add_paragraph(line)
                    
                # Apply consistent formatting
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        
        def format_date_for_image(extraction_data):                
            # First try to parse the formatted_date that we create
            if 'formatted_date' in extraction_data and extraction_data['formatted_date']:
                formatted_date_str = extraction_data['formatted_date']
                
                try:
                    import re
                    # Match pattern: "DayName, Month DD, YYYY at HH:MM TimeZone"
                    date_match = re.search(r'(\w+),\s+(\w+)\s+(\d+),\s+(\d+)', formatted_date_str)
                    if date_match:
                        day_name = date_match.group(1)
                        month_name = date_match.group(2)
                        day = int(date_match.group(3))
                        year = int(date_match.group(4))
                        
                        # Format the same way as request date: "Monday, April 7, 2025"
                        formatted_date = f"{day_name}, {month_name} {day}, {year}"
                        # Remove leading zeros from day (same as parse_request_date does)
                        formatted_date = formatted_date.replace(' 0', ' ')
                        return formatted_date
                                
                except Exception as e:
                    print(f"Error parsing formatted_date: {e}")
            
            # Fallback: try to use the raw extraction_date and format it the same way
            if 'extraction_date' not in extraction_data:
                return ""
            
            try:
                # Get the raw extraction date from the log (this is the original format)
                date_str = extraction_data['extraction_date']
                
                # Handle TX1 date format: Mon Apr  7 10:05:07 2025 (UTC-0500)
                if hasattr(self, 'extraction_type') and self.extraction_type == "TX1":
                    # Remove timezone info in parentheses for parsing
                    cleaned_date = re.sub(r'\s*\([^)]*\)', '', date_str).strip()
                    date_obj = datetime.strptime(cleaned_date, '%a %b %d %H:%M:%S %Y')
                    # Format the same way as request date: "Monday, April 7, 2025"
                    formatted_date = date_obj.strftime("%A, %B %d, %Y")
                    formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
                    return formatted_date
                
                # Handle FTK date format: Thu Mar 27 10:10:41 2025
                elif hasattr(self, 'extraction_type') and self.extraction_type == "FTK":
                    date_obj = datetime.strptime(date_str, '%a %b %d %H:%M:%S %Y')
                    # Format the same way as request date: "Monday, April 7, 2025"
                    formatted_date = date_obj.strftime("%A, %B %d, %Y")
                    formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
                    return formatted_date
                
                return ""
                
            except Exception as e:
                print(f"Error formatting raw date for PY_IMAGEDATE: {e}")
                return ""
        
        # Get the raw log content for PY_ACQUIRE
        log_content = ""
        if hasattr(self, 'extraction_file') and self.extraction_file:
            try:
                with open(self.extraction_file, 'r', encoding='utf-8', errors='ignore') as file:
                    log_content = file.read()
            except Exception as e:
                print(f"Error reading log file for PY_ACQUIRE: {e}")
                log_content = ""
        
         # Get extraction data for date formatting
        extraction_data = self.parse_extraction_file() if hasattr(self, 'parse_extraction_file') else {}
        
        # Format the date as mm/dd/yyyy
        image_date = format_date_for_image(extraction_data)
        
        # Get time frame dates based on role
        time_frame_start = ""
        time_frame_end = ""
        
        if self.role_type.get() == "Agency Assist":
            if (self.legal_authority.get() == 'Search Warrant' and 
                hasattr(self, 'time_frame_var') and self.time_frame_var.get() == 1):
                try:
                    if hasattr(self, 'time_frame_start') and self.time_frame_start.get().strip():
                        time_frame_start = self.parse_request_date(self.time_frame_start.get())
                    if hasattr(self, 'time_frame_end') and self.time_frame_end.get().strip():
                        time_frame_end = self.parse_request_date(self.time_frame_end.get())
                except Exception as e:
                    print(f"Error parsing time frame dates for Agency Assist: {e}")
        else:  # Case Agent
            if (self.legal_self.get() == 'Search Warrant' and 
                hasattr(self, 'case_agent_time_frame_var') and self.case_agent_time_frame_var.get() == 1):
                try:
                    if hasattr(self, 'case_agent_time_frame_start') and self.case_agent_time_frame_start.get().strip():
                        time_frame_start = self.parse_request_date(self.case_agent_time_frame_start.get())
                    if hasattr(self, 'case_agent_time_frame_end') and self.case_agent_time_frame_end.get().strip():
                        time_frame_end = self.parse_request_date(self.case_agent_time_frame_end.get())
                except Exception as e:
                    print(f"Error parsing time frame dates for Case Agent: {e}")
        
        # Map search strings to data values
        replacement_map = {
            "PY_DFR": data.get('DFR_Num', ''),
            "PY_CASENUMBER": data.get('Case_Number', ''),
            "PY_EVIDENCE": data.get('evidence_ID', ''),
            "PY_REQDATE": data.get('Request_Date', ''),
            "PY_OWNER": data.get('Device_Owner', ''),
            "PY_REQAGENCY": self.format_request_agency_full(data),  # Changed this line
            "PY_REQOFF": self.format_case_officer(data),
            "PY_LIMITSTART": time_frame_start,
            "PY_LIMITEND": time_frame_end,
            "PY_EXAMINER": f"{data.get('Examiner_Title', '')} {data.get('Examiner_Name', '')}".strip(),
            "PY_IMAGEDATE": image_date,

            "PY_DEVMAKE": data.get('device_PCMan', ''),
            "PY_DEVMODEL": data.get('device_PCMod', ''),
            "PY_PCSERIAL": data.get('device_PCSerial', ''),
            "PY_COLOR": data.get('Device_Color', ''),
            "PY_PASSCODE": data.get('device_password', ''),
            
            "PY_HDMAKE": data.get('hd_make', ''),
            "PY_HDMODEL": data.get('hd_model', ''),
            "PY_HDSERIAL": data.get('hd_serial', ''),
            "PY_CAPACITY": data.get('Device_Capacity', ''),
            
            "PY_FTKVER": data.get('FTK_OS', ''),
            "PY_TX1VER": data.get('TX1_OS', ''),
            "PY_XWVER": data.get('xways_OS', ''),
            "PY_DCVER": data.get('DC_OS', ''),
        }

        # Only add PC-specific replacements for non-computer devices
        if self.device_type.get() != "Computer":
            replacement_map.update({
                "PY_PCMAN": data.get('device_PCMan', ''),
                "PY_PCMOD": data.get('device_PCMod', ''),
            })

        # Only add hard drive replacements for computer devices
        if self.device_type.get() == "Computer":
            replacement_map.update({
                "PY_HDMAKE": data.get('hd_make', ''),
                "PY_HDMODEL": data.get('hd_model', ''),
                "PY_HDSERIAL": data.get('hd_serial', ''),
            })
        
        print("\nGenerating replacement content:")
        
        # Create content for each search string
        for search_string, doc in search_docs.items():
            if search_string == "PY_TEXT":
                # These are handled by generate_paragraphs
                para_count = len(doc.paragraphs) if doc.paragraphs else 0
                print(f"  {search_string}: {para_count} paragraphs already generated")
                continue 
            
            # Special handling for PY_ACQUIRE to preserve formatting
            if search_string == "PY_ACQUIRE":
                print(f"  {search_string}: preserving log file formatting ({len(log_content)} characters)")
                add_formatted_log_to_doc(doc, log_content)
                continue
            
            value = replacement_map.get(search_string, '')
            
            # If no value exists, add a space to effectively erase the placeholder
            if not value:
                value = ''
                print(f"  {search_string}: (empty)")
            else:
                print(f"  {search_string}: '{value[:50]}...'")
            
            add_text_to_doc(doc, value)
    
    def format_time_frame_date(self, date_string):
        """
        Format time frame dates using the same format as other dates in the application
        """
        if not date_string or not date_string.strip():
            return ""
        
        try:
            # Use the existing parse_request_date method to format consistently
            return self.parse_request_date(date_string.strip())
        except ValueError as e:
            # If parsing fails, return the original string
            print(f"Warning: Could not format time frame date '{date_string}': {e}")
            return date_string
    
    def format_software_name(self, software_name):
        if not software_name:
            return ''
        
        # If the software name contains "AccessData" and "FTK", simplify it to just "FTK Imager"
        if 'AccessData' in software_name and 'FTK' in software_name:
            return 'FTK Imager'
        
        # If the software name contains "X-Ways", simplify it to just "X-Ways Forensics"
        if 'X-Ways' in software_name:
            return 'X-Ways Forensics'
        
        return software_name

    def format_case_officer(self, data):
        parts = [] 
        # For Case Agent role, use examiner information
        if self.role_type.get() == "Case Agent":
            if data.get('Examiner_Title'):
                parts.append(data['Examiner_Title'])
            if data.get('Examiner_Name'):
                parts.append(data['Examiner_Name'])
        else:  # Agency Assist role           
            if data.get('Request_Title'):
                abbreviated_title = self.get_title_abbreviation(data['Request_Title'])
                parts.append(abbreviated_title)
            
            if data.get('Request_Officer'):
                parts.append(data['Request_Officer'])

        return ' '.join(parts) if parts else ''

    def format_request_agency(self, data):
        # For Case Agent role, use examiner agency
        if self.role_type.get() == "Case Agent":
            return data.get('Examiner_Agency_Abbr', '')
        else:  # Agency Assist role
            return data.get('Request_Agency_Abbr', '')

    def get_title_abbreviation(self, full_title):
        title_abbreviations = {
            "officer": "Off.",
            "deputy": "Dep.",
            "special agent": "SA",
            "supervisory special agent": "SSA",
            "special assistant attorney general": "SAAG",
            "saag": "SAAG",
            "intel analyst": "IA",
            "detective": "Det.",
            "investigator": "Inv.",
            "trooper": "Trooper",
            "computer forensic examiner": "CFE",
        }
        
        # Convert to lowercase for comparison
        title_lower = full_title.lower().strip()
        
        # Return the abbreviation if found, otherwise return the original title
        return title_abbreviations.get(title_lower, full_title)      

    def get_examiner_agency(self):
        agency = (self.examiner_agency_type.get() or "").strip()
        if agency == "South Dakota DCI":
            return "South Dakota Division of Criminal Investigation"
        return agency

    def get_request_title(self):
        return (self.request_title_type.get() or "").strip()

    def generate_paragraphs(self, data, new_doc):
        def add_paragraph_with_style(doc, text):
            text = fill_paragraph(text, data)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p

        # Determine which paragraph to use based on device type and role
        device_type = data.get('Device_Type', '').lower()
        
        # Select the appropriate paragraph based on device type and transfer status
        if device_type == 'computer':
            if self.role_type.get() == "Case Agent":
                paragraph_key = 'one_a_computer'
            elif self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_computer'
            else:
                paragraph_key = 'one_c_computer'
        elif device_type == 'loose hard drive':
            if self.role_type.get() == "Case Agent":
                paragraph_key = 'one_a_loose'
            elif self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_loose'
            else:
                paragraph_key = 'one_c_loose'
        else:  # USB Drive, Memory Card, Other Storage Device
            if self.role_type.get() == "Case Agent":
                paragraph_key = 'one_a_storage'
            elif self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_storage'
            else:
                paragraph_key = 'one_c_storage'

        # Add the selected paragraph
        add_paragraph_with_style(new_doc, self.paragraphs[paragraph_key])

        # Add Authority Paragraph (only if not empty)
        authority_paragraph = self.get_authority_paragraph()
        if authority_paragraph:  # Only add if not empty
            add_paragraph_with_style(new_doc, authority_paragraph)

        # Add Paragraph Three
        if self.role_type.get() == "Case Agent":
            add_paragraph_with_style(new_doc, self.paragraphs['three_b'])
        else:  # Agency Assist
            add_paragraph_with_style(new_doc, self.paragraphs['three_a'])

        # Add Paragraph Four - Forensic Extraction Header
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['four'])

        # Add Paragraph Five - based on extraction type
        if hasattr(self, 'extraction_type'):
            if self.extraction_type == "TX1":
                add_paragraph_with_style(new_doc, self.paragraphs['five_tx1'])
            elif self.extraction_type == "FTK":
                add_paragraph_with_style(new_doc, self.paragraphs['five_ftk'])
            elif self.extraction_type == "XWAYS":
                add_paragraph_with_style(new_doc, self.paragraphs['five_xways'])
            elif self.extraction_type == "DC":
                add_paragraph_with_style(new_doc, self.paragraphs.get('five_dc', self.paragraphs['five_ftk']))

        # Add Paragraph Six
        add_paragraph_with_style(new_doc, self.paragraphs['six'])

        # Add Paragraph Seven - Forensic Processing Header
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['seven'])

        # Add Paragraph Eight - Processing paragraph
        add_paragraph_with_style(new_doc, self.paragraphs['eight'])

        # Add Griffeye paragraph if Griffeye is selected
        if hasattr(self, 'griffeye_var') and self.griffeye_var.get() == 1:
            add_paragraph_with_style(new_doc, self.paragraphs['griffeye_para'])

        # Check if "No Evidence Found" is checked
        if self.no_evidence_var.get() == 1:
            # Add FINDINGS OF FORENSIC EXAMINATION Header
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['nine'])
            
            # Add No Evidence paragraph instead of AXIOM report
            add_paragraph_with_style(new_doc, self.paragraphs['Paragraph_NoEv'])
        else:
            # Add remaining paragraphs for normal flow
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['nine'])
            add_paragraph_with_style(new_doc, self.paragraphs['ten'])
            add_paragraph_with_style(new_doc, self.paragraphs['eleventeen'])

            # Add AXIOM Digital Report section if Axiom is selected
            if hasattr(self, 'axiom_var') and self.axiom_var.get() == 1:
                add_paragraph_with_style(new_doc, self.paragraphs['twelveteen'])
                
                # Add AXIOM Digital Report Header
                self.add_bold_underline_paragraph(new_doc, self.paragraphs['thirteen'])
                
                # Add standard AXIOM paragraphs
                add_paragraph_with_style(new_doc, self.paragraphs['fourteen'])
                add_paragraph_with_style(new_doc, self.paragraphs['fifteen'])
                add_paragraph_with_style(new_doc, self.paragraphs['sixteen'])
                
                # Add Artifacts Header
                self.add_bold_underline_paragraph(new_doc, self.paragraphs['seventeen'])
                
                # Add artifacts if selected
                if hasattr(self, 'selected_artifacts') and self.selected_artifacts:
                    self.generate_artifact_paragraphs(new_doc, self.selected_artifacts)
                else:
                    # If no artifacts were selected, add a message
                    p = new_doc.add_paragraph()
                    p.style = new_doc.styles['Normal']
                    run = p.add_run("No artifacts were selected.")
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.italic = True
    
    def save_document(self, doc):
        output_filename = self.output_filename.get().strip()
        
        # If no filename provided, generate default filename
        if not output_filename:
            # Get DFR Number
            dfr_number = self.DFR_Num.get().strip()
            
            # Get Device Owner (apply title case)
            device_owner = self.device_owner.get().strip().title()
            
            # Get Device Model from PC device fields (apply title case)
            device_model = self.device_PCMod.get().strip().title()
            
            # Build the filename: "(DFR Number) - (Device Owner) (Device Model)"
            filename_parts = []
            
            if dfr_number:
                filename_parts.append(dfr_number)
            
            if device_owner:
                if device_model:
                    filename_parts.append(f"{device_owner} {device_model}")
                else:
                    filename_parts.append(device_owner)
            elif device_model:
                filename_parts.append(device_model)
            
            if filename_parts:
                output_filename = " - ".join(filename_parts)
            else:
                # Fallback to template filename if no data available
                template_filename = os.path.basename(self.template_file)
                output_filename = os.path.splitext(template_filename)[0]
        
        # Remove invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            output_filename = output_filename.replace(char, '')
        
        # Ensure .docx extension
        if not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        # Get save location from the entry field
        save_location = self.save_location.get().strip()
        
        # Validate save location exists, fallback to desktop if not
        if not save_location or not os.path.exists(save_location):
            save_location = os.path.join(os.path.expanduser("~"), "Desktop")
            # Update the entry to show the fallback location
            self.save_location.delete(0, tk.END)
            self.save_location.insert(0, save_location)
            messagebox.showwarning("Save Location", f"Invalid save location. Using default: {save_location}")
        
        # Build full output path
        output_path = os.path.join(save_location, output_filename)
        
        # Check if file exists and append number if needed
        counter = 1
        base_name = os.path.splitext(output_filename)[0]
        extension = os.path.splitext(output_filename)[1]
        
        while os.path.exists(output_path):
            new_filename = f"{base_name} ({counter}){extension}"
            output_path = os.path.join(save_location, new_filename)
            counter += 1
        
        try:
            doc.save(output_path)
            remember_titles_from_form(self)
            remember_agencies_from_form(self)
            messagebox.showinfo("Success", f"Report generated and saved as:\n{output_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save report: {str(e)}")
        
    def get_authority_paragraph(self):
        authority_map = {
            'Search Warrant': 'auth_sw',
            'Consent': 'auth_c',
            'Implied Consent': 'auth_i',
            'Parole': 'auth_p'
        }
        
        # Get the appropriate authority field based on role
        if self.role_type.get() == "Case Agent":
            selected_authority = self.legal_self.get()
            
            # For Case Agent with Search Warrant, don't include the auth_sw paragraph
            # because there's no requesting officer information to reference
            if selected_authority == 'Search Warrant':
                return ""  # Return empty string to skip the authority paragraph
        else:
            selected_authority = self.legal_authority.get()
        
        if selected_authority not in authority_map:
            raise ValueError(f"Invalid legal authority selected: {selected_authority}")
        return self.paragraphs[authority_map[selected_authority]]

    def format_agency(self, agency, return_abbreviation=False):
        if not agency:
            return "" if not return_abbreviation else ("", "")
                
        # Dictionary mapping normalized lowercase inputs to (full_name, abbreviation)
        agencies = {
            "federal bureau of investigation": ("Federal Bureau of Investigation", "FBI"),
            "fbi": ("Federal Bureau of Investigation", "FBI"),
            "drug enforcement administration": ("Drug Enforcement Administration", "DEA"),
            "dea": ("Drug Enforcement Administration", "DEA"),
            "bureau of alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "atf": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "department of homeland security": ("Department of Homeland Security", "DHS"),
            "dhs": ("Department of Homeland Security", "DHS"),
            "united states marshals service": ("United States Marshals Service", "USMS"),
            "usms": ("United States Marshals Service", "USMS"),
            "south dakota division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sddci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sd dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "bureau of indian affairs": ("Bureau of Indian Affairs", "BIA"),
            "bia": ("Bureau of Indian Affairs", "BIA"),
            "south dakota highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "sdhp": ("South Dakota Highway Patrol", "SDHP"),
            "national security agency": ("National Security Agency", "NSA"),
            "nsa": ("National Security Agency", "NSA"),
            "federal emergency management agency": ("Federal Emergency Management Agency", "FEMA"),
            "fema": ("Federal Emergency Management Agency", "FEMA"),
            "internal revenue service": ("Internal Revenue Service", "IRS"),
            "irs": ("Internal Revenue Service", "IRS"),
            "department of defense": ("Department of Defense", "DOD"),
            "dod": ("Department of Defense", "DOD"),
            "department of justice": ("Department of Justice", "DOJ"),
            "doj": ("Department of Justice", "DOJ")
        }
        
        agency_lower = agency.lower().strip()
        
        # Check for direct match in the agencies dictionary
        if agency_lower in agencies:
            full_name, abbr = agencies[agency_lower]
            if return_abbreviation:
                return full_name, abbr
            else:
                return full_name
        
        # Preserve original for PD/SO handling
        original_agency = agency
        original_lower = original_agency.lower().strip()
        
        # Handle PD and SO expansion and abbreviation creation
        if original_lower.endswith(" pd"):
            # Extract the location part (everything before " pd")
            location = original_agency[:-3].strip()
            agency = f"{location} Police Department"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + PD
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "PD"
        elif original_lower.endswith(" so"):
            # Extract the location part (everything before " so")
            location = original_agency[:-3].strip()
            agency = f"{location} Sheriff's Office"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + SO
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "SO"
        else:
            # Default to None if not PD/SO
            abbreviation = None
        
        # Format the agency name with title case
        formatted_name = title_agency_words(agency)
        
        # For short agencies (4 or fewer chars), use uppercase
        if len(agency) <= 4:
            formatted_name = agency.upper()
        
        # If abbreviation wasn't set (non-PD/SO), calculate it
        if abbreviation is None:
            # Check for substring matches in agencies dict (fallback for partial matches)
            found_match = False
            for full_lower, (full_name, abbr) in agencies.items():
                if full_lower in agency_lower:
                    formatted_name = full_name
                    abbreviation = abbr
                    found_match = True
                    break
            
            # If still no match, generate abbreviation dynamically
            if not found_match:
                if len(agency) <= 4:
                    abbreviation = agency.upper()
                else:
                    words = agency.split()
                    skip_words = ['of', 'the', 'and', 'in', 'on', 'at', 'by', 'for', 'with', 'a', 'an']
                    abbr_chars = [word[0].upper() for word in words if word.lower() not in skip_words and word]
                    abbreviation = ''.join(abbr_chars)
        
        if return_abbreviation:
            return formatted_name, abbreviation
        else:
            return formatted_name    
    
    def format_title(self, title):
        if len(title) <= 2:
            return title.upper()
        return ' '.join(word.capitalize() for word in title.split())

    def parse_request_date(self, date_string):
        if not date_string or not isinstance(date_string, str):
            raise ValueError("Date string cannot be empty")
        
        date_string = date_string.strip()
    
        formats = [
            '%m/%d/%Y', '%m/%d/%y',  # With leading zeros, slash separator
            '%-m/%-d/%y', '%-m/%-d/%Y',  # Without leading zeros, slash separator
            '%m/%-d/%y', '%m/%-d/%Y',  # Mixed format, slash separator
            '%-m/%d/%y', '%-m/%d/%Y',  # Mixed format, slash separator
            '%m-%d-%Y', '%m-%d-%y',  # With leading zeros, dash separator
            '%-m-%d-%y', '%-m-%d-%Y',  # Without leading zeros, dash separator
            '%m-%d-%y', '%m-%d-%Y',  # Mixed format, dash separator
            '%-m-%d-%y', '%-m-%d-%Y'  # Mixed format, dash separator
        ]
    
        for fmt in formats:
            try:
                date_obj = datetime.strptime(date_string, fmt)
                # Format the date with full month name and day of week
                formatted_date = date_obj.strftime("%A, %B %d, %Y")
                # Remove leading zeros from day while preserving proper date formatting
                formatted_date = formatted_date.replace(' 0', ' ')
                return formatted_date
            except ValueError:
                continue
    
        raise ValueError(f"Invalid date format: {date_string}. Please use MM/DD/YYYY or MM-DD-YYYY format.")

    def back_to_start(self):
        close_and_return(self)
        
    def on_closing(self):
        close_and_return(self)


# Ω Digital Forensics Report Writer Ω (ver. 1.0.1) © 2026 #