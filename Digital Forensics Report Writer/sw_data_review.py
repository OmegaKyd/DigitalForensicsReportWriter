import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt
import os
import re
from tkinterdnd2 import DND_FILES, TkinterDnD
from settings_manager import SettingsManager
from ui_theme import apply_theme, add_header_bar, add_action_bar, pack_right_actions, size_window, parse_mdy_date, add_date_entry, COLORS, close_and_return
from app_menu import attach_app_menu
from paragraphs_manager import fill_paragraph, load_paragraphs
from report_common import (
    current_dfr_prefix,
    is_complete_dfr_number,
    add_template_picker,
    sync_template_choice,
    load_request_titles,
    setup_title_combobox,
    remember_titles_from_form,
    setup_agency_combobox,
    remember_agencies_from_form,
    title_agency_words,
    saved_examiner_agency,
    bind_prefix_typeahead,
    refresh_request_title_values,
)
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

#ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)

class WarrantDataReturns(TkinterDnD.Tk):
    def __init__(self, master=None):
        super().__init__()
        
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.master = master
        apply_theme(self)
        attach_app_menu(self)
        add_header_bar(self, "Warrant Data Returns", "Warrant, subpoena, and service-provider reports")
        self.action_bar = add_action_bar(self)
        self.generate_button = ttk.Button(self.action_bar, text="Generate Report", command=self.generate_report)
        self.exit_button = ttk.Button(self.action_bar, text="Exit", command=self.on_closing)
        pack_right_actions(self.generate_button, self.exit_button)
        
        self.settings_manager = SettingsManager()
        self.current_settings = self.settings_manager.load_settings()
        self._save_timer = None
        
        self.paragraphs = {}
        self._paragraph_kind = "warrant"
        self.initialize_warrant_paragraphs()
        
        self.title("Ω Digital Forensics Report Writer - Warrant Data Returns Ω")
        
        size_window(self, 1120, 720, min_width=980, min_height=660)
        
        self.main_frame = ttk.Frame(self)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        self.columns_frame = ttk.Frame(self.main_frame)
        self.columns_frame.pack(fill=tk.X, expand=False, padx=8, pady=8)
        
        # Left Column 
        self.left_frame = ttk.Frame(self.columns_frame, padding=0)
        self.left_frame.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        
        self.scrollable_frame = ttk.Frame(self.left_frame, padding=(2, 2, 2, 2))
        self.scrollable_frame.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # Middle Column
        self.middle_frame = ttk.Frame(self.columns_frame, padding=0)
        self.middle_frame.grid(row=0, column=1, sticky="nsew", padx=0, pady=0)
        
        # Right Column
        self.right_frame = ttk.Frame(self.columns_frame, padding=0)
        self.right_frame.grid(row=0, column=2, sticky="nsew", padx=0, pady=0)
        
        self.columns_frame.columnconfigure(0, weight=1, minsize=320)
        self.columns_frame.columnconfigure(1, weight=1, minsize=320)
        self.columns_frame.columnconfigure(2, weight=1, minsize=320)
        self.columns_frame.rowconfigure(0, weight=0)
        
        self.create_widgets()

    def create_widgets(self):
        self.create_left_column_widgets()
        self.create_middle_column_widgets()
        self.create_right_column_widgets()

# LEFT PANE
    def create_left_column_widgets(self):
        self.create_request_information_frame()
        self.create_warrant_return_dates_frame()
        self.create_examiner_information_frame()

    def create_warrant_return_dates_frame(self):
        warrant_frame = ttk.LabelFrame(self.scrollable_frame, text="Warrant Return Information", padding=6)
        warrant_frame.pack(fill=tk.X, padx=5, pady=4)

        ttk.Label(warrant_frame, text="Warrant Service Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.warrant_service_date = add_date_entry(warrant_frame, row=0, column=1)

        ttk.Label(warrant_frame, text="Data Return Date:").grid(row=1, column=0, sticky="w", pady=2)
        self.data_return_date = add_date_entry(warrant_frame, row=1, column=1)

        self.time_frame_label = ttk.Label(warrant_frame, text="Time Limit Limited?:")
        self.time_frame_label.grid(row=2, column=0, sticky="w", pady=2)

        self.time_frame_checkbox_frame = ttk.Frame(warrant_frame)
        self.time_frame_checkbox_frame.grid(row=2, column=1, sticky="ew", pady=2)

        self.time_frame_limited_var = tk.IntVar(warrant_frame)
        self.time_frame_checkbox = ttk.Checkbutton(
            self.time_frame_checkbox_frame,
            text="Yes",
            variable=self.time_frame_limited_var,
            onvalue=1, offvalue=0,
            command=self.toggle_time_frame_fields
        )
        self.time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 8))

        self.time_frame_dates_frame = ttk.Frame(warrant_frame)
        self.time_frame_dates_frame.grid(row=3, column=0, columnspan=2, sticky="ew", pady=2)
        self.time_frame_dates_frame.grid_remove()

        ttk.Label(self.time_frame_dates_frame, text="Time Frame Start Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.time_frame_start_date = add_date_entry(self.time_frame_dates_frame, row=0, column=1)

        ttk.Label(self.time_frame_dates_frame, text="Time Frame End Date:").grid(row=1, column=0, sticky="w", pady=2)
        self.time_frame_end_date = add_date_entry(self.time_frame_dates_frame, row=1, column=1)

        warrant_frame.columnconfigure(1, weight=1)

    def create_request_information_frame(self):
        request_frame = ttk.LabelFrame(self.scrollable_frame, text="Request Information", padding=(4, 2))
        request_frame.pack(fill=tk.X, padx=2, pady=2)

        ttk.Label(request_frame, text="Requesting Agency:").grid(row=0, column=0, sticky="w", pady=2)
        self.requesting_agency = ttk.Combobox(request_frame)
        self.requesting_agency.grid(row=0, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.requesting_agency)

        ttk.Label(request_frame, text="Requesting Officer Title:").grid(row=1, column=0, sticky="w", pady=2)
        self.requesting_officer_title = ttk.Combobox(request_frame, values=load_request_titles())
        self.requesting_officer_title.grid(row=1, column=1, sticky="ew", pady=2)
        self.requesting_officer_title.set("")
        bind_prefix_typeahead(self.requesting_officer_title)

        self.requesting_officer_title_other = ttk.Entry(request_frame)
        self.requesting_officer_title_other.grid(row=2, column=1, sticky="ew", pady=2)
        self.requesting_officer_title_other.grid_remove()

        ttk.Label(request_frame, text="Requesting Officer:").grid(row=3, column=0, sticky="w", pady=2)
        self.requesting_officer = ttk.Entry(request_frame)
        self.requesting_officer.grid(row=3, column=1, sticky="ew", pady=2)

        ttk.Label(request_frame, text="Primary Case Offense:").grid(row=4, column=0, sticky="w", pady=2)
        self.offense_var = tk.StringVar()  # ← new variable for tracing
        self.primary_case_offense = ttk.Entry(request_frame, textvariable=self.offense_var)
        self.primary_case_offense.grid(row=4, column=1, sticky="ew", pady=2)
        self.offense_var.trace_add("write", lambda *args: self.format_offense_lowercase())

        request_frame.columnconfigure(1, weight=1)

    def format_offense_lowercase(self):
        current = self.offense_var.get()
        if current:
            formatted = ''.join(c.lower() if c.isalpha() else c for c in current)
            if formatted != current:
                self.offense_var.set(formatted)

    def toggle_requesting_officer_title_other(self, event=None):
        if self.requesting_officer_title.get() == "Other":
            self.requesting_officer_title_other.grid()
        else:
            self.requesting_officer_title_other.grid_remove()
            self.requesting_officer_title_other.delete(0, tk.END)

    def create_examiner_information_frame(self):
        examiner_frame = ttk.LabelFrame(self.scrollable_frame, text="Examiner Information", padding=(4, 2))
        examiner_frame.pack(fill=tk.X, padx=2, pady=2)

        ttk.Label(examiner_frame, text="Examiner Agency:").grid(row=0, column=0, sticky="w", pady=2)
        self.examiner_agency = ttk.Combobox(examiner_frame)
        self.examiner_agency.grid(row=0, column=1, sticky="ew", pady=2)
        setup_agency_combobox(
            self.examiner_agency,
            saved=saved_examiner_agency(self.current_settings),
            on_change=self.auto_save_settings,
        )

        self.examiner_agency_other = ttk.Entry(examiner_frame)
        self.examiner_agency_other.grid(row=1, column=1, sticky="ew", pady=2)
        self.examiner_agency_other.grid_remove()

        ttk.Label(examiner_frame, text="Examiner Title:").grid(row=2, column=0, sticky="w", pady=2)
        self.examiner_title = ttk.Combobox(examiner_frame, values=load_request_titles())
        self.examiner_title.grid(row=2, column=1, sticky="ew", pady=2)
        setup_title_combobox(
            self.examiner_title,
            saved=self.current_settings.get("examiner_title", ""),
            on_change=self.auto_save_settings,
        )

        self.examiner_title_other = ttk.Entry(examiner_frame)
        self.examiner_title_other.grid(row=3, column=1, sticky="ew", pady=2)
        self.examiner_title_other.grid_remove()

        ttk.Label(examiner_frame, text="Examiner Name:").grid(row=4, column=0, sticky="w", pady=2)
        self.examiner_name = ttk.Entry(examiner_frame)
        saved_name = self.current_settings.get("examiner_name", "")
        self.examiner_name.insert(0, saved_name)
        self.examiner_name.grid(row=4, column=1, sticky="ew", pady=2)

        ttk.Label(examiner_frame, text="Case Number:").grid(row=5, column=0, sticky="w", pady=2)
        self.case_var = tk.StringVar() 
        self.case_number = ttk.Entry(examiner_frame, textvariable=self.case_var)
        self.case_number.grid(row=5, column=1, sticky="ew", pady=2)
        self.case_var.trace_add("write", lambda *args: self.format_case_number_upper())

        ttk.Label(examiner_frame, text="DFR Number:").grid(row=6, column=0, sticky="w", pady=2)
        self.dfr_number = ttk.Entry(examiner_frame)

        self.dfr_number.insert(0, current_dfr_prefix())
        self.dfr_number.grid(row=6, column=1, sticky="ew", pady=2)

        examiner_frame.columnconfigure(1, weight=1)

    def format_case_number_upper(self):
        current = self.case_var.get()
        if current:
            formatted = ''.join(c.upper() if c.isalpha() else c for c in current)
            if formatted != current:
                self.case_var.set(formatted)

    def toggle_examiner_agency_other(self, event=None):
        return

    def toggle_examiner_title_other(self, event=None):
        if self.examiner_title.get() == "Other":
            self.examiner_title_other.grid()
        else:
            self.examiner_title_other.grid_remove()
            self.examiner_title_other.delete(0, tk.END)

    def toggle_time_frame_fields(self):
        if self.time_frame_limited_var.get() == 1:
            self.time_frame_dates_frame.grid()
        else:
            self.time_frame_dates_frame.grid_remove()
            if hasattr(self, 'time_frame_start_date'):
                self.time_frame_start_date.delete(0, tk.END)
            if hasattr(self, 'time_frame_end_date'):
                self.time_frame_end_date.delete(0, tk.END)

# MIDDLE PANE
    def create_middle_column_widgets(self):
        self.create_account_information_frame()
        self.create_forensic_software_frame()

    def create_account_information_frame(self):
        account_frame = ttk.LabelFrame(self.middle_frame, text="Account Information", padding=6)
        account_frame.pack(fill=tk.X, padx=5, pady=4, expand=False)   # ← no expand

        ttk.Label(account_frame, text="Service Provider:").grid(row=0, column=0, sticky="w", pady=2)
        self.service_provider = ttk.Entry(account_frame)
        self.service_provider.grid(row=0, column=1, sticky="ew", pady=2)

        ttk.Label(account_frame, text="Account Identifier:").grid(row=1, column=0, sticky="w", pady=2)
        self.account_identifier = ttk.Entry(account_frame)
        self.account_identifier.grid(row=1, column=1, sticky="ew", pady=2)

        ttk.Label(account_frame, text="GB of Data:").grid(row=2, column=0, sticky="w", pady=2)
        self.account_data_size = ttk.Entry(account_frame)
        self.account_data_size.grid(row=2, column=1, sticky="ew", pady=2)

        account_frame.columnconfigure(1, weight=1)

    def create_forensic_software_frame(self):
        forensic_frame = ttk.LabelFrame(self.middle_frame, text="Forensic Processing Software", padding=6)
        forensic_frame.pack(fill=tk.X, padx=5, pady=4, expand=False)

        self.cb_cellebrite_var = tk.IntVar(forensic_frame)
        self.cb_axiom_var      = tk.IntVar(forensic_frame)
        self.cb_griffeye_var   = tk.IntVar(forensic_frame)
        self.cb_manual_var     = tk.IntVar(forensic_frame)

        # Place all checkboxes in one row
        ttk.Checkbutton(
            forensic_frame, text="Cellebrite", variable=self.cb_cellebrite_var,
            onvalue=1, offvalue=0
        ).grid(row=0, column=0, sticky="w", padx=(0, 0), pady=4)

        ttk.Checkbutton(
            forensic_frame, text="AXIOM", variable=self.cb_axiom_var,
            onvalue=1, offvalue=0
        ).grid(row=0, column=1, sticky="w", padx=(0, 0), pady=4)

        ttk.Checkbutton(
            forensic_frame, text="Griffeye", variable=self.cb_griffeye_var,
            onvalue=1, offvalue=0
        ).grid(row=0, column=2, sticky="w", padx=(0, 0), pady=4)

        ttk.Checkbutton(
            forensic_frame, text="Manual Exam", variable=self.cb_manual_var,
            onvalue=1, offvalue=0
        ).grid(row=0, column=3, sticky="w", padx=(0, 0), pady=4)

        for col in range(4):
            forensic_frame.columnconfigure(col, weight=1)

    def get_selected_forensic_software(self):
        selected = []
        if hasattr(self, 'cb_cellebrite_var') and self.cb_cellebrite_var.get() == 1:
            selected.append("Cellebrite")
        if hasattr(self, 'cb_axiom_var') and self.cb_axiom_var.get() == 1:
            selected.append("AXIOM")
        if hasattr(self, 'cb_griffeye_var') and self.cb_griffeye_var.get() == 1:
            selected.append("Griffeye")
        if hasattr(self, 'cb_manual_var') and self.cb_manual_var.get() == 1:
            selected.append("Manual Exam")
        return selected

    def create_output_file_frame(self):
        output_frame = ttk.LabelFrame(self.right_frame, text="Output File", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)
        content = ttk.Frame(output_frame)
        content.pack(fill=tk.X, expand=True)

        # Output File Name
        ttk.Label(content, text="Output File Name:").grid(row=0, column=0, sticky="w", pady=2)
        self.output_filename = ttk.Entry(content)
        self.output_filename.grid(row=1, column=0, columnspan=2, sticky="ew", pady=2)

        # Hint text
        hint_style = ttk.Style()
        hint_style.configure("Hint.TLabel", font=('Arial', 8))
        hint_label = ttk.Label(content,
                               text="*Default File Name \"(DFR #) - (Provider) (AccountID) Return\"",
                               style="Hint.TLabel", wraplength=350)
        hint_label.grid(row=2, column=0, columnspan=2, sticky="w", pady=(0, 6))

        # Save Location label
        ttk.Label(content, text="Save Location:").grid(row=3, column=0, sticky="w", pady=2)

        # Save Location entry (full width, no sub-frame)
        self.save_location = ttk.Entry(content)
        self.save_location.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        self.save_location.insert(0, os.path.join(os.path.expanduser("~"), "Desktop"))

        # Browse button on its own row below the entry
        ttk.Button(content, text="Browse", command=self.browse_save_location).grid(row=5, column=0, columnspan=1, pady=4, sticky="ew")

        content.columnconfigure(1, weight=1)

    def create_template_file_frame(self):
        template_frame = ttk.LabelFrame(self.right_frame, text="Template File", padding=6)
        template_frame.pack(fill=tk.X, expand=False, padx=5, pady=4, anchor="n")
        add_template_picker(self, template_frame, preferred="DFR SW Return (2026).docx", keywords=("sw", "warrant", "return"))
         
    def on_template_drop(self, event):
        """Handle drag-and-drop of DFR template file."""
        try:
            file_path = self.tk.splitlist(event.data)[0]

            if not file_path.lower().endswith('.docx'):
                messagebox.showerror("Invalid File", "Please drop a .docx file.")
                return

            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_status_label.configure(text=f"Selected: {os.path.basename(file_path)}")

        except Exception as e:
            messagebox.showerror("Drop Error", f"Failed to process drop:\n{str(e)}")

    def browse_template(self, event=None):
        """Browse for DFR template file."""
        path = filedialog.askopenfilename(
            title="Select DFR Template",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")]
        )
        if path:
            self.template_file = path
            sync_template_choice(self, path)
            self.template_status_label.configure(text=f"Selected: {os.path.basename(path)}")

    def create_right_column_widgets(self):
        self.create_template_file_frame()
        self.create_output_file_frame()

    # ================== HELPERS ==================
    def toggle_title_entry(self, event=None):
        self.toggle_requesting_officer_title_other(event)

    def toggle_examiner_title_entry(self, event=None):
        self.toggle_examiner_title_other(event)

    def on_agency_type_changed(self, event=None):
        self.toggle_examiner_agency_other(event)
        self.auto_save_settings()

    def toggle_agency_entry(self, event=None):
        self.toggle_examiner_agency_other(event)

    def auto_save_settings(self, event=None):
        if hasattr(self, '_save_timer') and self._save_timer:
            self.after_cancel(self._save_timer)
        self._save_timer = self.after(500, self._perform_auto_save)

    def _perform_auto_save(self):
        try:
            # Extract ONLY the prefix from the current DFR field
            current_dfr = self.dfr_number.get().strip()
            prefix = current_dfr

            # If there's a hyphen, take everything up to and including the hyphen
            if '-' in current_dfr:
                prefix = current_dfr.rsplit('-', 1)[0] + '-'

            # Clean and validate prefix
            prefix = prefix.strip()
            if not prefix.endswith('-'):
                prefix += '-'

            # Ensure it's a valid DFR prefix (DFR + year + -)
            if not (prefix.startswith('DFR') and len(prefix) >= 8 and prefix[3:7].isdigit()):
                current_year = datetime.now().year
                prefix = f"DFR{current_year}-"

            settings_to_save = {
                "examiner_agency_type": self.examiner_agency.get(),
                "examiner_agency_custom": "",
                "examiner_title": self.examiner_title.get(),
                "examiner_title_custom": self.examiner_title_other.get() if self.examiner_title.get() == "Other" else "",
                "examiner_name": self.examiner_name.get(),
                "dfr_number_prefix": prefix,  # always just the prefix (e.g. "DFR2026-")
            }
            self.settings_manager.save_settings(settings_to_save)
            self.current_settings.update(settings_to_save)
            remember_agencies_from_form(self)
        except Exception as e:
            print(f"Error auto-saving: {e}")

    def get_dfr_prefix(self):
        dfr = self.dfr_number.get() if hasattr(self, "dfr_number") else ""
        match = re.match(r'(DFR\d{4}-)', dfr)
        return match.group(1) if match else f"DFR{datetime.now().year}-"

    def setup_auto_complete_dropdown(self, combobox):
        combobox.bind('<KeyRelease>', lambda event, cb=combobox: self.auto_complete(event, cb))

    def auto_complete(self, event, combobox):
        if event.keysym in ('Up', 'Down', 'Left', 'Right', 'Return', 'Tab', 'BackSpace', 'Delete'):
            return
        value = combobox.get().lower()
        if not value:
            return
        for option in combobox['values']:
            if option.lower().startswith(value):
                combobox.set(option)
                combobox.icursor(len(option))
                combobox.selection_range(len(value), len(option))
                break

    def browse_save_location(self):
        current = self.save_location.get().strip() or os.path.join(os.path.expanduser("~"), "Desktop")
        selected = filedialog.askdirectory(title="Select Save Location", initialdir=current)
        if selected:
            self.save_location.delete(0, tk.END)
            self.save_location.insert(0, selected)

###UPDATE PARAGRAPHS###

    def initialize_warrant_paragraphs(self):
        self.paragraphs = load_paragraphs("warrant")

    def generate_paragraphs(self, data, new_doc):
        def add(text):
            p = new_doc.add_paragraph(fill_paragraph(text, data))
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        add(self.paragraphs['intro_self'])
        #add(self.paragraphs['conclusion'])

###PARAGRAPHS###

    def validate_fields(self):
        missing = []

        # Examiner Agency
        examiner_agency = self.examiner_agency.get().strip()
        if not examiner_agency:
            missing.append("Examiner Agency")

        # Examiner Title
        examiner_title = self.examiner_title.get().strip()
        if not examiner_title:
            missing.append("Examiner Title")

        # Examiner Name
        if not self.examiner_name.get().strip():
            missing.append("Examiner Name")

        # Case Number
        if not self.case_number.get().strip():
            missing.append("Case Number")

        # DFR Number
        if not is_complete_dfr_number(self.dfr_number.get()):
            missing.append("DFR Number")

        # Requesting Agency
        if not self.requesting_agency.get().strip():
            missing.append("Requesting Agency")

        # Requesting Officer Title
        req_title = self.requesting_officer_title.get().strip()
        if not req_title:
            missing.append("Requesting Officer Title")

        # Requesting Officer
        if not self.requesting_officer.get().strip():
            missing.append("Requesting Officer")

        # Primary Case Offense
        if not self.primary_case_offense.get().strip():
            missing.append("Primary Case Offense")

        # Warrant Service Date
        if not self.warrant_service_date.get().strip():
            missing.append("Warrant Service Date")

        # Data Return Date
        if not self.data_return_date.get().strip():
            missing.append("Data Return Date")

        # Forensic Software
        if not self.get_selected_forensic_software():
            missing.append("At least one Forensic Processing Software")

        # Template File - use the correct attribute name
        template_path = getattr(self, 'template_file', None)
        if not template_path or not os.path.exists(template_path):
            missing.append(f"DFR Template File (current path: {template_path or 'None'})")

        return missing

    def generate_report(self):
        missing = self.validate_fields()
        if missing:
            messagebox.showerror("Missing Fields", f"Please fill in the following:\n\n" + "\n".join(f"• {f}" for f in missing))
            return

        # Load the template
        if not hasattr(self, 'template_file') or not self.template_file:
            messagebox.showerror("Error", "No template file selected.")
            return

        doc = Document(self.template_file)

        # Gather data using current widgets
        data = {
            'PY_DFR': self.dfr_number.get().strip().upper(),
            'PY_CASENUMBER': self.case_number.get().strip().upper(),
            'PY_EXAMINER': f"{self.format_title(self.examiner_title.get())} {self.examiner_name.get().strip().title()}",
            'PY_REQAGENCY': self.format_agency(self.requesting_agency.get().strip()),
            'PY_REQOFF': f"{self.format_title(self.requesting_officer_title.get())} {self.requesting_officer.get().strip().title()}",
            'PY_SERVEDATE': self.parse_request_date(self.warrant_service_date.get().strip()),
            'PY_RETURNDATE': self.parse_request_date(self.data_return_date.get().strip()),
            'PY_PROVIDER': self.service_provider.get().strip(),
            'PY_ACCOUNTID': self.account_identifier.get().strip(),
            'PY_DATASIZE': self.account_data_size.get().strip(),
        }

        # Optional time frame if checked
        if self.time_frame_limited_var.get() == 1:
            data['PY_LIMITSTART'] = self.parse_request_date(self.time_frame_start_date.get().strip())
            data['PY_LIMITEND'] = self.parse_request_date(self.time_frame_end_date.get().strip())
        else:
            data['PY_LIMITSTART'] = ""
            data['PY_LIMITEND'] = ""

        # Create search docs for replacement
        search_docs = {}
        for key in ['PY_DFR', 'PY_CASENUMBER', 'PY_EXAMINER', 'PY_REQAGENCY', 'PY_REQOFF', 'PY_SERVEDATE', 'PY_RETURNDATE', 'PY_PROVIDER', 'PY_ACCOUNTID', 'PY_DATASIZE', 'PY_LIMITSTART', 'PY_LIMITEND']:
            search_docs[key] = Document()
            p = search_docs[key].add_paragraph(data.get(key, ''))
            p.style = search_docs[key].styles['Normal']

        # Add examination summary/details (from your PY_TEXT or other)
        # Assuming you have a summary field or generate it
        summary_doc = Document()
        summary_doc.add_paragraph("Examination Summary: [Add details here if needed]")
        search_docs['PY_TEXT'] = summary_doc

        # Perform replacements
        self.search_and_replace_content_controls_simple(doc, search_docs)
        self.search_and_replace_split_placeholders(doc, search_docs)

        # Save
        self.save_document(doc)

    def generate_replacements(self, data, search_docs):
        def add_text(doc, text):
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
        for key, doc in search_docs.items():
            if key == "PY_TEXT":
                continue
            add_text(doc, data.get(key, ''))

    def save_document(self, doc):
        # Build the suggested filename
        dfr = self.dfr_number.get().strip().upper()
        provider = self.service_provider.get().strip()
        account_id = self.account_identifier.get().strip()

        safe_provider = re.sub(r'[^\w\-]', '_', provider) if provider else "UnknownProvider"
        safe_account = re.sub(r'[^\w\-]', '_', account_id) if account_id else "UnknownAccount"

        filename = f"{dfr} - {safe_provider} ({safe_account}) Return.docx"

        # Get default save directory from the UI Entry widget (or fallback)
        default_dir = None
        if hasattr(self, 'save_location'):
            default_dir = self.save_location.get().strip()
        if not default_dir:
            default_dir = self.current_settings.get("default_save_path", os.path.expanduser("~/Desktop"))

        # Full default path
        default_path = os.path.join(default_dir, filename)

        # Auto-save if valid directory exists
        if default_dir and os.path.isdir(default_dir):
            try:
                final_path = default_path
                # Avoid overwrite by adding suffix if file exists
                if os.path.exists(final_path):
                    base, ext = os.path.splitext(final_path)
                    i = 1
                    while os.path.exists(final_path):
                        final_path = f"{base}_{i}{ext}"
                        i += 1

                doc.save(final_path)
                remember_titles_from_form(self)
                remember_agencies_from_form(self)
                messagebox.showinfo("Success", f"Report auto-saved to:\n{final_path}")

                # Optional: remember this directory for next time
                self.current_settings["default_save_path"] = default_dir
                self.settings_manager.save_settings(self.current_settings)
                return
            except Exception as e:
                messagebox.showwarning("Auto-save Failed", f"Auto-save error:\n{str(e)}\n\nFalling back to manual selection.")
                # Continue to dialog

        # Fallback: manual save dialog
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialdir=default_dir,
            initialfile=filename,
            title="Save Warrant Data Return Report"
        )

        if save_path:
            try:
                doc.save(save_path)
                remember_titles_from_form(self)
                remember_agencies_from_form(self)
                messagebox.showinfo("Success", f"Report saved to:\n{save_path}")

                # Remember the chosen directory
                chosen_dir = os.path.dirname(save_path)
                self.current_settings["default_save_path"] = chosen_dir
                self.settings_manager.save_settings(self.current_settings)
            except Exception as e:
                messagebox.showerror("Save Error", f"Could not save:\n{str(e)}")
        else:
            messagebox.showinfo("Cancelled", "Save cancelled.")

    def format_agency(self, agency, return_abbreviation=False):
        if not agency:
            return "" if not return_abbreviation else ("", "")
        
        agencies = {
            "federal bureau of investigation": ("Federal Bureau of Investigation", "FBI"),
            "fbi": ("Federal Bureau of Investigation", "FBI"),
            "drug enforcement administration": ("Drug Enforcement Administration", "DEA"),
            "dea": ("Drug Enforcement Administration", "DEA"),
            "bureau of alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "atf": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "department of homeland security": ("Department of Homeland Security", "DHS"),
            "dhs": ("Department of Homeland Security", "DHS"),
            "united states marshals service": ("United States Marshals Service", "USMS"),
            "usms": ("United States Marshals Service", "USMS"),
            "south dakota division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "bureau of indian affairs": ("Bureau of Indian Affairs", "BIA"),
            "bia": ("Bureau of Indian Affairs", "BIA"),
            "south dakota highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "sdhp": ("South Dakota Highway Patrol", "SDHP"),
            "national security agency": ("National Security Agency", "NSA"),
            "nsa": ("National Security Agency", "NSA"),
            "federal emergency management agency": ("Federal Emergency Management Agency", "FEMA"),
            "fema": ("Federal Emergency Management Agency", "FEMA"),
            "internal revenue service": ("Internal Revenue Service", "IRS"),
            "irs": ("Internal Revenue Service", "IRS"),
        }
        
        agency_lower = agency.lower().strip()
        
        if agency_lower in agencies:
            full_name, abbr = agencies[agency_lower]
            return (full_name, abbr) if return_abbreviation else full_name
        
        original_agency = agency.strip()
        original_lower = original_agency.lower()
        
        abbreviation = None
        if original_lower.endswith(" pd"):
            location = original_agency[:-3].strip()
            agency = f"{location} Police Department"
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "PD"
        elif original_lower.endswith(" so"):
            location = original_agency[:-3].strip()
            agency = f"{location} Sheriff's Office"
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "SO"
        
        formatted_name = title_agency_words(agency)
        if len(agency) <= 4:
            formatted_name = agency.upper()
        
        if abbreviation is None:
            found_match = False
            for full_lower, (full_name, abbr) in agencies.items():
                if full_lower in agency_lower:
                    formatted_name = full_name
                    abbreviation = abbr
                    found_match = True
                    break
            if not found_match:
                if len(agency) <= 4:
                    abbreviation = agency.upper()
                else:
                    words = agency.split()
                    skip_words = ['of', 'the', 'and', 'in', 'on', 'at', 'by', 'for', 'with', 'a', 'an']
                    abbr_chars = [word[0].upper() for word in words if word.lower() not in skip_words and word]
                    abbreviation = ''.join(abbr_chars)
        
        if return_abbreviation:
            return formatted_name, abbreviation
        return formatted_name

    def format_title(self, title):
        if len(title) <= 2:
            return title.upper()
        return ' '.join(word.capitalize() for word in title.split())

    def get_request_agency_formatted(self):
        agency_text = self.requesting_agency.get().strip()
        if agency_text:
            return self.format_agency(agency_text)
        return ""

    def get_request_title(self):
        return (self.requesting_officer_title.get() or "").strip()

    def get_examiner_agency(self):
        agency = (self.examiner_agency.get() or "").strip()
        if agency == "South Dakota DCI":
            return "South Dakota Division of Criminal Investigation"
        return agency

    def parse_request_date(self, date_string):
        return parse_mdy_date(date_string, empty_ok=True)

    def format_time_frame_date(self, date_string):
        return self.parse_request_date(date_string)

    def search_and_replace_content_controls_simple(self, doc, search_docs):
        replaced_strings = {}
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False

        doc_xml_str = etree.tostring(doc._element, encoding='unicode')

        for search_string, replacement_doc in search_docs.items():
            if search_string in doc_xml_str:
                if search_string == "PY_TEXT":
                    replacement_xml_parts = []
                    for para in replacement_doc.paragraphs:
                        para_text = para.text or ""
                        text_parts = para_text.split('\n')
                        for text_part in text_parts:
                            if not text_part and len(text_parts) > 1:
                                replacement_xml_parts.append('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr></w:pPr></w:p>')
                                continue
                            para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr></w:pPr>'
                            if text_part:
                                para_xml += f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="22"/></w:rPr><w:t xml:space="preserve">{text_part.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")}</w:t></w:r>'
                            para_xml += '</w:p>'
                            replacement_xml_parts.append(para_xml)
                    replacement_text = ''.join(replacement_xml_parts)
                else:
                    replacement_text = replacement_doc.paragraphs[0].text if replacement_doc.paragraphs else ""
                    replacement_text = replacement_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                doc_xml_str = doc_xml_str.replace(search_string, replacement_text)
                replaced_strings[search_string] = True

        new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
        doc._element.clear()
        for child in new_element:
            doc._element.append(child)
        
        return replaced_strings

    def search_and_replace_split_placeholders(self, doc, search_docs):
        replaced_strings = {}
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False

        try:
            doc_xml_str = etree.tostring(doc._element, encoding='unicode')
            original_xml_str = doc_xml_str

            target_placeholders = ['PY_DFR', 'PY_EXAMINER', 'PY_REQAGENCY', 'PY_REQOFF', 'PY_SERVEDATE', 'PY_RETURNDATE', 'PY_PROVIDER', 'PY_ACCOUNTID', 'PY_DATASIZE']

            for search_string in target_placeholders:
                if replaced_strings.get(search_string):
                    continue
                replacement_doc = search_docs.get(search_string)
                if not replacement_doc or not replacement_doc.paragraphs:
                    continue
                replacement_text = replacement_doc.paragraphs[0].text or ""
                replacement_text = replacement_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&apos;')

                pattern_parts = [re.escape(char) for char in search_string]
                pattern = r'(?:<[^>]*>)*'.join(pattern_parts)

                matches = list(re.finditer(pattern, doc_xml_str))
                if matches:
                    for match in reversed(matches):
                        doc_xml_str = doc_xml_str[:match.start()] + replacement_text + doc_xml_str[match.end():]
                    replaced_strings[search_string] = True

            if doc_xml_str != original_xml_str:
                new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
                doc._element.clear()
                for child in new_element:
                    doc._element.append(child)
        except:
            pass
        return replaced_strings

    def on_closing(self):
        close_and_return(self)

if __name__ == "__main__":
    app = WarrantDataReturns()
    app.mainloop()
    
# Ω Digital Forensics Report Writer Ω (ver. 1.0.1) © 2026 #