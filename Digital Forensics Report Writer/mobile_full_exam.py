import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt
import os
import sys
import time
import ctypes
import re
import PyPDF2
from tkinter.scrolledtext import ScrolledText
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from settings_manager import SettingsManager
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree
from ui_theme import apply_theme, add_header_bar, add_action_bar, pack_right_actions, size_window, build_extracted_info_pane, parse_mdy_date, add_date_entry, COLORS, FormTabs, close_and_return
from app_menu import attach_app_menu
from paragraphs_manager import fill_paragraph, load_paragraphs
from cellebrite_pdf import process_pdf_selection, parse_cellebrite_pdfs_only
from report_common import (
    prefer_gui_over_parsed,
    ask_open_file,
    ask_directory,
    default_export_dir,
    remember_folder,
    suggested_report_filename,
    unique_output_path,
    apply_suggested_filename,
    seed_save_location,
    show_placeholder_preview,
    mobile_preview_rows,
    require_device_identity,
    current_dfr_prefix,
    is_complete_dfr_number,
    add_template_picker,
    sync_template_choice,
    load_request_titles,
    setup_title_combobox,
    remember_titles_from_form,
    setup_agency_combobox,
    remember_agencies_from_form,
    title_agency_words,
    saved_examiner_agency,
    bind_prefix_typeahead,
    refresh_request_title_values,
)

#ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)

class MobileFullExam(TkinterDnD.Tk):
    def __init__(self, master=None):
        super().__init__()
        
        self.role_type = None  # Initialize to avoid AttributeError

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.master = master
        apply_theme(self)
        attach_app_menu(self)
        add_header_bar(self, "Full Mobile Exam", "For Use With Cellebrite and Graykey Extractions")
        self.action_bar = add_action_bar(self)
        self.preview_button = ttk.Button(self.action_bar, text="Preview", command=self.preview_placeholders)
        self.generate_button = ttk.Button(self.action_bar, text="Generate Report", command=self.generate_report)
        self.exit_button = ttk.Button(self.action_bar, text="Exit", command=self.on_closing)
        pack_right_actions(self.preview_button, self.generate_button, self.exit_button)
        
        # Initialize settings manager
        self.settings_manager = SettingsManager()
        self.current_settings = self.settings_manager.load_settings()
        self._save_timer = None  # For auto-save debouncing
        
        # Initialize mobile exam specific paragraphs
        self.paragraphs = {}
        self._paragraph_kind = "mobile_full"
        self.initialize_mobile_exam_paragraphs()
        
        self.title("Ω Digital Forensics Report Writer - Full Mobile Exam Ω")
        
        size_window(self, 1260, 700)
        
        self.main_frame = ttk.Frame(self)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.paned_window = ttk.PanedWindow(self.main_frame, orient=tk.HORIZONTAL)
        self.paned_window.pack(fill=tk.BOTH, expand=True)

        self.left_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.left_frame, weight=3)

        self.form_tabs = FormTabs(self.left_frame)
        self.tab_request = self.form_tabs.add_tab("request", "Request Info")
        self.tab_examiner = self.form_tabs.add_tab("examiner", "Examiner Info")
        self.tab_device = self.form_tabs.add_tab("device", "Device Info")
        self.tab_output = self.form_tabs.add_tab("output", "Output Info")
        self.scrollable_frame = self.tab_request
        self.middle_frame = self.tab_device

        self.right_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.right_frame, weight=2)
        
        # Create the widgets and info display
        self.create_widgets()
        self.create_info_display()
        
        # Bind mousewheel to scrolling
        self.bind_all("<MouseWheel>", self._on_mousewheel)
        self.bind_tab_status_events()
        self.refresh_tab_status()

    def _on_mousewheel(self, event):
        if hasattr(self, "form_tabs"):
            self.form_tabs.on_mousewheel(event)

    def create_widgets(self):
        style = ttk.Style()
        style.configure('Header.TLabel', font=('Arial', 12, 'bold'))
        style.configure('Section.TFrame', padding='10')

        # Create each section in separate methods
        self.create_left_column_widgets()
        self.create_middle_column_widgets() 
        self.create_right_column_widgets()

    def _field_filled(self, widget):
        if widget is None:
            return False
        try:
            return bool(widget.get().strip())
        except Exception:
            return False

    def bind_tab_status_events(self):
        widgets = [
            getattr(self, name, None)
            for name in (
                "role_type", "request_date", "request_agency", "request_officer", "case_type",
                "offense_type", "legal_self", "sw_service_date", "transfer_title", "transfer_officer",
                "transfer_agency", "transfer_date", "examiner_agency_type", "examiner_agency_entry",
                "examiner_title_type", "examiner_title_entry", "examiner_name", "DFR_Num",
                "device_owner", "output_filename", "save_location",
            )
        ]
        for widget in widgets:
            if widget is None:
                continue
            try:
                widget.bind("<KeyRelease>", lambda e: self.refresh_tab_status(), add="+")
                widget.bind("<<ComboboxSelected>>", lambda e: self.refresh_tab_status(), add="+")
            except Exception:
                pass

    def refresh_tab_status(self):
        if not hasattr(self, "form_tabs"):
            return
        request_ok = self._field_filled(getattr(self, "role_type", None))
        if getattr(self, "role_type", None) and self.role_type.get() == "Agency Assist":
            request_ok = all(self._field_filled(getattr(self, name, None)) for name in (
                "request_date", "request_agency", "request_officer", "case_type"
            ))
        elif getattr(self, "role_type", None) and self.role_type.get() == "Case Agent":
            request_ok = self._field_filled(getattr(self, "offense_type", None))
            if getattr(self, "legal_self", None) and self.legal_self.get() == "Search Warrant":
                request_ok = request_ok and self._field_filled(getattr(self, "sw_service_date", None))
        if getattr(self, "device_transfer_var", None) and self.device_transfer_var.get() == 1:
            request_ok = request_ok and all(self._field_filled(getattr(self, name, None)) for name in (
                "transfer_title", "transfer_officer", "transfer_agency", "transfer_date"
            ))

        examiner_ok = self._field_filled(getattr(self, "examiner_name", None)) and is_complete_dfr_number(self.DFR_Num.get() if hasattr(self, "DFR_Num") else "")
        if getattr(self, "examiner_agency_type", None):
            examiner_ok = examiner_ok and self._field_filled(self.examiner_agency_type)
        if getattr(self, "examiner_title_type", None):
            examiner_ok = examiner_ok and self._field_filled(self.examiner_title_type)

        device_ok = self._field_filled(getattr(self, "device_owner", None))
        try:
            device_ok = device_ok and bool(self.get_selected_forensic_software())
        except Exception:
            pass

        output_ok = bool(self._field_filled(getattr(self, "save_location", None)))
        if output_ok:
            try:
                output_ok = os.path.isdir(self.save_location.get().strip())
            except Exception:
                output_ok = False

        self.form_tabs.set_complete("request", request_ok)
        self.form_tabs.set_complete("examiner", examiner_ok)
        self.form_tabs.set_complete("device", device_ok)
        self.form_tabs.set_complete("output", output_ok)

    def create_left_column_widgets(self):
        self.create_role_selection_frame()
        self.create_request_information_frame()
        self.create_transfer_information_frame()
        self.create_examiner_information_frame()
        


    def create_role_selection_frame(self):
        role_frame = ttk.LabelFrame(self.scrollable_frame, text="Examination Type", padding="10")
        role_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(role_frame, text="Select Role:").grid(row=0, column=0, sticky="w", pady=2)
        self.role_type = ttk.Combobox(role_frame, 
            values=["Agency Assist", "Case Agent"], state="readonly")
        self.role_type.grid(row=0, column=1, sticky="ew", pady=2)
        self.role_type.set("Agency Assist")  
        self.role_type.bind('<<ComboboxSelected>>', self.toggle_request_content)
        self.setup_auto_complete_dropdown(self.role_type)

        role_frame.columnconfigure(1, weight=1)

    def create_request_information_frame(self):
        request_frame = ttk.LabelFrame(self.scrollable_frame, text="Request Information", padding="10")
        request_frame.pack(fill=tk.X, padx=5, pady=5)

        # Create two frames for different sets of fields
        self.agency_assist_frame = ttk.Frame(request_frame)
        self.agency_assist_frame.pack(fill=tk.X, padx=5, pady=5)

        self.case_agent_frame = ttk.Frame(request_frame)
        self.case_agent_frame.pack(fill=tk.X, padx=5, pady=5)
        self.case_agent_frame.pack_forget()  # Hide initially

        self.create_agency_assist_fields()
        self.create_case_agent_fields()

    def create_agency_assist_fields(self):
        # Request Date
        ttk.Label(self.agency_assist_frame, text="Exam Request Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.request_date = add_date_entry(self.agency_assist_frame, row=0, column=1)

        # Request Agency
        ttk.Label(self.agency_assist_frame, text="Requesting Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.request_agency = ttk.Combobox(self.agency_assist_frame)
        self.request_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.request_agency)

        # Request Title (with dropdown and custom entry)
        ttk.Label(self.agency_assist_frame, text="Requesting Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        request_title_frame = ttk.Frame(self.agency_assist_frame)
        request_title_frame.grid(row=2, column=1, sticky="ew", pady=2)

        self.request_title_type = ttk.Combobox(request_title_frame, values=load_request_titles())
        self.request_title_type.pack(fill=tk.X, expand=True)
        self.request_title_type.set("")
        bind_prefix_typeahead(self.request_title_type)

        self.request_title_entry = ttk.Entry(request_title_frame)
        self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        self.request_title_entry.pack_forget()

        # Request Officer
        ttk.Label(self.agency_assist_frame, text="Requesting Officer:").grid(row=3, column=0, sticky="w", pady=2)
        self.request_officer = ttk.Entry(self.agency_assist_frame)
        self.request_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Offense Type
        ttk.Label(self.agency_assist_frame, text="Primary Case Offense:").grid(row=4, column=0, sticky="w", pady=2)
        self.case_type = ttk.Entry(self.agency_assist_frame)
        self.case_type.grid(row=4, column=1, sticky="ew", pady=2)

        # Legal Authority
        ttk.Label(self.agency_assist_frame, text="Legal Authority:").grid(row=5, column=0, sticky="w", pady=2)
        self.legal_authority = ttk.Combobox(self.agency_assist_frame, 
            values=['Search Warrant', 'Consent', 'Parole', 'Implied Consent'], state="readonly")
        self.legal_authority.grid(row=5, column=1, sticky="ew", pady=2)
        self.legal_authority.set('Search Warrant')
        self.legal_authority.bind('<<ComboboxSelected>>', self.toggle_time_frame_section)
        self.setup_auto_complete_dropdown(self.legal_authority)

        # Time Frame Section (only for Search Warrant)
        self.time_frame_label = ttk.Label(self.agency_assist_frame, text="Time Frame Limited?:")
        self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)


        
        # Create checkbox frame
        self.time_frame_checkbox_frame = ttk.Frame(self.agency_assist_frame)
        self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)
        
        # Create checkbox variable
        self.time_frame_limited_var = tk.IntVar(self.agency_assist_frame)
        
        # Create checkbox
        self.time_frame_checkbox = ttk.Checkbutton(
            self.time_frame_checkbox_frame,
            text="Yes",
            variable=self.time_frame_limited_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_time_frame_fields
        )
        self.time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Set default to unchecked
        self.time_frame_limited_var.set(0)

        # Time Frame Date Fields (initially hidden)
        self.time_frame_dates_frame = ttk.Frame(self.agency_assist_frame)
        self.time_frame_dates_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        self.time_frame_dates_frame.grid_remove()

        ttk.Label(self.time_frame_dates_frame, text="Time Frame Start Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.time_frame_start_date = add_date_entry(self.time_frame_dates_frame, row=0, column=1)

        ttk.Label(self.time_frame_dates_frame, text="Time Frame End Date:").grid(row=1, column=0, sticky="w", pady=2)
        self.time_frame_end_date = add_date_entry(self.time_frame_dates_frame, row=1, column=1)


        # Configure grid weights
        self.time_frame_dates_frame.columnconfigure(1, weight=1)
        self.agency_assist_frame.columnconfigure(1, weight=1)

    def create_case_agent_fields(self):
        # Offense Type
        ttk.Label(self.case_agent_frame, text="Primary Case Offense:").grid(row=0, column=0, sticky="w", pady=2)
        self.offense_type = ttk.Entry(self.case_agent_frame)
        self.offense_type.grid(row=0, column=1, sticky="ew", pady=2)

        # Legal Authority
        ttk.Label(self.case_agent_frame, text="Legal Authority:").grid(row=1, column=0, sticky="w", pady=2)
        self.legal_self = ttk.Combobox(self.case_agent_frame, 
            values=['Search Warrant', 'Consent', 'Parole', 'Implied Consent'], state="readonly")
        self.legal_self.grid(row=1, column=1, sticky="ew", pady=2)
        self.legal_self.set('Search Warrant')
        self.legal_self.bind('<<ComboboxSelected>>', self.toggle_sw_date_and_time_frame)
        self.setup_auto_complete_dropdown(self.legal_self)

        # Search Warrant Date (initially shown since default is Search Warrant)
        self.sw_date_frame = ttk.Frame(self.case_agent_frame)
        self.sw_date_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=2)

        ttk.Label(self.sw_date_frame, text="Search Warrant Service Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.sw_service_date = add_date_entry(self.sw_date_frame, row=0, column=1)

        # Time Frame Section (only for Search Warrant)
        self.case_time_frame_label = ttk.Label(self.case_agent_frame, text="Time Frame Limited?:")
        self.case_time_frame_label.grid(row=3, column=0, sticky="w", pady=2)
        
        # Create checkbox frame
        self.case_time_frame_checkbox_frame = ttk.Frame(self.case_agent_frame)
        self.case_time_frame_checkbox_frame.grid(row=3, column=1, sticky="ew", pady=2)
        
        # Create checkbox variable
        self.case_time_frame_limited_var = tk.IntVar(self.case_agent_frame)
        
        # Create checkbox
        self.case_time_frame_checkbox = ttk.Checkbutton(
            self.case_time_frame_checkbox_frame,
            text="Yes",
            variable=self.case_time_frame_limited_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_case_time_frame_fields
        )
        self.case_time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Set default to unchecked
        self.case_time_frame_limited_var.set(0)

        # Time Frame Date Fields (initially hidden)
        self.case_time_frame_dates_frame = ttk.Frame(self.case_agent_frame)
        self.case_time_frame_dates_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        self.case_time_frame_dates_frame.grid_remove()

        ttk.Label(self.case_time_frame_dates_frame, text="Time Frame Start Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.case_time_frame_start_date = add_date_entry(self.case_time_frame_dates_frame, row=0, column=1)

        ttk.Label(self.case_time_frame_dates_frame, text="Time Frame End Date:").grid(row=1, column=0, sticky="w", pady=2)
        self.case_time_frame_end_date = add_date_entry(self.case_time_frame_dates_frame, row=1, column=1)

        # Configure grid weights
        self.sw_date_frame.columnconfigure(1, weight=1)
        self.case_time_frame_dates_frame.columnconfigure(1, weight=1)
        self.case_agent_frame.columnconfigure(1, weight=1)

    def toggle_time_frame_section(self, event=None):
        """Shows/hides the entire time frame section based on legal authority selection"""
        if self.legal_authority.get() == 'Search Warrant':
            self.time_frame_label.grid()
            self.time_frame_checkbox_frame.grid()
            # If checkbox is checked, also show the date fields
            if self.time_frame_limited_var.get() == 1:
                self.time_frame_dates_frame.grid()
        else:
            # Hide everything and clear fields
            self.time_frame_label.grid_remove()
            self.time_frame_checkbox_frame.grid_remove()
            self.time_frame_dates_frame.grid_remove()
            # Reset checkbox and clear date fields
            self.time_frame_limited_var.set(0)
            if hasattr(self, 'time_frame_start_date'):
                self.time_frame_start_date.delete(0, tk.END)
            if hasattr(self, 'time_frame_end_date'):
                self.time_frame_end_date.delete(0, tk.END)

                
    def toggle_time_frame_fields(self, event=None):
        """Shows/hides the date fields based on checkbox state"""
        if self.time_frame_limited_var.get() == 1:
            self.time_frame_dates_frame.grid()
        else:
            self.time_frame_dates_frame.grid_remove()
            # Clear date fields when hiding
            if hasattr(self, 'time_frame_start_date'):
                self.time_frame_start_date.delete(0, tk.END)
            if hasattr(self, 'time_frame_end_date'):
                self.time_frame_end_date.delete(0, tk.END)

    
    def toggle_sw_date_and_time_frame(self, event=None):
        """Handle both SW date and time frame toggling for Case Agent"""
        # Call existing SW date toggle
        self.toggle_sw_date()
        # Call time frame toggle
        self.toggle_case_time_frame_section()

    def toggle_case_time_frame_section(self, event=None):
        """Show/hide time frame section based on legal authority selection (Case Agent)"""
        if hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant':
            # Show time frame section
            self.case_time_frame_label.grid(row=3, column=0, sticky="w", pady=2)
            self.case_time_frame_checkbox_frame.grid(row=3, column=1, sticky="ew", pady=2)
            # Show date fields if checkbox is checked
            if self.case_time_frame_limited_var.get() == 1:
                self.case_time_frame_dates_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            # Hide time frame section and clear data
            self.case_time_frame_label.grid_remove()
            self.case_time_frame_checkbox_frame.grid_remove()
            self.case_time_frame_dates_frame.grid_remove()
            # Reset checkbox and clear date fields
            self.case_time_frame_limited_var.set(0)
            if hasattr(self, 'case_time_frame_start_date'):
                self.case_time_frame_start_date.delete(0, tk.END)
            if hasattr(self, 'case_time_frame_end_date'):
                self.case_time_frame_end_date.delete(0, tk.END)

    # MODIFY: Update toggle_case_time_frame_fields for Case Agent
    def toggle_case_time_frame_fields(self, event=None):
        """Show/hide time frame date fields based on checkbox state (Case Agent)"""
        if self.case_time_frame_limited_var.get() == 1:
            self.case_time_frame_dates_frame.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            self.case_time_frame_dates_frame.grid_remove()
            # Clear date fields when hiding
            if hasattr(self, 'case_time_frame_start_date'):
                self.case_time_frame_start_date.delete(0, tk.END)
            if hasattr(self, 'case_time_frame_end_date'):
                self.case_time_frame_end_date.delete(0, tk.END)


    def create_device_information_frame(self):
        device_frame = ttk.LabelFrame(self.tab_device, text="Device Information", padding="10")
        device_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Owner
        ttk.Label(device_frame, text="Device Owner:").grid(row=0, column=0, sticky="w", pady=2)
        self.device_owner = ttk.Entry(device_frame)
        self.device_owner.grid(row=0, column=1, sticky="ew", pady=2)

        # Airplane Mode
        ttk.Label(device_frame, text="Airplane Mode When Received:").grid(row=1, column=0, sticky="w", pady=2)
        self.airplane_mode = ttk.Combobox(device_frame, values=['Yes', 'No'], state="readonly")
        self.airplane_mode.grid(row=1, column=1, sticky="ew", pady=2)
        self.airplane_mode.set('Yes')
        self.setup_auto_complete_dropdown(self.airplane_mode)

        # Passcode
        ttk.Label(device_frame, text="Passcode:").grid(row=2, column=0, sticky="w", pady=2)
        self.device_passcode = ttk.Entry(device_frame)
        self.device_passcode.grid(row=2, column=1, sticky="ew", pady=2)

        # Color
        ttk.Label(device_frame, text="Color:").grid(row=3, column=0, sticky="w", pady=2)
        self.device_color = ttk.Entry(device_frame)
        self.device_color.grid(row=3, column=1, sticky="ew", pady=2)

        # Capacity
        ttk.Label(device_frame, text="Capacity:").grid(row=4, column=0, sticky="w", pady=2)
        self.device_capacity = ttk.Entry(device_frame)
        self.device_capacity.grid(row=4, column=1, sticky="ew", pady=2)

        # ICCID
        ttk.Label(device_frame, text="ICCID:").grid(row=5, column=0, sticky="w", pady=2)
        self.device_iccid = ttk.Entry(device_frame)
        self.device_iccid.grid(row=5, column=1, sticky="ew", pady=2)

        # Carrier
        ttk.Label(device_frame, text="Carrier:").grid(row=6, column=0, sticky="w", pady=2)
        self.device_carrier = ttk.Entry(device_frame)
        self.device_carrier.grid(row=6, column=1, sticky="ew", pady=2)

        device_frame.columnconfigure(1, weight=1)

    def create_transfer_information_frame(self):
        transfer_frame = ttk.LabelFrame(self.tab_request, text="Transfer Information", padding="10")
        transfer_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Transfer - Changed from Combobox to Single Yes Checkbox
        ttk.Label(transfer_frame, text="Device Transferred by Another Officer:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(transfer_frame)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variable
        self.device_transfer_var = tk.IntVar(transfer_frame)
        
        # Create single Yes checkbox
        self.transfer_yes_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Yes",
            variable=self.device_transfer_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_transfer_fields
        )
        self.transfer_yes_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Set default to unchecked (0 = No)
        self.device_transfer_var.set(0)

        # Transfer Fields (initially hidden)
        self.transfer_fields = ttk.Frame(transfer_frame)
        self.transfer_fields.grid(row=1, column=0, columnspan=2, sticky="ew", pady=2)
        self.transfer_fields.grid_remove()

        ttk.Label(self.transfer_fields, text="Transfer Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.transfer_date = add_date_entry(self.transfer_fields, row=0, column=1)

        ttk.Label(self.transfer_fields, text="Transfer Officer Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.transfer_agency = ttk.Combobox(self.transfer_fields)
        self.transfer_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.transfer_agency)

        ttk.Label(self.transfer_fields, text="Transfer Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        self.transfer_title = ttk.Combobox(self.transfer_fields, values=load_request_titles())
        self.transfer_title.grid(row=2, column=1, sticky="ew", pady=2)
        setup_title_combobox(self.transfer_title)

        ttk.Label(self.transfer_fields, text="Transfer Officer Name:").grid(row=3, column=0, sticky="w", pady=2)
        self.transfer_officer = ttk.Entry(self.transfer_fields)
        self.transfer_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Configure grid weights
        transfer_frame.columnconfigure(1, weight=1)
        self.transfer_fields.columnconfigure(1, weight=1)

    def create_examiner_information_frame(self):
        examiner_frame = ttk.LabelFrame(self.tab_examiner, text="Examiner Information", padding="10")
        examiner_frame.pack(fill=tk.X, padx=5, pady=5)

        # Examiner Agency
        ttk.Label(examiner_frame, text="Examiner Agency:").grid(row=0, column=0, sticky="w", pady=2)
        examiner_agency_frame = ttk.Frame(examiner_frame)
        examiner_agency_frame.grid(row=0, column=1, sticky="ew", pady=2)

        self.examiner_agency_type = ttk.Combobox(examiner_agency_frame)
        self.examiner_agency_type.pack(fill=tk.X, expand=True)
        setup_agency_combobox(
            self.examiner_agency_type,
            saved=saved_examiner_agency(self.current_settings),
            on_change=self.auto_save_settings,
        )

        self.examiner_agency_entry = ttk.Entry(examiner_agency_frame)
        self.examiner_agency_entry.pack_forget()

        # Examiner Title
        ttk.Label(examiner_frame, text="Examiner Title:").grid(row=1, column=0, sticky="w", pady=2)
        examiner_title_frame = ttk.Frame(examiner_frame)
        examiner_title_frame.grid(row=1, column=1, sticky="ew", pady=2)

        self.examiner_title_type = ttk.Combobox(examiner_title_frame, values=load_request_titles())
        self.examiner_title_type.pack(fill=tk.X, expand=True)
        setup_title_combobox(
            self.examiner_title_type,
            saved=self.current_settings.get("examiner_title", ""),
            on_change=self.auto_save_settings,
        )

        self.examiner_title_entry = ttk.Entry(examiner_title_frame)
        self.examiner_title_entry.pack_forget()

        # Examiner Name
        ttk.Label(examiner_frame, text="Examiner Name:").grid(row=2, column=0, sticky="w", pady=2)
        self.examiner_name = ttk.Entry(examiner_frame)
        self.examiner_name.grid(row=2, column=1, sticky="ew", pady=2)
        # Load saved setting
        if self.current_settings.get("examiner_name"):
            self.examiner_name.insert(0, self.current_settings["examiner_name"])
        
        # Bind auto-save events
        self.examiner_name.bind('<KeyRelease>', self.auto_save_settings)
        self.examiner_name.bind('<FocusOut>', self.auto_save_settings)

        # Case Number
        ttk.Label(examiner_frame, text="Case Number:").grid(row=3, column=0, sticky="w", pady=2)
        self.case_number = ttk.Entry(examiner_frame)
        self.case_number.grid(row=3, column=1, sticky="ew", pady=2)
        
        #Evidence Number
        ttk.Label(examiner_frame, text="Evidence Number:").grid(row=4, column=0, sticky="w", pady=2)
        self.evidence_number = ttk.Entry(examiner_frame)
        self.evidence_number.grid(row=4, column=1, sticky="ew", pady=2)

        # DFR Report Number
        ttk.Label(examiner_frame, text="Report Number:").grid(row=5, column=0, sticky="w", pady=2)
        self.DFR_Num = ttk.Entry(examiner_frame)
        self.DFR_Num.grid(row=5, column=1, sticky="ew", pady=2)
        
        # Load saved prefix or use default
        self.DFR_Num.insert(0, current_dfr_prefix())

        examiner_frame.columnconfigure(1, weight=1)


    def toggle_examiner_title_entry(self, event=None):
        if self.examiner_title_type.get() == "Other (specify)":
            self.examiner_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.examiner_title_entry.pack_forget()
            self.examiner_title_entry.delete(0, tk.END)
        
        # Auto-save when selection changes
        self.auto_save_settings()

    def on_agency_type_changed(self, event=None):
        """Handle agency type change and auto-save"""
        # Call the original toggle method
        self.toggle_agency_entry(event)
        # Auto-save the change
        self.auto_save_settings()

    def auto_save_settings(self, event=None):
        """Automatically save examiner settings when any field changes"""
        # Use after_idle to prevent excessive saving during rapid typing
        if hasattr(self, '_save_timer') and self._save_timer:
            self.after_cancel(self._save_timer)
        
        # Schedule save for after user stops typing (500ms delay)
        self._save_timer = self.after(500, self._perform_auto_save)

    def _perform_auto_save(self):
        try:
            # Get examiner title value based on dropdown selection
            examiner_title_value = self.examiner_title_type.get()
                
            settings_to_save = {
                "examiner_agency_type": self.examiner_agency_type.get(),
                "examiner_agency_custom": "",
                "examiner_title": examiner_title_value,
                "examiner_name": self.examiner_name.get(),
                "dfr_number_prefix": self.get_dfr_prefix(),
                "version": "1.0.1"
            }
            
            self.settings_manager.save_settings(settings_to_save)
            # Update current_settings to reflect the save
            self.current_settings.update(settings_to_save)
            remember_agencies_from_form(self)
            
        except Exception as e:
            print(f"Error auto-saving examiner settings: {e}")

    def get_dfr_prefix(self):
        """Extract the DFR prefix from the current DFR number field"""
        dfr_value = self.DFR_Num.get()
        # Find the last occurrence of "DFR" and include everything up to and including any year/dash pattern
        import re
        match = re.match(r'(DFR\d{4}-)', dfr_value)
        if match:
            return match.group(1)
        else:
            # Fallback to current year
            from datetime import datetime
            current_year = datetime.now().year
            return f"DFR{current_year}-"

    def create_middle_column_widgets(self):
        self.create_device_information_frame()
        self.create_forensic_software_frame()
        self.create_output_file_frame()

    def create_right_column_widgets(self):
        self.create_file_upload_frame()

    def create_forensic_software_frame(self):
        forensic_software_frame = ttk.LabelFrame(self.tab_device, text="Forensic Software", padding="10")
        forensic_software_frame.pack(fill=tk.X, padx=5, pady=5)

        forensic_software_content = ttk.Frame(forensic_software_frame)
        forensic_software_content.pack(fill=tk.X, expand=True)

        # Forensic Processing Software - Using checkboxes all on same line
        ttk.Label(forensic_software_content, text="Forensic Processing Software:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(forensic_software_content)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variables - matching artifacts checkbox style exactly
        self.axiom_var = tk.IntVar(forensic_software_frame)
        self.cellebrite_var = tk.IntVar(forensic_software_frame)
        self.griffeye_var = tk.IntVar(forensic_software_frame)
        
        # Create checkboxes exactly like artifacts checkboxes
        self.axiom_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Axiom",
            variable=self.axiom_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.axiom_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        self.cellebrite_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Cellebrite",
            variable=self.cellebrite_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.cellebrite_checkbox.grid(row=0, column=1, sticky="w", padx=(0, 10))
        
        self.griffeye_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Griffeye",
            variable=self.griffeye_var,
            onvalue=1,
            offvalue=0,
            command=self.on_forensic_software_change
        )
        self.griffeye_checkbox.grid(row=0, column=2, sticky="w", padx=(0, 10))
        
        # Don't set any default selection - all checkboxes start unchecked
        
        # Create artifacts button frame
        self.artifacts_button_frame = ttk.Frame(forensic_software_content)
        self.artifacts_button_frame.grid(row=1, column=0, columnspan=2, sticky="w", pady=5)

        self.selected_artifacts = []
        self.select_artifacts_button = ttk.Button(
            self.artifacts_button_frame, 
            text="Select Artifacts", 
            command=self.open_artifacts_popup
        )

        # Don't show artifacts button initially since nothing is selected by default
        # self.select_artifacts_button.pack(side=tk.LEFT, padx=5)

        forensic_software_content.columnconfigure(1, weight=1)

    def on_forensic_software_change(self):
        # Show/hide artifacts button based on Axiom selection
        if self.axiom_var.get() == 1:
            self.select_artifacts_button.pack(side=tk.LEFT, padx=5)
        else:
            self.select_artifacts_button.pack_forget()
            # Clear selected artifacts when Axiom is unchecked
            self.selected_artifacts = []
        self.refresh_tab_status()
            
    def get_selected_forensic_software(self):
        """Get list of selected forensic software"""
        selected = []
        if self.axiom_var.get() == 1:
            selected.append("Axiom")
        if self.cellebrite_var.get() == 1:
            selected.append("Cellebrite")
        if self.griffeye_var.get() == 1:
            selected.append("Griffeye")
        return selected

    def create_file_upload_frame(self):
        file_frame = ttk.LabelFrame(self.right_frame, text="File Upload", padding="10")
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        add_template_picker(self, file_frame, preferred="DFR Mobile (2026).docx", keywords=("mobile",))

        # Extraction File Drag-and-Drop
        self.extraction_drop_label = ttk.Label(file_frame, text="Drag & Drop GrayKey PDF, Cellebrite Summary/Quick View PDF or UFD File — or click to browse",
                                      relief="solid", borderwidth=2, padding=10)
        self.extraction_drop_label.pack(fill=tk.X, pady=10, padx=5)
        self.extraction_drop_label.bind('<Button-1>', self.browse_extraction_file)
        self.extraction_drop_label.drop_target_register(DND_FILES)
        self.extraction_drop_label.dnd_bind('<<Drop>>', self.drop_extraction_file)

    def open_artifacts_popup(self):
        if hasattr(self, 'extraction_file') and self.extraction_file:
            extraction_data = self.parse_extraction_file()
            device_type = self.determine_device_type(extraction_data)
            self.selected_artifacts = self.select_artifacts_popup(device_type)
            if self.selected_artifacts:
                messagebox.showinfo("Artifacts Selected", f"{len(self.selected_artifacts)} artifacts selected successfully.")
        else:
            messagebox.showwarning("No Extraction File", "Please upload an extraction file first to determine device type.")

    def create_output_file_frame(self):
        output_frame = ttk.LabelFrame(self.tab_output, text="Output File", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        output_content = ttk.Frame(output_frame)
        output_content.pack(fill=tk.X, expand=True)
        
        # Output filename
        ttk.Label(output_content, text="Output File Name:").grid(row=0, column=0, sticky="w", pady=2)
        self.output_filename = ttk.Entry(output_content)
        self.output_filename.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Hint for filename (directly under the filename field)
        hint_style = ttk.Style()
        hint_style.configure("Hint.TLabel", font=('Arial', 7))  
        hint_label = ttk.Label(output_content, 
                              text="*Default File Name \"(DFR #) - (Owner Name) (Device Model)\"", 
                              style="Hint.TLabel",
                              wraplength=350)
        hint_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 5))
        
        # Save location
        ttk.Label(output_content, text="Save Location:").grid(row=2, column=0, sticky="w", pady=2)
        
        # Frame for save location entry and browse button
        location_frame = ttk.Frame(output_content)
        location_frame.grid(row=2, column=1, sticky="ew", pady=2)
        location_frame.columnconfigure(0, weight=1)
        
        # Save location entry (initialize with desktop path)
        self.save_location = ttk.Entry(location_frame)
        self.save_location.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        # Set default to desktop
        self.save_location.insert(0, default_export_dir())
        
        # Browse button
        browse_button = ttk.Button(location_frame, text="Browse", command=self.browse_save_location)
        browse_button.grid(row=0, column=1)
        
        output_content.columnconfigure(1, weight=1)


    def browse_save_location(self):
        """Allow user to browse and select a save location"""
        try:
            # Get current path from the entry, default to desktop if empty
            current_path = self.save_location.get().strip()
            selected_path = ask_directory("export", current_path, "Select Save Location")
            if selected_path:
                self.save_location.delete(0, tk.END)
                self.save_location.insert(0, selected_path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not browse for save location: {str(e)}")


    def toggle_request_content(self, event=None):
        # Clear all request information fields when role is changed
        if hasattr(self, 'request_date'):
            self.request_date.delete(0, tk.END)
        if hasattr(self, 'request_agency'):
            self.request_agency.delete(0, tk.END)
        if hasattr(self, 'request_title_type'):
            self.request_title_type.set("")
        if hasattr(self, 'request_title_entry'):
            self.request_title_entry.delete(0, tk.END)
            self.request_title_entry.pack_forget()
        if hasattr(self, 'request_officer'):
            self.request_officer.delete(0, tk.END)
        if hasattr(self, 'case_type'):
            self.case_type.delete(0, tk.END)
        if hasattr(self, 'legal_authority'):
            self.legal_authority.set('Search Warrant')
        if hasattr(self, 'offense_type'):
            self.offense_type.delete(0, tk.END)
        if hasattr(self, 'legal_self'):
            self.legal_self.set('Search Warrant')
        if hasattr(self, 'sw_service_date'):
            self.sw_service_date.delete(0, tk.END)
        
        # Clear time frame fields for both Agency Assist and Case Agent
        if hasattr(self, 'time_frame_limited_var'):
            self.time_frame_limited_var.set(0)
        if hasattr(self, 'time_frame_start_date'):
            self.time_frame_start_date.delete(0, tk.END)
        if hasattr(self, 'time_frame_end_date'):
            self.time_frame_end_date.delete(0, tk.END)
        if hasattr(self, 'case_time_frame_limited_var'):
            self.case_time_frame_limited_var.set(0)
        if hasattr(self, 'case_time_frame_start_date'):
            self.case_time_frame_start_date.delete(0, tk.END)
        if hasattr(self, 'case_time_frame_end_date'):
            self.case_time_frame_end_date.delete(0, tk.END)
        
        # Hide search warrant date fields when switching roles
        if hasattr(self, 'sw_date_frame'):
            for widget in self.sw_date_frame.winfo_children():
                widget.grid_remove()
        
        # Hide time frame sections
        if hasattr(self, 'time_frame_dates_frame'):
            self.time_frame_dates_frame.grid_remove()
        if hasattr(self, 'case_time_frame_dates_frame'):
            self.case_time_frame_dates_frame.grid_remove()
        
        # Now show/hide the appropriate frame
        if self.role_type.get() == "Agency Assist":
            self.agency_assist_frame.pack(fill=tk.X, padx=5, pady=5)
            self.case_agent_frame.pack_forget()
            # Show time frame section since default is Search Warrant
            if hasattr(self, 'time_frame_label'):
                self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)
                self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)
        else:  # Case Agent
            self.agency_assist_frame.pack_forget()
            self.case_agent_frame.pack(fill=tk.X, padx=5, pady=5)
            # Show the SW date field and time frame section if Search Warrant is selected
            self.toggle_sw_date()
            self.toggle_case_time_frame_section()

    def toggle_sw_date(self, event=None):
        if hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant':
            # Show the SW date frame
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            # Hide the SW date frame and clear its value
            if hasattr(self, 'sw_date_frame'):
                self.sw_date_frame.grid_remove()
                
                # Clear the SW service date field when hiding it
                if hasattr(self, 'sw_service_date'):
                    self.sw_service_date.delete(0, tk.END)

    def toggle_title_entry(self, event=None):
        if self.request_title_type.get() == "Other (specify)":
            self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.request_title_entry.pack_forget()

    def toggle_agency_entry(self, event=None):
        return

    def toggle_transfer_fields(self, event=None):
        if self.device_transfer_var.get() == 1:
            self.transfer_fields.grid()
        else:
            self.transfer_fields.grid_remove()
        self.refresh_tab_status()

    def setup_auto_complete_dropdown(self, combobox):
        combobox.bind('<KeyRelease>', lambda event, cb=combobox: self.auto_complete(event, cb))

    def auto_complete(self, event, combobox):
        if event.keysym in ('Up', 'Down', 'Left', 'Right', 'Return', 'Tab', 'BackSpace', 'Delete'):
            return
        
        value = combobox.get()
        
        if not value:
            return
        
        all_values = combobox['values']
        
        previous_value = combobox.get()
        
        for option in all_values:
            if option.lower().startswith(value.lower()):
                combobox.set(option)
                combobox.icursor(len(option))
                combobox.selection_range(len(value), len(option))
                                
                break

    def drop_extraction_file(self, event):
        file_paths = list(self.tk.splitlist(event.data))
        if not file_paths:
            return
        remember_folder("extraction", file_paths[0])
        self._load_extraction_files(file_paths)

    def drop_template_file(self, event):
        file_path = self.tk.splitlist(event.data)[0]
        if file_path.endswith('.docx'):
            self.template_file = file_path
            sync_template_choice(self, file_path)
            remember_folder("template", file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        else:
            messagebox.showerror("Error", "Please select a Word document template")

    def browse_extraction_file(self, event):
        file_path = ask_open_file(
            [("Extraction files", "*.ufd;*.pdf"), ("UFD files", "*.ufd"), ("PDF files", "*.pdf")],
            folder_kind="extraction",
            title="Select extraction file",
        )
        if file_path:
            self._load_extraction_files([file_path])
            
    def browse_template_file(self, event):
        file_path = ask_open_file(
            [("Word documents", "*.docx")],
            folder_kind="template",
            title="Select DFR template",
        )
        if file_path:
            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")

    def add_bold_underline_paragraph(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.underline = True
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        return p

    def _load_extraction_files(self, file_paths):
        pdfs = [path for path in file_paths if path.lower().endswith(".pdf")]
        ufds = [path for path in file_paths if path.lower().endswith(".ufd")]
        if ufds:
            self.extraction_file = ufds[0]
            self.extraction_type = "Cellebrite"
            self.cellebrite_companion_file = None
            self.extraction_drop_label.configure(text=f"Selected: {os.path.basename(self.extraction_file)} (Cellebrite UFD)")
            extraction_data = self.parse_ufd_file()
            if not pdfs:
                if messagebox.askyesno(
                    "Cellebrite PDFs",
                    "A UFD was loaded. Do you also have a Cellebrite Summary Report or Quick View PDF to add device identifiers (IMEI, ICCID, phone, carrier)?",
                ):
                    extra = ask_open_file(
                        [("PDF files", "*.pdf")],
                        folder_kind="extraction",
                        title="Select Cellebrite Summary or Quick View PDF",
                    )
                    if extra:
                        pdfs = [extra]
            if pdfs:
                extraction_data = self._merge_ufd_with_pdfs(extraction_data, pdfs)
                names = [os.path.basename(self.extraction_file)] + [os.path.basename(p) for p in pdfs]
                self.extraction_drop_label.configure(text=f"Selected: {', '.join(names)} (UFD + Cellebrite PDF)")
            self.update_info_display(data=extraction_data)
            apply_suggested_filename(self, "Mobile", extraction_data.get("device_model", ""))
            response = messagebox.askyesno("Multiple Extractions", "Were additional extractions performed?")
            self.multiple_extractions = response
            if response:
                messagebox.showinfo("Note", "UPDATE OUTPUT FILE TO INCLUDE ADDITIONAL EXTRACTIONS")
            return
        if pdfs:
            self._cellebrite_companion_asked = False
            extraction_data = process_pdf_selection(self, pdfs)
            self.update_info_display(data=extraction_data)
            apply_suggested_filename(self, "Mobile", (extraction_data or {}).get("device_model", ""))
            return
        messagebox.showerror("Error", "Please select a UFD, GrayKey PDF, or Cellebrite Summary/Quick View PDF")
        self.update_info_display("Error: Please select a valid file format (UFD or PDF).")

    def _merge_ufd_with_pdfs(self, ufd_data, pdf_paths):
        pdf_data, companions = parse_cellebrite_pdfs_only(pdf_paths)
        if companions:
            self.cellebrite_companion_file = companions[0]
        merged = prefer_gui_over_parsed(ufd_data, pdf_data)
        from cellebrite_pdf import _fill_empty_device_fields
        _fill_empty_device_fields(self, merged)
        return merged

    def parse_extraction_file(self):
        if self.extraction_type == "Cellebrite":
            extraction_file = getattr(self, "extraction_file", "") or ""
            companion = getattr(self, "cellebrite_companion_file", None)
            if extraction_file.lower().endswith(".ufd"):
                data = self.parse_ufd_file()
                pdfs = [path for path in (companion, getattr(self, "cellebrite_companion_file_2", None)) if path]
                if pdfs:
                    data = self._merge_ufd_with_pdfs(data, pdfs)
                return data
            if extraction_file.lower().endswith(".pdf") or companion:
                return process_pdf_selection(self, [extraction_file] if extraction_file else [])
            return self.parse_ufd_file()
        elif self.extraction_type == "GrayKey":
            return self.parse_pdf_file()
        else:
            raise ValueError("Unknown extraction type")
            
    def parse_ufd_file(self):
        try:
            extraction_data = {
                'extraction_date': '',
                'extraction_tool': '',
                'extraction_type': '',
                'device_model': '',
                'Device_OS': ''
            }
            
            with open(self.extraction_file, "r") as file:
                for line in file:
                    try:
                        if '=' in line:
                            key, value = line.strip().split('=', 1)
                            if key == 'Date':
                                date_time_parts = value.split()
                                if len(date_time_parts) >= 3:  
                                    date_part = date_time_parts[0]
                                    full_time_part = date_time_parts[1]
                                    time_parts = full_time_part.split(':')
                                    time_part = f"{time_parts[0]}:{time_parts[1]}"
                                    
                                    # Parse the date - UFD files use D/M/Y format
                                    date_parts = date_part.split('/')
                                    if len(date_parts) == 3:
                                        day, month, year = date_parts  # Note the order: D/M/Y
                                        
                                        # Convert to datetime object
                                        date_obj = datetime(int(year), int(month), int(day), 
                                                        int(time_parts[0]), int(time_parts[1]))
                                        
                                        # Use the system's local timezone instead of the one in the file
                                        import time
                                        
                                        # Determine if DST is in effect for this date
                                        target_timestamp = date_obj.timestamp()
                                        is_dst = time.localtime(target_timestamp).tm_isdst > 0
                                        
                                        # Get the local timezone information
                                        # Get local UTC offset in seconds
                                        if is_dst:
                                            local_utc_offset = -time.altzone / 3600  # in hours
                                        else:
                                            local_utc_offset = -time.timezone / 3600  # in hours
                                        
                                        # Get the timezone name based on local settings
                                        tz_name = self.get_local_timezone_name(local_utc_offset, is_dst)
                                        
                                        # Format the date with timezone
                                        formatted_date = date_obj.strftime("%A, %B %d, %Y").replace(' 0', ' ')
                                        extraction_data['formatted_date'] = f"{formatted_date} at {time_part} {tz_name}"
                            elif key == 'AcquisitionTool':
                                extraction_data['extraction_tool'] = value
                            elif key == 'ExtractionType':
                                extraction_data['extraction_type'] = value
                            elif key == 'Model':
                                extraction_data['device_model'] = value
                            elif key == 'Vendor':
                                extraction_data['device_manufacturer'] = value.title()
                            elif key == 'OS':  
                                extraction_data['Device_OS'] = value
                            elif key == 'InternalBuild':
                                extraction_data['cellebrite_version'] = value
                    except ValueError:
                        continue
                        
            # Determine the appropriate article for the manufacturer
            manufacturer = extraction_data.get('device_manufacturer', 'Unknown')
            extraction_data['article'] = self.get_article_for_manufacturer(manufacturer)
                
            return extraction_data
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to parse UFD file: {str(e)}")
            return {}

    def get_local_timezone_name(self, utc_offset, is_dst):
        import time
        
        # Get the local timezone name directly from the system
        timezone_name = time.tzname[1] if is_dst else time.tzname[0]
        
        # Dictionary of timezone abbreviations to full names
        timezone_names = {
            "PST": "Pacific Standard Time",
            "PDT": "Pacific Daylight Time",
            "MST": "Mountain Standard Time",
            "MDT": "Mountain Daylight Time", 
            "CST": "Central Standard Time",
            "CDT": "Central Daylight Time",
            "EST": "Eastern Standard Time",
            "EDT": "Eastern Daylight Time",
            "HST": "Hawaii Standard Time",
            "AKST": "Alaska Standard Time",
            "AKDT": "Alaska Daylight Time",
            "AST": "Atlantic Standard Time",
            "ADT": "Atlantic Daylight Time"
        }
        
        return timezone_names.get(timezone_name, timezone_name)

    def extract_info_from_pdf(self, pdf_path):
        extracted_info = {}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                pdf_text = ""
                for page in pdf_reader.pages:
                    pdf_text += page.extract_text()
                
                header_info = self.extract_header_info(pdf_text)
                extracted_info.update(header_info)
                
                device_info = self.extract_device_table_info(pdf_text)
                extracted_info.update(device_info)
                
                event_log_info = self.extract_event_log_info(pdf_text)
                extracted_info.update(event_log_info)
                
                extraction_type_info = self.extract_extraction_type(pdf_text)
                extracted_info.update(extraction_type_info)
                
                return extracted_info
        
        except Exception as e:
            messagebox.showerror("PDF Processing Error", f"Error processing PDF: {e}")
            return extracted_info

    def extract_header_info(self, pdf_text):
        """Pull GrayKey/Graykey OS version from old and new Progress Report layouts."""
        header_info = {}
        if not pdf_text:
            return header_info

        version_patterns = [
            # New format (2026): "Graykey Software: OS Version: 1.29.0.35291825, App Bundle: ..."
            r'Gray\s*[Kk]ey\s+Software:\s*OS Version:\s*([0-9]+(?:\.[0-9]+)*)',
            # Older format: "GrayKey Software:" then OS Version on the same or following line
            r'Gray\s*[Kk]ey\s+Software:.*?OS Version:\s*([0-9]+(?:\.[0-9]+)*)',
            r'Gray\s*[Kk]ey\s+Software Version[:\s]+([0-9]+(?:\.[0-9]+)*)',
            r'Gray\s*[Kk]ey\s+OS Version[:\s]+([0-9]+(?:\.[0-9]+)*)',
            # Event-log fallback on both layouts
            r'On-device agent started[\s\S]{0,240}?OS Version:\s*([0-9]+(?:\.[0-9]+)*)',
        ]

        flags = re.IGNORECASE | re.DOTALL
        for pattern in version_patterns:
            software_match = re.search(pattern, pdf_text, flags)
            if software_match:
                header_info['GrayKey_OS'] = software_match.group(1).strip()
                break

        return header_info

    def extract_device_table_info(self, pdf_text):
        device_info = {}
        device_section_start = pdf_text.find("Target Device Information")
        if device_section_start == -1:
            return device_info
        
        # Extract the section after "Target Device Information" and before "Event Log"
        event_log_start = pdf_text.find("Event Log", device_section_start)
        if event_log_start == -1:
            device_section = pdf_text[device_section_start:]
        else:
            device_section = pdf_text[device_section_start:event_log_start]
        
        # Look for specific key-value pairs
        key_value_patterns = [
            (r'Software Version\s+(.*?)(?:\n|$)', 'Device_OS'),
            (r'Model\s+(.*?)(?:\n|$)', 'Device_Model'),
            (r'Serial Number\s+(.*?)(?:\n|$)', 'Serial_Number'),
            (r'Phone Number\s+(.*?)(?:\n|$)', 'Phone_Number'),
            (r'IMEI\s+(.*?)(?:\n|$)', 'DEV_IMEI'),
            (r'Device Name\s+(.*?)(?:\n|$)', 'Device_Name'),
            (r'Accounts\s+(.*?)(?:\n|$)', 'Device_Account'),
        ]
        
        for pattern, key in key_value_patterns:
            match = re.search(pattern, pdf_text)
            if match:
                value = match.group(1).strip()
                if key == 'Phone_Number':
                    value = self.format_phone_number(value)
                device_info[key] = value
        
        return device_info

    def format_phone_number(self, phone_str):
        if not phone_str:
            return phone_str
        
        # Remove all non-digit characters
        digits_only = re.sub(r'\D', '', phone_str)
        
        # If 11 digits and starts with 1, remove the leading 1
        if len(digits_only) == 11 and digits_only.startswith('1'):
            digits_only = digits_only[1:]
        
        # If we have exactly 10 digits, format them
        if len(digits_only) == 10:
            return f"{digits_only[:3]}-{digits_only[3:6]}-{digits_only[6:]}"
        
        # If not 10 digits after processing, return the original
        return phone_str

    def extract_event_log_info(self, pdf_text):
        event_log_info = {}
        
        # Find the start of the Event Log section
        event_log_start = pdf_text.find("Event Log")
        if event_log_start == -1:
            return event_log_info
        
        # Extract the section after "Event Log"
        event_log_section = pdf_text[event_log_start:]
        
        # Find the first event log entry (timestamp)
        first_event_match = re.search(r'Event Log\s*\n([0-9]{4}-[0-9]{2}-[0-9]{2})\s+([0-9]{2}:[0-9]{2}):[0-9]{2}\s+UTC', event_log_section)
        if first_event_match:
            # Extract components and reformat
            date_part = first_event_match.group(1).strip()  # YYYY-MM-DD
            time_part = first_event_match.group(2).strip()  # HH:MM
            
            # Parse the date into components
            year, month, day = date_part.split('-')
            hour, minute = time_part.split(':')
            
            # Convert to datetime object
            date_obj = datetime(int(year), int(month), int(day), int(hour), int(minute))
            
            # Use the system's local timezone instead of calculating from UTC offset
            import time
            
            # Determine if DST is in effect for this date
            target_timestamp = date_obj.timestamp()
            is_dst = time.localtime(target_timestamp).tm_isdst > 0
            
            # Get the local timezone information
            # Get local UTC offset in seconds
            if is_dst:
                local_utc_offset = -time.altzone / 3600  # in hours
            else:
                local_utc_offset = -time.timezone / 3600  # in hours
            
            # Convert UTC time to local time
            local_hour = int(hour) + int(local_utc_offset)
            
            # Handle day changes when crossing midnight
            local_day = int(day)
            if local_hour < 0:
                local_hour += 24
                # Need to adjust the day backward
                from datetime import timedelta
                prev_day = date_obj - timedelta(days=1)
                local_day = prev_day.day
                local_month = prev_day.month
                local_year = prev_day.year
            elif local_hour >= 24:
                local_hour -= 24
                # Need to adjust the day forward
                from datetime import timedelta
                next_day = date_obj + timedelta(days=1)
                local_day = next_day.day
                local_month = next_day.month
                local_year = next_day.year
            else:
                local_month = int(month)
                local_year = int(year)
            
            # Get the timezone name based on local settings
            tz_name = self.get_local_timezone_name(local_utc_offset, is_dst)
            
            # Format the date with proper DST handling - now with the same format as UFD files
            local_date_obj = datetime(local_year, local_month, local_day, local_hour, int(minute))
            formatted_date = local_date_obj.strftime("%A, %B %d, %Y").replace(' 0', ' ')
            event_log_info['First_Event_Time'] = f"{formatted_date} at {local_hour:02d}:{minute} {tz_name}"
        
        return event_log_info

    def extract_extraction_type(self, pdf_text):
        extraction_info = {}
        
        # Find the start of the Extraction Result Summary section
        summary_start = pdf_text.find("Extraction Result Summary")
        if summary_start == -1:
            return extraction_info
        
        # Look for the first line after the heading which contains the extraction type
        match = re.search(r'Extraction Result Summary\s*\n(.*?)(?:\s+\d{4}-\d{2}-\d{2}|\s+Extraction size)', pdf_text[summary_start:])
        if match:
            extraction_info['Extraction_Type'] = match.group(1).strip()
        
        return extraction_info

    def parse_pdf_file(self):
        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            messagebox.showerror("Error", "No PDF file selected")
            return {}
        
        # Extract data from PDF
        pdf_data = self.extract_info_from_pdf(self.extraction_file)
        
        # Debug: Print the keys in pdf_data to see what's being extracted
        print("Keys in pdf_data after extraction:", list(pdf_data.keys()))
        
        # Set default values for extraction tool and type
        pdf_data['extraction_tool'] = "GrayKey"

        # Only set a default if the extraction type wasn't found
        if 'Extraction_Type' in pdf_data:
            pdf_data['extraction_type'] = pdf_data['Extraction_Type']
            # Remove duplicate to prevent double display
            del pdf_data['Extraction_Type']
        else:
            pdf_data['extraction_type'] = "data" 
        
        # Get the first event time for the formatted date
        if 'First_Event_Time' in pdf_data:
            pdf_data['formatted_date'] = pdf_data['First_Event_Time']
            # Remove duplicate to prevent double display
            del pdf_data['First_Event_Time']
        else:
            pdf_data['formatted_date'] = "(INSERT DATE)"
        
        # Try to determine manufacturer from model
        if 'Device_Model' in pdf_data:
            device_model = pdf_data['Device_Model']
            pdf_data['device_model'] = device_model
            # Remove duplicate to prevent double display
            del pdf_data['Device_Model']
            
            # Dictionary mapping manufacturer keywords to manufacturer names
            manufacturer_keywords = {
                'apple': ['iphone', 'ipad', 'ipod'],
                'samsung': ['samsung', 'galaxy', 'sm-', 'gt-'],
                'google': ['google', 'pixel', 'nexus'],
                'lg': ['lg', 'lm-', 'vs', 'h815', 'h811', 'h870'],
                'motorola': ['motorola', 'moto', 'xt', 'droid']
            }
            
            # Check for each manufacturer's keywords in the device model
            device_model_lower = device_model.lower()
            found_manufacturer = False
            
            for manufacturer, keywords in manufacturer_keywords.items():
                if any(keyword in device_model_lower for keyword in keywords):
                    pdf_data['device_manufacturer'] = manufacturer.capitalize()
                    found_manufacturer = True
                    break
            
            # If no specific manufacturer was identified
            if not found_manufacturer:
                pdf_data['device_manufacturer'] = "Unknown"
        else:
            pdf_data['device_model'] = "Unknown device"
            pdf_data['device_manufacturer'] = "Unknown"
        
        # Make sure Device_OS has a value if not already set
        if 'Device_OS' in pdf_data:
            # Keep the existing value but don't create duplicate
            pass
        else:
            pdf_data['Device_OS'] = "Unknown OS"
            
        # Determine the appropriate article for the manufacturer
        manufacturer = pdf_data.get('device_manufacturer', 'Unknown')
        pdf_data['article'] = self.get_article_for_manufacturer(manufacturer)

        # Debug: Print the keys in pdf_data after processing
        print("Keys in pdf_data after processing:", list(pdf_data.keys()))
        
        return pdf_data

    def get_article_for_manufacturer(self, manufacturer):
        # List of manufacturers that should use 'an'
        an_manufacturers = ['apple', 'lg', 'htc', 'alcatel']
    
        # Convert to lowercase for comparison
        manufacturer_lower = manufacturer.lower()
    
        # Check if manufacturer starts with a vowel sound
        if manufacturer_lower in an_manufacturers or manufacturer_lower[0] in 'aeiou':
            return 'an'
        else:
            return 'a'

    def create_info_display(self):
        self.info_frame, self.info_display = build_extracted_info_pane(
            self.right_frame,
            "Add a UFD or PDF file to see the extracted information.",
        )
        self.no_evidence_var = tk.IntVar(self)
        self.no_evidence_checkbox = ttk.Checkbutton(
            self.info_frame,
            text="No Evidence Found",
            variable=self.no_evidence_var,
            onvalue=1,
            offvalue=0,
        )
        self.no_evidence_checkbox.pack(anchor="w", pady=(6, 0))

    def update_info_display(self, message=None, data=None):
        self.info_display.config(state=tk.NORMAL)
        self.info_display.delete(1.0, tk.END)
        
        if message:
            self.info_display.insert(tk.END, message)
        
        if data:
            self.info_display.insert(tk.END, "Extracted Information:\n\n")
            
            display_data = data.copy()
            
            exclude_keys = ['article']
            
            for key in exclude_keys:
                if key in display_data:
                    del display_data[key]
            
            is_cellebrite = self.extraction_type == "Cellebrite" if hasattr(self, "extraction_type") else False
            
            if is_cellebrite:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('extraction_tool', 'Extraction Tool'),
                    ('cellebrite_version', 'Cellebrite Version'),
                    ('extraction_type', 'Extraction Type'),
                    ('device_manufacturer', 'Device Manufacturer'),
                    ('device_model', 'Device Model'),
                    ('Device_OS', 'Device OS'),
                    ('Device_Name', 'Device Name'),
                    ('DEV_IMEI', 'Device IMEI'),
                    ('Phone_Number', 'Device Phone Number'),
                    ('device_iccid', 'ICCID'),
                    ('device_carrier', 'Carrier'),
                    ('device_passcode', 'Passcode'),
                    ('Device_Account', 'Account(s)'),
                ]
            else:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('extraction_tool', 'Extraction Tool'),
                    ('GrayKey_OS', 'GrayKey OS'),
                    ('extraction_type', 'Extraction Type'),
                    ('device_manufacturer', 'Device Manufacturer'),
                    ('device_model', 'Device Model'),
                    ('Serial_Number', 'Device Serial Number'),
                    ('Device_OS', 'Device OS'),
                    ('Device_Name', 'Device Name'),
                    ('DEV_IMEI', 'Device IMEI'),
                    ('Phone_Number', 'Device Phone Number'),
                    ('Device_Account', 'Account(s)'),
                ]
            
            displayed_keys = set()
            
            for key_pair in display_order:
                key, display_name = key_pair
                if key in display_data and display_data[key] and key not in displayed_keys:
                    value = display_data[key]
                    self.info_display.insert(tk.END, f"{display_name}: ", "key")
                    self.info_display.insert(tk.END, f"{value}\n", "value")
                    displayed_keys.add(key)
            
            additional_fields = [key for key in display_data.keys() 
                               if key not in displayed_keys and 
                               not key.startswith('_') and 
                               display_data[key]]
            
            if additional_fields:
                self.info_display.insert(tk.END, "\nAdditional Information:\n", "section")
                for key in sorted(additional_fields):
                    value = display_data[key]
                    display_key = key.replace('_', ' ').title()
                    self.info_display.insert(tk.END, f"{display_key}: ", "key")
                    self.info_display.insert(tk.END, f"{value}\n", "value")
            
            self.info_display.tag_configure("key", font=("Arial", 10, "bold"))
            self.info_display.tag_configure("value", font=("Arial", 10))
            self.info_display.tag_configure("section", font=("Arial", 11, "bold"))
        
        self.info_display.config(state=tk.DISABLED)

    def browse_template_file(self, event):
        try:
            file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx")])
            if file_path:
                # Verify the file exists and is readable
                with open(file_path, 'rb') as test_file:
                    # Just checking if we can open it
                    pass
                self.template_file = file_path
                sync_template_choice(self, file_path)
                self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        except IOError as e:
            messagebox.showerror("File Error", f"Could not read the selected file: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

### PARAGRAPHS ###      
    def initialize_mobile_exam_paragraphs(self):
        self.paragraphs = load_paragraphs("mobile_full")

    def determine_device_type(self, data):
        # Check manufacturer and model info
        manufacturer = data.get('device_manufacturer', '').lower()
        model = data.get('device_model', '').lower()
        os = data.get('Device_OS', '').lower()
        # Check for Apple devices
        if 'apple' in manufacturer or 'iphone' in model or 'ipad' in model or 'ipod' in model or 'ios' in os:
            return "iOS"
        # Check for Android devices
        elif any(android_term in os for android_term in ['android', 'samsung']):
            return "Android"
        # Default to Android if can't determine (you can change this default if needed)
        else:
            return "Android"

    def select_artifacts_popup(self, device_type):
        artifacts = []
        previously_selected = self.selected_artifacts if hasattr(self, 'selected_artifacts') else []
        
        # Determine which artifact list to show based on device type
        if device_type == "iOS":
            # Define Left Column iOS Artifacts
            left_artifacts = [
                "iOS Device Information",
                "Owner Information",
                "Apple Accounts",
                "Google Accounts",
                "SIM Card Activity",
                "Chrome Downloads",
                "Chrome Keyword Search Terms",
                "Chrome Web History",
                "Chrome Web Visits",
                "Apple Maps Trips",
                "Cached Locations",
                "Significant Locations",
            ]
            
            # Define Right Column iOS Artifacts
            right_artifacts = [
                "Apple Notes",
                "Apple Notes - Voice",
                "Live Photos",
                "Pictures",
                "Videos"
            ]
        else:  # Android
            # Define Left Column Android Artifacts
            left_artifacts = [
                "Android Device Information",
                "Google Accounts",
                "Accounts Information",
                "Android Downloads",
                "Chrome Downloads",
                "Chrome Keyword Search Terms",
                "Chrome Web History",
                "Chrome Web Visits",
                "Google Maps Search History",
            ]
            
            # Define Right Column Android Artifacts
            right_artifacts = [
                "Samsung Notes",
                "Pictures",
                "Videos"
            ]
        
        # Create artifact selection dialog
        selection_window = tk.Toplevel(self)
        selection_window.title(f"Select {device_type} Artifacts to Include")
        
        # Calculate screen dimensions and window position
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        # Set width and height for the popup window
        popup_width = 500
        popup_height = 500
        
        # Calculate position to center the window
        x_position = (screen_width - popup_width) // 2
        y_position = (screen_height - popup_height) // 2
        
        # Set window geometry: width x height + x_position + y_position
        selection_window.geometry(f"{popup_width}x{popup_height}+{x_position}+{y_position}")
        
        # Make window modal
        selection_window.transient(self)
        selection_window.grab_set()
        
        # Create main container frame
        main_frame = ttk.Frame(selection_window)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Add instructional text at the top
        instruction_frame = ttk.Frame(main_frame)
        instruction_frame.pack(fill="x", pady=10)

        instruction_text = ttk.Label(
            instruction_frame,
            text="Select from these common artifacts tagged during your examination:",
            wraplength=400,
            justify="left",
            font=("Arial", 10)
        )
        instruction_text.pack(anchor="w")

        # Add a separator for visual clarity
        separator = ttk.Separator(main_frame, orient="horizontal")
        separator.pack(fill="x", pady=10)
        
        # Create a single scrollable frame for both columns
        canvas_frame = ttk.Frame(main_frame)
        canvas_frame.pack(fill="both", expand=True)
        
        canvas = tk.Canvas(canvas_frame)
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Pack the canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Create a frame inside the scrollable area to hold the two columns
        columns_container = ttk.Frame(scrollable_frame)
        columns_container.pack(fill="both", expand=True)
        
        # Create left and right column frames
        left_column = ttk.Frame(columns_container)
        left_column.grid(row=0, column=0, sticky="n", padx=(0, 30))
        
        right_column = ttk.Frame(columns_container)
        right_column.grid(row=0, column=1, sticky="n", padx=(30, 0))
        
        # Configure grid to distribute space evenly
        columns_container.columnconfigure(0, weight=1)
        columns_container.columnconfigure(1, weight=1)
        
        # Dictionary to store variables and frames
        artifact_vars = {}
        checkbox_widgets = {}
        subtag_frames = {}
        
        # Function to create artifact checkboxes in a given column
        def create_artifact_checkboxes(column_frame, artifact_list):
            for artifact in artifact_list:
                # Create a frame for this artifact and its potential subtags
                artifact_frame = ttk.Frame(column_frame)
                artifact_frame.pack(anchor="w", fill="x", pady=2)
                
                # Create the main checkbox
                var = tk.IntVar(selection_window)
                
                # Set checkbox to checked if previously selected
                if artifact in previously_selected:
                    var.set(1)
                else:
                    var.set(0)
                
                cb = ttk.Checkbutton(
                    artifact_frame,
                    text=artifact,
                    variable=var,
                    onvalue=1,
                    offvalue=0
                )
                cb.pack(anchor="w", padx=10, pady=2)
                
                # Store references
                artifact_vars[artifact] = var
                checkbox_widgets[artifact] = cb
                
                # For Pictures and Videos, create a subtag frame
                if artifact in ["Pictures", "Videos"]:
                    # Create a container for subtags
                    subtag_container = ttk.Frame(artifact_frame)
                    subtag_frames[artifact] = subtag_container
                    
                    # Add the subtags
                    subtags = ["Child Pornography", "Child Erotica", "Age Difficult"]
                    for subtag in subtags:
                        subtag_name = f"{artifact} -- {subtag}"
                        subvar = tk.IntVar(selection_window)
                        
                        # Set subtag checkbox to checked if previously selected
                        if subtag_name in previously_selected:
                            subvar.set(1)
                        else:
                            subvar.set(0)
                        
                        sub_cb = ttk.Checkbutton(
                            subtag_container,
                            text=f"• {subtag}",
                            variable=subvar,
                            onvalue=1,
                            offvalue=0
                        )
                        sub_cb.pack(anchor="w", padx=30, pady=2)
                        
                        artifact_vars[subtag_name] = subvar
                    
                    # Show subtags if parent was previously selected
                    if artifact in previously_selected:
                        subtag_container.pack(anchor="w", fill="x")
                    else:
                        subtag_container.pack_forget()
                    
                    def make_toggle_function(artifact_name, container):
                        def toggle_func(*args):
                            if artifact_vars[artifact_name].get() == 1:
                                container.pack(anchor="w", fill="x")
                            else:
                                container.pack_forget()
                                # Uncheck all subtags when parent is unchecked
                                for key, var in artifact_vars.items():
                                    if key.startswith(f"{artifact_name} -- "):
                                        var.set(0)
                        return toggle_func
                    
                    toggle_func = make_toggle_function(artifact, subtag_container)
                    var.trace("w", toggle_func)
        
        # Create checkboxes in both columns
        create_artifact_checkboxes(left_column, left_artifacts)
        create_artifact_checkboxes(right_column, right_artifacts)
        
        # Button frame at the bottom
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=10)
        
        # Select All / Deselect All buttons
        def select_all():
            for key, var in artifact_vars.items():
                if not key.startswith("Pictures -- ") and not key.startswith("Videos -- "):
                    var.set(1)
            # Make sure to show subtag frames
            for artifact in ["Pictures", "Videos"]:
                if artifact in subtag_frames:
                    subtag_frames[artifact].pack(anchor="w", fill="x")
        
        def deselect_all():
            for key, var in artifact_vars.items():
                var.set(0)
            # Hide subtag frames
            for artifact in ["Pictures", "Videos"]:
                if artifact in subtag_frames:
                    subtag_frames[artifact].pack_forget()
        
        select_all_button = ttk.Button(button_frame, text="Select All", command=select_all)
        select_all_button.pack(side="left", padx=5)
        
        deselect_all_button = ttk.Button(button_frame, text="Deselect All", command=deselect_all)
        deselect_all_button.pack(side="left", padx=5)
        
        # Done button
        def on_done():
            nonlocal artifacts
            artifacts = []  # Clear the list first
            
            for artifact_name, var in artifact_vars.items():
                # Only add artifacts with checked boxes (value = 1)
                if var.get() == 1:
                    artifacts.append(artifact_name)
            
            selection_window.destroy()
        
        done_button = ttk.Button(button_frame, text="Done", command=on_done)
        done_button.pack(side="right", padx=5)
        
        # Add a function to handle mousewheel scrolling
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        # Bind mousewheel event to the canvas
        canvas.bind_all("<MouseWheel>", on_mousewheel)
        
        # Wait for window to close
        self.wait_window(selection_window)
        
        # Unbind mousewheel event when window is closed
        try:
            canvas.unbind_all("<MouseWheel>")
        except:
            pass
        
        return artifacts

    def generate_artifact_paragraphs(self, doc, selected_artifacts):
        # Map of artifact names to their descriptions and tag names
        ios_artifacts = {
            "iOS Device Information": {
                "description": "iOS Device Information contains information about the physical device, such as model information, the software version, timezone information, and the IMEI.",
                "tag": "iOS Device Information"
            },
            "Owner Information": {
                "description": "Owner Information contains information about the iOS device and the device owner. Information includes the device name, the phone number associated with the phone, and other details associated with iOS.",
                "tag": "Owner Information"
            },
            "Apple Accounts": {
                "description": "Apple Accounts contains information about the Apple ID accounts used on the macOS computer. The account details contained can help investigators recover and correlate account information across applications, and provide information on what accounts to review and get more information from.",
                "tag": "Apple Accounts"
            },
            "Google Accounts": {
                "description": "Google Accounts contains the Google accounts that are currently signed in on any Google application on the device.",
                "tag": "Google Accounts"
            },
            "SIM Card Activity": {
                "description": "SIM Card Activity contains information about the SIM cards that have been used in an iOS device.",
                "tag": "SIM Card Activity"
            },
            "Chrome Downloads": {
                "description": "Chrome Downloads contains information about the files that a user downloads from the Internet.",
                "tag": "Chrome Downloads"
            },
            "Chrome Keyword Search Terms": {
                "description": "Chrome Keyword Search Terms contains information about the keyword search terms that a user enters.",
                "tag": "Chrome Keyword Search Terms"
            },
            "Chrome Web History": {
                "description": "Chrome Web History contains a history of the websites that the user visits (includes unique visits only).",
                "tag": "Chrome Web History"
            },
            "Chrome Web Visits": {
                "description": "Chrome Web Visits contains a history of the websites that the user visits (includes all visits).",
                "tag": "Chrome Web Visits"
            },
            "Apple Maps Trips": {
                "description": "Apple Maps Trips contains trips generated by Apple Maps. Instances of this artifact can be suggested routes as well as trips that the user actually takes.",
                "tag": "Apple Maps Trips"
            },
            "Cached Locations": {
                "description": "Cached Locations stores a sample of locations that the iOS device has cached. Each instance contains the location, speed, and direction of travel at that particular point in time. The frequency that location samples are cached can vary depending on device usage, where some applications (such as Maps) can result in samples being cached very frequently.",
                "tag": "Cached Locations"
            },"Significant Locations": {
                "description": "Significant Locations contains information about places that are deemed to be significant in some way to the user. These locations can be manually added by the user (such as a home or work address) or are automatically added by Apple. This data is used to help make more personalized predictions.",
                "tag": "Significant Locations"
            },"Apple Notes": {
                "description": "Apple Notes contains information about the notes that a user has created on their Apple device.",
                "tag": "Apple Notes"
            },
            "Apple Notes - Voice": {
                "description": "Apple Notes Voice contains the recovered voice notes from an iOS device.",
                "tag": "Apple Notes - Voice"
            },
            "Live Photos": {
                "description": "Live Photos contains Live Photos that were retrieved using parsing. Support exists for all versions of iOS.",
                "tag": "Live Photos",
            },
        }
        
        android_artifacts = {
            "Android Device Information": {
                "description": "Android Device Information contains the phone identification values.",
                "tag": "Android Device Information"
            },
            "Google Accounts": {
                "description": "Google Accounts contains the Google accounts that are currently signed in on any Google application on the device.",
                "tag": "Google Accounts"
            },
            "Accounts Information": {
                "description": "Contains the login information and tokens for accounts on the Android device.",
                "tag": "Accounts Information"
            },
            "Android Downloads": {
                "description": "Android Downloads contains file download information from a recovered Android device.",
                "tag": "Android Downloads"
            },
            "Chrome Downloads": {
                "description": "Chrome Downloads contains information about the files that a user downloads from the internet.",
                "tag": "Chrome Downloads"
            },
            "Chrome Keyword Search Terms": {
                "description": "Chrome Keyword Search Terms contains information about the keyword search terms that a user enters.",
                "tag": "Chrome Keyword Search Terms"
            },
            "Chrome Web History": {
                "description": "Chrome Web History contains a history of the websites that the user visits (includes unique visits only).",
                "tag": "Chrome Web History"
            },
            "Chrome Web Visits": {
                "description": "Chrome Web Visits contains a history of the websites that the user visits (includes all visits).",
                "tag": "Chrome Web Visits"
            },
            "Samsung Notes": {
                "description": "Samsung Notes contains information about the notes that a user has created on their Samsung device.",
                "tag": "Samsung Notes"
            },
            "Google Maps Search History": {
                "description": "Google Maps Search History contains information about locations that the user may have searched for using Google Maps. This includes searches for businesses, addresses, and points of interest. These searches will synchronize between different devices on Google Maps when signed into the same account.",
                "tag": "Google Maps Search History"
            }
        }
        
        # Artifacts for both iOS and Android
        common_artifacts = {
            "Pictures": {
                "description": "Pictures contains pictures that were retrieved using either carving or parsing techniques. The supported picture formats are JPEG (.jpeg, .jpg, .jpe), PNG (.png), Bitmaps (.bmp), Graphics Interchange Format (.gif), Icons (.ico), and Tagged Image File Format (.tif, .tiff).",
                "tag": "Pictures"
            },
            "Pictures -- Child Pornography": {
                "description": "The image files contained in this tag appear to depict child pornography. Child pornography consists of any visual depiction, including photographs, film, videos, or pictures depicting sexually explicit conduct. \"Sexually explicit conduct\" means material depicting any person under the age of 18 years engaged in graphic sexual intercourse, including genital-genital, oral-genital, anal-genital, or oral-anal whether between persons of the same or opposite sex, or lascivious simulated sexual intercourse where the genitals, breast or pubic area of any person is exhibited. In addition, any material depicting a minor involved in bestiality; masturbation; sadistic or masochistic abuse; or lascivious exhibition of the genitals or pubic area.",
                "tag": "Pictures -- Child Pornography"
            },
            "Pictures -- Child Erotica": {
                "description": "The image files contained in this tag do not meet the statutory requirement to be considered child pornography. These image files depict juvenile subjects who are shown wearing sexually suggestive clothing, posing in sexually suggestive positions, or that appear to be possessed for a sexual purpose. Image files of this nature are often referred to as \"child erotica\" by forensic examiners and investigators.",
                "tag": "Pictures -- Child Erotica"
            },
            "Pictures -- Age Difficult": {
                "description": "The image files in this tag are pornographic in nature and depict younger looking subjects who may be juveniles under the age of 18, or may be young adults who are 18 years of age or older. Therefore, without positive identification of the subjects shown in the image files in this tag, I cannot make an accurate determination regarding the legal or illegal nature of the image files. Pornographic image files and video files of this nature are often referred to as \"age difficult\" pornography by forensic examiners and investigators.",
                "tag": "Pictures -- Age Difficult"
            },
            "Videos": {
                "description": "Videos contains videos that are recovered using parsing or carving. Supported formats for parsing include AVI, MP4, MOV, MPEG, DIVX, A3GP, ASF, WMV, DVR-MS, MKV, VOB, MOD, and WEBM. Supported carving formats include AVI, MP4, DIVX, A3GP, M4A, QT, and WEBM.",
                "tag": "Videos"
            },
            "Videos -- Child Pornography": {
                "description": "The video files contained in this tag appear to depict child pornography. Child pornography consists of any visual depiction, including photographs, film, videos, or pictures depicting sexually explicit conduct. \"Sexually explicit conduct\" means material depicting any person under the age of 18 years engaged in graphic sexual intercourse, including genital-genital, oral-genital, anal-genital, or oral-anal whether between persons of the same or opposite sex, or lascivious simulated sexual intercourse where the genitals, breast or pubic area of any person is exhibited. In addition, any material depicting a minor involved in bestiality; masturbation; sadistic or masochistic abuse; or lascivious exhibition of the genitals or pubic area.",
                "tag": "Videos -- Child Pornography"
            },
            "Videos -- Child Erotica": {
                "description": "The video files contained in this tag do not meet the statutory requirement to be considered child pornography. These video files depict juvenile subjects who are shown wearing sexually suggestive clothing, posing in sexually suggestive positions, or that appear to be possessed for a sexual purpose. Video files of this nature are often referred to as \"child erotica\" by forensic examiners and investigators.",
                "tag": "Videos -- Child Erotica"
            },
            "Videos -- Age Difficult": {
                "description": "The video files in this tag are pornographic in nature and depict younger looking subjects who may be juveniles under the age of 18, or may be young adults who are 18 years of age or older. Therefore, without positive identification of the subjects shown in the video files in this tag, I cannot make an accurate determination regarding the legal or illegal nature of the video files. Video files of this nature are often referred to as \"age difficult\" pornography by forensic examiners and investigators.",
                "tag": "Videos -- Age Difficult"
            }
        }
        
        # Combine the appropriate artifacts based on the selected ones
        artifacts_map = {**common_artifacts}
        if any(artifact.startswith("iOS") or artifact.startswith("Apple") or artifact.startswith("Safari") or artifact.startswith("iPhone") for artifact in selected_artifacts):
            artifacts_map.update(ios_artifacts)
        else:
            artifacts_map.update(android_artifacts)
        
        # Organize selected artifacts to ensure proper ordering
        # This organizes main categories followed by their subcategories
        organized_artifacts = []
        
        # First, add main categories that don't have subcategories
        for artifact in selected_artifacts:
            if " -- " not in artifact and artifact not in ["Pictures", "Videos"]:
                organized_artifacts.append(artifact)
        
        # Next, add Pictures and its subcategories if selected
        if "Pictures" in selected_artifacts:
            organized_artifacts.append("Pictures")
            for artifact in selected_artifacts:
                if artifact.startswith("Pictures -- "):
                    organized_artifacts.append(artifact)
        
        # Then, add Videos and its subcategories if selected
        if "Videos" in selected_artifacts:
            organized_artifacts.append("Videos")
            for artifact in selected_artifacts:
                if artifact.startswith("Videos -- "):
                    organized_artifacts.append(artifact)
        
        # Add any remaining artifacts that weren't processed
        for artifact in selected_artifacts:
            if artifact not in organized_artifacts:
                organized_artifacts.append(artifact)
        
        # Number counter for main artifacts only
        counter = 1
        
        # Add selected artifacts to the document
        for artifact in organized_artifacts:
            if artifact in artifacts_map:
                # Check if it's a subtag (for Pictures or Videos)
                is_subtag = " -- " in artifact
                
                # For main categories
                if not is_subtag:
                    # Add item number and artifact name with proper formatting
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(f"{counter}) ")
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    run = p.add_run(f"{artifact.upper()}")
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    # Add blank line
                    doc.add_paragraph()
                    
                    # For all main categories, add description first
                    # Add description first
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(artifacts_map[artifact]["description"])
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    # Add blank line after description
                    doc.add_paragraph()
                    
                    # Then add tag information with underlining
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(f"Tag: {artifacts_map[artifact]['tag']}")
                    run.font.underline = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    # Increment counter for main artifacts
                    counter += 1
                else:
                    # For subtags under Pictures and Videos, keep tag first then description
                    # Add tag information FIRST
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(f"Tag: {artifacts_map[artifact]['tag']}")
                    run.font.underline = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
                    # Add blank line after tag
                    doc.add_paragraph()
                    
                    # Then add description
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run(artifacts_map[artifact]["description"])
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                
                # Add blank line after the description
                doc.add_paragraph()
                
                # Add DEVICE ARTIFACTS indication
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run("(DEVICE ARTIFACTS)\n")
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                
                # Add blank line after each artifact
                doc.add_paragraph()
    
    def validate_fields(self):
        missing_fields = []

        examiner_agency = (self.examiner_agency_type.get() or "").strip()
        
        # Get examiner title value based on dropdown selection
        examiner_title = self.examiner_title_type.get()
        if not examiner_title.strip():
            missing_fields.append("Examiner Title")

        # Base fields that are always required
        base_fields = {
            "Device Owner": self.device_owner.get(),
            "Examiner Agency": examiner_agency,
            "Examiner Title": examiner_title,
            "Examiner Name": self.examiner_name.get(),
            "Case Number": self.case_number.get(),
            "Evidence Number": self.evidence_number.get(),
            "DFR Number": self.DFR_Num.get() if is_complete_dfr_number(self.DFR_Num.get()) else "",
        }

        # Check role-specific fields
        if self.role_type.get() == "Agency Assist":
            role_fields = {
                "Exam Request Date": self.request_date.get(),
                "Requesting Agency": self.request_agency.get(),
                "Requesting Officer": self.request_officer.get(),
                "Case Type": self.case_type.get(),
            }
            
            if not self.get_request_title():
                missing_fields.append("Requesting Officer Title")
                
            # Check time frame fields for Agency Assist
            if (hasattr(self, 'legal_authority') and self.legal_authority.get() == 'Search Warrant' and 
                hasattr(self, 'time_frame_limited_var') and self.time_frame_limited_var.get() == 1):
                if not self.time_frame_start_date.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.time_frame_end_date.get().strip():
                    missing_fields.append("Time Frame End Date")
        else:  # Case Agent
            role_fields = {
                "Case Type": self.offense_type.get(),
            }
            
            if self.legal_self.get() == 'Search Warrant' and not self.sw_service_date.get().strip():
                missing_fields.append("Search Warrant Service Date")
                
            # Check time frame fields for Case Agent
            if (hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant' and 
                hasattr(self, 'case_time_frame_limited_var') and self.case_time_frame_limited_var.get() == 1):
                if not self.case_time_frame_start_date.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.case_time_frame_end_date.get().strip():
                    missing_fields.append("Time Frame End Date")

        # Combine all fields
        fields = {**base_fields, **role_fields}


        for field_name, field_value in fields.items():
            if not field_value.strip():
                missing_fields.append(field_name)

        # Fixed: Use device_transfer_var instead of device_transfer
        if self.device_transfer_var.get() == 1:
            if not self.transfer_title.get().strip():
                missing_fields.append("Transfer Officer Title")
            if not self.transfer_officer.get().strip():
                missing_fields.append("Transfer Officer Name")
            if not self.transfer_agency.get().strip():
                missing_fields.append("Transfer Officer Agency")
            if not self.transfer_date.get().strip():
                missing_fields.append("Transfer Date")

        selected_software = self.get_selected_forensic_software()
        if not selected_software:
            missing_fields.append("Forensic Processing Software (at least one must be selected)")

        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            missing_fields.append("Extraction File (UFD or PDF)")
        if not hasattr(self, 'template_file') or not self.template_file:
            missing_fields.append("Template File")
        if not hasattr(self, 'extraction_type') or not self.extraction_type:
            missing_fields.append("Extraction Type (UFD or PDF file required)")

        return missing_fields
  
    def search_and_replace_content_controls_simple(self, doc, search_docs):
        """
        Enhanced version that handles multi-paragraph replacements for PY_TEXT
        with proper newline handling
        """
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        # Get the document XML as bytes
        doc_xml_str = etree.tostring(doc._element, encoding='unicode')
        
        # Perform replacements directly in the XML string
        for search_string, replacement_doc in search_docs.items():
            if search_string in doc_xml_str:
                # Special handling for PY_TEXT - multi-paragraph content
                if search_string == "PY_TEXT":
                    # Build the complete replacement XML for all paragraphs
                    replacement_xml_parts = []
                    
                    for para in replacement_doc.paragraphs:
                        # Split the paragraph text by newlines to create separate paragraphs
                        para_text = para.text
                        text_parts = para_text.split('\n')
                        
                        for text_part in text_parts:
                            # Skip empty parts that aren't intentional blank paragraphs
                            if not text_part and len(text_parts) > 1:
                                # Add empty paragraph for blank lines
                                replacement_xml_parts.append('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:pPr></w:pPr></w:p>')
                                continue
                            
                            # Create paragraph XML with proper formatting
                            para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            
                            # Add paragraph properties if needed (for alignment, spacing, etc.)
                            para_xml += '<w:pPr></w:pPr>'
                            
                            # Check if this entire paragraph should be bold/underline
                            # (for headers like "FORENSIC EXTRACTION")
                            is_header = False
                            for run in para.runs:
                                if run.font.bold and run.font.underline:
                                    is_header = True
                                    break
                            
                            if text_part:  # Only add run if there's text
                                para_xml += '<w:r>'
                                para_xml += '<w:rPr>'
                                
                                # Apply formatting based on the original run formatting
                                if is_header or (para.runs and para.runs[0].font.bold):
                                    para_xml += '<w:b/>'
                                if is_header or (para.runs and para.runs[0].font.underline):
                                    para_xml += '<w:u w:val="single"/>'
                                if para.runs and para.runs[0].font.italic:
                                    para_xml += '<w:i/>'
                                
                                # Font settings
                                font_name = 'Arial'
                                font_size = 22  # 11pt * 2 (half-points)
                                if para.runs and para.runs[0].font.name:
                                    font_name = para.runs[0].font.name
                                if para.runs and para.runs[0].font.size:
                                    font_size = para.runs[0].font.size.pt * 2 if hasattr(para.runs[0].font.size, 'pt') else 22
                                
                                para_xml += f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
                                para_xml += f'<w:sz w:val="{int(font_size)}"/>'
                                
                                para_xml += '</w:rPr>'
                                
                                # Add the text content (escape XML special characters)
                                text_content = text_part
                                text_content = text_content.replace('&', '&amp;')
                                text_content = text_content.replace('<', '&lt;')
                                text_content = text_content.replace('>', '&gt;')
                                text_content = text_content.replace('"', '&quot;')
                                text_content = text_content.replace("'", '&apos;')
                                
                                para_xml += f'<w:t xml:space="preserve">{text_content}</w:t>'
                                para_xml += '</w:r>'
                            
                            para_xml += '</w:p>'
                            replacement_xml_parts.append(para_xml)
                    
                    # Join all paragraphs
                    replacement_text = ''.join(replacement_xml_parts)
                    
                else:
                    # Handle single paragraph replacements (existing logic)
                    if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                        replacement_text = replacement_doc.paragraphs[0].text
                    else:
                        replacement_text = ""
                    
                    # Escape XML special characters in replacement text
                    replacement_text = replacement_text.replace('&', '&amp;')
                    replacement_text = replacement_text.replace('<', '&lt;')
                    replacement_text = replacement_text.replace('>', '&gt;')
                    replacement_text = replacement_text.replace('"', '&quot;')
                    replacement_text = replacement_text.replace("'", '&apos;')
                
                # Replace in the XML string
                doc_xml_str = doc_xml_str.replace(search_string, replacement_text)
                replaced_strings[search_string] = True
                
                if search_string == "PY_TEXT":
                    print(f"Replaced '{search_string}' with formatted paragraphs")
                else:
                    print(f"Replaced '{search_string}' with '{replacement_text[:50]}...'")
        
        # Parse the modified XML back into the document
        new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
        doc._element.clear()
        for child in new_element:
            doc._element.append(child)
        
        return replaced_strings

    # Add this method for handling split placeholders
    def search_and_replace_split_placeholders(self, doc, search_docs):
        """
        Handle placeholders that are split across XML elements
        """
        import re  # Move the import to the top
        
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        try:
            # Get the document XML as string
            doc_xml_str = etree.tostring(doc._element, encoding='unicode')
            original_xml_str = doc_xml_str
            
            print("\nHandling split placeholders...")
            
            # Look for placeholders that might be split by XML tags
            target_placeholders = ['PY_DFR', 'PY_OWNER', 'PY_EXAMINER']
            
            for search_string in target_placeholders:
                if replaced_strings[search_string]:
                    continue
                    
                replacement_doc = search_docs[search_string]
                if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                    replacement_text = replacement_doc.paragraphs[0].text
                else:
                    replacement_text = ""
                
                # Escape XML special characters
                replacement_text = replacement_text.replace('&', '&amp;')
                replacement_text = replacement_text.replace('<', '&lt;')
                replacement_text = replacement_text.replace('>', '&gt;')
                replacement_text = replacement_text.replace('"', '&quot;')
                replacement_text = replacement_text.replace("'", '&apos;')
                
                # Create pattern allowing XML tags between each character
                pattern_parts = []
                for char in search_string:
                    pattern_parts.append(re.escape(char))
                
                # Join with pattern that allows XML tags between characters
                pattern = r'(?:<[^>]*>)*'.join(pattern_parts)
                
                # Find and replace the split placeholder
                matches = list(re.finditer(pattern, doc_xml_str))
                
                if matches:
                    # Replace from last match to first to avoid position shifts
                    for match in reversed(matches):
                        start_pos = match.start()
                        end_pos = match.end()
                        doc_xml_str = doc_xml_str[:start_pos] + replacement_text + doc_xml_str[end_pos:]
                    
                    replaced_strings[search_string] = True
                    print(f"  ✓ Replaced split placeholder '{search_string}' ({len(matches)} occurrences)")
                else:
                    print(f"  ✗ Could not find split placeholder '{search_string}'")
            
            # Parse back the document if changes were made
            if doc_xml_str != original_xml_str:
                new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
                doc._element.clear()
                for child in new_element:
                    doc._element.append(child)
            
        except Exception as e:
            print(f"Split placeholder replacement failed: {e}")
        
        return replaced_strings

    # Add this method for formatting time frame dates
    def format_time_frame_date(self, date_string):
        """
        Format time frame dates using the same format as other dates in the application
        """
        if not date_string or not date_string.strip():
            return ""
        
        try:
            # Use the existing parse_request_date method to format consistently
            return self.parse_request_date(date_string.strip())
        except ValueError as e:
            # If parsing fails, return the original string
            print(f"Warning: Could not format time frame date '{date_string}': {e}")
            return date_string
  
    def preview_placeholders(self):
        self.generate_report(preview_only=True)

    def generate_report(self, preview_only=False):
        try:
            missing_fields = self.validate_fields()
            if missing_fields and not preview_only:
                messagebox.showerror(
                    "Missing Fields",
                    f"The following fields are missing:\n\n{', '.join(missing_fields)}\n\nPlease complete them before proceeding."
                )
                return

            examiner_agency = self.get_examiner_agency()
            examiner_agency_formatted, examiner_agency_abbr = self.format_agency(examiner_agency, return_abbreviation=True)

            examiner_title_value = self.examiner_title_type.get()

            # Handle officer name / requesting agency based on role.
            # Case Agent: the examiner IS the requesting officer, so PY_REQOFF / PY_REQAGENCY
            # and narrative placeholders still resolve through the existing replace path.
            if self.role_type.get() == "Agency Assist":
                officer_name = self.request_officer.get().strip()
                officer_last_name = officer_name.split()[-1] if officer_name and ' ' in officer_name else officer_name
                request_title_value = self.format_title(self.get_request_title())
                request_agency_formatted, request_agency_abbr = self.format_agency(self.request_agency.get(), return_abbreviation=True)
            else:
                officer_name = self.examiner_name.get().strip()
                officer_last_name = officer_name.split()[-1] if officer_name and ' ' in officer_name else officer_name
                request_title_value = self.format_title(examiner_title_value)
                request_agency_formatted, request_agency_abbr = examiner_agency_formatted, examiner_agency_abbr

            if self.role_type.get() == "Agency Assist":
                request_date = self.parse_request_date(self.request_date.get())
            elif self.role_type.get() == "Case Agent" and self.legal_self.get() == 'Search Warrant':
                request_date = self.parse_request_date(self.sw_service_date.get())
            else:
                request_date = ""

            # Get the case type based on role
            if self.role_type.get() == "Agency Assist":
                case_type = self.case_type.get()
            else:  # Case Agent
                case_type = self.offense_type.get()

            # Get examiner title value based on dropdown selection
            examiner_title = self.examiner_title_type.get()

            data = {
                'Request_Date': request_date,
                'Request_Agency': request_agency_formatted,
                'Request_Agency_Abbr': request_agency_abbr,
                'Request_Title': request_title_value,
                'Request_Officer': officer_name.title() if officer_name else "",
                'Request_Officer_LastName': officer_last_name.title() if officer_last_name else '',
                'Request_Case': case_type,
                'Device_Owner': self.device_owner.get().title(),
                'Examiner_Agency': examiner_agency_formatted,
                'Examiner_Agency_Abbr': examiner_agency_abbr,
                'Examiner_Title': self.format_title(examiner_title),
                'Examiner_Name': self.examiner_name.get().title(),
                'Case_Number': self.case_number.get().upper(),
                'evidence_ID': self.evidence_number.get().upper(),
                'DFR_Num': self.DFR_Num.get().upper(),
            }

            data.update({
                'device_passcode': self.device_passcode.get().strip() if hasattr(self, 'device_passcode') else '',
                'device_color': self.device_color.get().strip() if hasattr(self, 'device_color') else '',
                'device_capacity': self.device_capacity.get().strip() if hasattr(self, 'device_capacity') else '',
                'device_iccid': self.device_iccid.get().strip() if hasattr(self, 'device_iccid') else '',
                'device_carrier': self.device_carrier.get().strip().title() if hasattr(self, 'device_carrier') else '',
            })

            if self.role_type.get() == "Case Agent" and self.legal_self.get() == 'Search Warrant':
                data['SW_Date'] = self.parse_request_date(self.sw_service_date.get())
            else:
                data['SW_Date'] = request_date or ""

            # Add transfer fields if applicable
            if self.device_transfer_var.get() == 1:
                transfer_agency_formatted, transfer_agency_abbr = self.format_agency(self.transfer_agency.get(), return_abbreviation=True)
                data.update({
                    'Transfer_Date': self.parse_request_date(self.transfer_date.get()),
                    'Transfer_Title': self.format_title(self.transfer_title.get()),
                    'Transfer_Officer': self.transfer_officer.get().title(),
                    'Transfer_Agency': transfer_agency_formatted,
                    'Transfer_Agency_Abbr': transfer_agency_abbr,
                })

            # Add time frame fields based on role
            if self.role_type.get() == "Agency Assist":
                if (hasattr(self, 'legal_authority') and self.legal_authority.get() == 'Search Warrant' and
                    hasattr(self, 'time_frame_limited_var') and self.time_frame_limited_var.get() == 1):
                    if hasattr(self, 'time_frame_start_date') and self.time_frame_start_date.get().strip():
                        data['PY_LIMITSTART'] = self.format_time_frame_date(self.time_frame_start_date.get())
                    if hasattr(self, 'time_frame_end_date') and self.time_frame_end_date.get().strip():
                        data['PY_LIMITEND'] = self.format_time_frame_date(self.time_frame_end_date.get())
            else:  # Case Agent
                if (hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant' and
                    hasattr(self, 'case_time_frame_limited_var') and self.case_time_frame_limited_var.get() == 1):
                    if hasattr(self, 'case_time_frame_start_date') and self.case_time_frame_start_date.get().strip():
                        data['PY_LIMITSTART'] = self.format_time_frame_date(self.case_time_frame_start_date.get())
                    if hasattr(self, 'case_time_frame_end_date') and self.case_time_frame_end_date.get().strip():
                        data['PY_LIMITEND'] = self.format_time_frame_date(self.case_time_frame_end_date.get())

            try:
                extraction_data = self.parse_extraction_file()
            except Exception:
                extraction_data = {}
            identity_missing = require_device_identity(extraction_data)
            if identity_missing and not preview_only:
                messagebox.showerror("Missing Fields", "\n".join(identity_missing))
                return
            data = prefer_gui_over_parsed(data, extraction_data)
            if preview_only:
                suggested = apply_suggested_filename(self, "Mobile", data.get("device_model", ""))
                show_placeholder_preview(
                    self,
                    mobile_preview_rows(
                        data,
                        officer_text=self.format_case_officer(data),
                        image_date=self.format_extraction_date(data.get("formatted_date", "")),
                    ),
                    suggested,
                )
                return

            # Load the template document
            doc = Document(self.template_file)

            # METHOD 1: Try direct paragraph replacement first
            py_text_found = False
            for para_index, para in enumerate(doc.paragraphs):
                if 'PY_TEXT' in para.text:
                    py_text_found = True
                    print(f"Found PY_TEXT in paragraph {para_index}")

                    # Clear the paragraph
                    para.clear()

                    # Generate all the content paragraphs
                    temp_doc = Document()
                    self.generate_paragraphs(data, temp_doc)

                    # Get the parent element
                    parent = para._element.getparent()
                    insert_index = parent.index(para._element)

                    # Insert all generated paragraphs at this position
                    for new_para in reversed(temp_doc.paragraphs):
                        # Create a copy of the paragraph in the main document
                        copied_para = doc.add_paragraph()

                        # Copy all runs with formatting
                        for run in new_para.runs:
                            new_run = copied_para.add_run(run.text)
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                            new_run.font.underline = run.font.underline
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size

                        # Move to correct position
                        parent.insert(insert_index, copied_para._element)

                    # Remove the original PY_TEXT paragraph
                    parent.remove(para._element)
                    break

            # Handle other replacements
            search_docs = {
                "PY_DFR": Document(),
                "PY_CASENUMBER": Document(),
                "PY_EVIDENCE": Document(),
                "PY_REQDATE": Document(),
                "PY_OWNER": Document(),
                "PY_REQAGENCY": Document(),
                "PY_REQOFF": Document(),
                "PY_LIMITSTART": Document(),
                "PY_LIMITEND": Document(),
                "PY_EXAMINER": Document(),
                "PY_IMAGEDATE": Document(),

                "PY_MAN": Document(),
                "PY_MOD": Document(),
                "PY_COLOR": Document(),
                "PY_PHONE": Document(),
                "PY_SERIAL": Document(),
                "PY_IMEI": Document(),
                "PY_CAPACITY": Document(),
                "PY_DEVNAME": Document(),
                "PY_ACCOUNT": Document(),
                "PY_ICCID": Document(),
                "PY_PASSCODE": Document(),

                "PY_CARRIER": Document(),
                "PY_OS": Document(),
                "PY_CBVER": Document(),
                "PY_GKVER": Document(),
            }

            # Generate replacement content for other fields
            self.generate_replacements(data, search_docs)

            # METHOD 2: If PY_TEXT wasn't found in paragraphs, it might be in a content control
            if not py_text_found:
                print("PY_TEXT not found in paragraphs, trying XML method...")
                # Create a document with all the generated paragraphs for PY_TEXT
                new_doc = Document()
                self.generate_paragraphs(data, new_doc)
                search_docs["PY_TEXT"] = new_doc

            # Replace remaining placeholders using XML method
            try:
                replaced_strings = self.search_and_replace_content_controls_simple(doc, search_docs)

                # Check for any remaining missing placeholders
                missing_strings = [s for s, replaced in replaced_strings.items() if not replaced]

                if missing_strings:
                    print(f"Trying split placeholder method for: {missing_strings}")
                    split_replaced = self.search_and_replace_split_placeholders(doc, search_docs)

                    # Update the replaced_strings with any newly found ones
                    for search_string, was_replaced in split_replaced.items():
                        if was_replaced:
                            replaced_strings[search_string] = True

            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"Replacement method failed: {e}")
                print(f"Error details: {error_details}")

            apply_suggested_filename(self, "Mobile", data.get("device_model", ""))
            self.save_document(doc)

            self.selected_artifacts = []

        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            messagebox.showerror("Error", f"Failed to generate report:\n\n{str(e)}\n\nDetails:\n{error_details[:500]}")

    def format_extraction_date(self, formatted_date_str):
        if not formatted_date_str:
            return ""
        
        try:
            # Split by " at " to get just the date part
            date_part = formatted_date_str.split(" at ")[0]
            return date_part
        except Exception:
            # If parsing fails, return the original string
            return formatted_date_str

    def generate_replacements(self, data, search_docs):
        def add_text_to_doc(doc, text):
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p
            
        # Get selected forensic software and format text
        selected_software = self.get_selected_forensic_software()
        if len(selected_software) == 1:
            forensic_software_text = selected_software[0]
        elif len(selected_software) == 2:
            forensic_software_text = " and ".join(selected_software)
        elif len(selected_software) > 2:
            forensic_software_text = ", ".join(selected_software[:-1]) + ", and " + selected_software[-1]
        else:
            forensic_software_text = ""
        
        # Get time frame dates based on role
        time_frame_start = ""
        time_frame_end = ""
        
        if self.role_type.get() == "Agency Assist":
            if (hasattr(self, 'legal_authority') and self.legal_authority.get() == 'Search Warrant' and 
                hasattr(self, 'time_frame_limited_var') and self.time_frame_limited_var.get() == 1):
                if hasattr(self, 'time_frame_start_date') and self.time_frame_start_date.get().strip():
                    time_frame_start = self.format_time_frame_date(self.time_frame_start_date.get())
                if hasattr(self, 'time_frame_end_date') and self.time_frame_end_date.get().strip():
                    time_frame_end = self.format_time_frame_date(self.time_frame_end_date.get())
        else:  # Case Agent
            if (hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant' and 
                hasattr(self, 'case_time_frame_limited_var') and self.case_time_frame_limited_var.get() == 1):
                if hasattr(self, 'case_time_frame_start_date') and self.case_time_frame_start_date.get().strip():
                    time_frame_start = self.format_time_frame_date(self.case_time_frame_start_date.get())
                if hasattr(self, 'case_time_frame_end_date') and self.case_time_frame_end_date.get().strip():
                    time_frame_end = self.format_time_frame_date(self.case_time_frame_end_date.get())
        
        # Map search strings to data values
        replacement_map = {
            "PY_DFR": data.get('DFR_Num', ''),
            "PY_CASENUMBER": data.get('Case_Number', '').strip(),
            "PY_EVIDENCE": data.get('evidence_ID', '').strip(),
            "PY_REQDATE": data.get('Request_Date', ''),
            "PY_OWNER": data.get('Device_Owner', ''),
            "PY_REQAGENCY": data.get('Request_Agency', ''),
            "PY_REQOFF": self.format_case_officer(data),
            "PY_LIMITSTART": time_frame_start,
            "PY_LIMITEND": time_frame_end,
            "PY_EXAMINER": f"{data.get('Examiner_Title', '')} {data.get('Examiner_Name', '')}".strip(),
            "PY_IMAGEDATE": self.format_extraction_date(data.get('formatted_date', '')),
            
            "PY_MAN": data.get('device_manufacturer', ''),
            "PY_MOD": data.get('device_model', ''),
            "PY_COLOR": data.get('device_color', ''),
            "PY_PHONE": data.get('Phone_Number', ''),
            "PY_SERIAL": data.get('Serial_Number', ''),
            "PY_IMEI": data.get('DEV_IMEI', ''),
            "PY_CAPACITY": data.get('device_capacity', ''),
            "PY_DEVNAME": data.get('Device_Name', ''),
            "PY_ACCOUNT": data.get('Device_Account', ''),
            "PY_ICCID": data.get('device_iccid', ''),
            "PY_PASSCODE": data.get('device_passcode', ''),
            
            "PY_CARRIER": data.get('device_carrier', ''),
            "PY_OS": data.get('Device_OS', ''),
            "PY_CBVER": data.get('cellebrite_version', ''),
            "PY_GKVER": data.get('GrayKey_OS', ''),
        }
        
        # Debug output
        print("\nGenerating replacement content:")
        
        # Create content for each search string
        for search_string, doc in search_docs.items():
            if search_string == "PY_TEXT":
                # Handled by generate_paragraphs
                para_count = len(doc.paragraphs)
                print(f"  {search_string}: {para_count} paragraphs already generated")
                continue
            
            value = replacement_map.get(search_string, '')
            
            # If no value exists, add empty string (space)
            if not value:
                value = ''
                print(f"  {search_string}: (empty)")
            else:
                print(f"  {search_string}: '{value[:50]}...'")
            
            add_text_to_doc(doc, value)

    def format_case_officer(self, data):
        parts = [] 
        # For Case Agent role, use examiner information
        if self.role_type.get() == "Case Agent":
            if data.get('Examiner_Title'):
                parts.append(data['Examiner_Title'])
            if data.get('Examiner_Name'):
                parts.append(data['Examiner_Name'])
        else:  # Agency Assist role           
            if data.get('Request_Title'):
                abbreviated_title = self.get_title_abbreviation(data['Request_Title'])
                parts.append(abbreviated_title)
            
            if data.get('Request_Officer'):
                parts.append(data['Request_Officer'])

        return ' '.join(parts) if parts else ''

    def format_request_agency(self, data):
        # For Case Agent role, use examiner agency
        if self.role_type.get() == "Case Agent":
            return data.get('Examiner_Agency_Abbr', '')
        else:  # Agency Assist role
            return data.get('Request_Agency_Abbr', '')

    def get_title_abbreviation(self, full_title):
        title_abbreviations = {
            "officer": "Off.",
            "deputy": "Dep.",
            "special agent": "SA",
            "supervisory special agent": "SSA",
            "special assistant attorney general": "SAAG",
            "saag": "SAAG",
            "intel analyst": "IA",
            "detective": "Det.",
            "investigator": "Inv.",
            "trooper": "Trooper",
            "computer forensic examiner": "CFE",
        }
        
        # Convert to lowercase for comparison
        title_lower = full_title.lower().strip()
        
        # Return the abbreviation if found, otherwise return the original title
        return title_abbreviations.get(title_lower, full_title)

    def get_examiner_agency(self):
        agency = (self.examiner_agency_type.get() or "").strip()
        if agency == "South Dakota DCI":
            return "South Dakota Division of Criminal Investigation"
        return agency

    def get_request_title(self):
        return (self.request_title_type.get() or "").strip()

    def validate_fields(self):
        missing_fields = []

        examiner_agency = (self.examiner_agency_type.get() or "").strip()
        
        # Get examiner title value based on dropdown selection
        examiner_title = self.examiner_title_type.get()
        if not examiner_title.strip():
            missing_fields.append("Examiner Title")

        # Base fields that are always required (REMOVED Case Number and Evidence Number)
        base_fields = {
            "Device Owner": self.device_owner.get(),
            "Examiner Agency": examiner_agency,
            "Examiner Title": examiner_title,
            "Examiner Name": self.examiner_name.get(),
            "DFR Number": self.DFR_Num.get() if is_complete_dfr_number(self.DFR_Num.get()) else "",
        }

        # Check role-specific fields
        if self.role_type.get() == "Agency Assist":
            role_fields = {
                "Exam Request Date": self.request_date.get(),
                "Requesting Agency": self.request_agency.get(),
                "Requesting Officer": self.request_officer.get(),
                "Case Type": self.case_type.get(),
            }
            
            if not self.get_request_title():
                missing_fields.append("Requesting Officer Title")
                
            # Check time frame fields for Agency Assist
            if (hasattr(self, 'legal_authority') and self.legal_authority.get() == 'Search Warrant' and 
                hasattr(self, 'time_frame_limited_var') and self.time_frame_limited_var.get() == 1):
                if not self.time_frame_start_date.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.time_frame_end_date.get().strip():
                    missing_fields.append("Time Frame End Date")
        else:  # Case Agent
            role_fields = {
                "Case Type": self.offense_type.get(),
            }
            
            if self.legal_self.get() == 'Search Warrant' and not self.sw_service_date.get().strip():
                missing_fields.append("Search Warrant Service Date")
                
            # Check time frame fields for Case Agent
            if (hasattr(self, 'legal_self') and self.legal_self.get() == 'Search Warrant' and 
                hasattr(self, 'case_time_frame_limited_var') and self.case_time_frame_limited_var.get() == 1):
                if not self.case_time_frame_start_date.get().strip():
                    missing_fields.append("Time Frame Start Date")
                if not self.case_time_frame_end_date.get().strip():
                    missing_fields.append("Time Frame End Date")

        # Combine all fields
        fields = {**base_fields, **role_fields}


        for field_name, field_value in fields.items():
            if not field_value.strip():
                missing_fields.append(field_name)

        # Fixed: Use device_transfer_var instead of device_transfer
        if self.device_transfer_var.get() == 1:
            if not self.transfer_title.get().strip():
                missing_fields.append("Transfer Officer Title")
            if not self.transfer_officer.get().strip():
                missing_fields.append("Transfer Officer Name")
            if not self.transfer_agency.get().strip():
                missing_fields.append("Transfer Officer Agency")
            if not self.transfer_date.get().strip():
                missing_fields.append("Transfer Date")

        selected_software = self.get_selected_forensic_software()
        if not selected_software:
            missing_fields.append("Forensic Processing Software (at least one must be selected)")

        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            missing_fields.append("Extraction File (UFD or PDF)")
        if not hasattr(self, 'template_file') or not self.template_file:
            missing_fields.append("Template File")
        if not hasattr(self, 'extraction_type') or not self.extraction_type:
            missing_fields.append("Extraction Type (UFD or PDF file required)")

        return missing_fields

    def generate_paragraphs(self, data, new_doc):
        def add_paragraph_with_style(doc, text):
            text = fill_paragraph(text, data)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p

        # Add Paragraph One - based on role type
        if self.role_type.get() == "Case Agent":
            # Case Agent always uses one_a
            add_paragraph_with_style(new_doc, self.paragraphs['one_a'])
        else:
            # Agency Assist uses one_b or one_c based on transfer
            # Fixed: Use device_transfer_var instead of device_transfer
            if self.device_transfer_var.get() == 0:
                add_paragraph_with_style(new_doc, self.paragraphs['one_b'])
            else:
                data.update({
                    'Transfer_Date': self.parse_request_date(self.transfer_date.get()),
                    'Transfer_Title': self.format_title(self.transfer_title.get()),
                    'Transfer_Officer': self.transfer_officer.get().title(),
                    'Transfer_Agency': self.format_agency(self.transfer_agency.get())
                })
                add_paragraph_with_style(new_doc, self.paragraphs['one_c'])

        # Add Paragraph Two - Airplane Mode
        if self.airplane_mode.get() == 'Yes':
            add_paragraph_with_style(new_doc, self.paragraphs['two_a'])
        else:
            add_paragraph_with_style(new_doc, self.paragraphs['two_b'])

        # Add Authority Paragraph (only if not empty)
        authority_paragraph = self.get_authority_paragraph()
        if authority_paragraph:  # Only add if not empty
            add_paragraph_with_style(new_doc, authority_paragraph)

        # Add Paragraph Three
        if self.role_type.get() == "Case Agent":
            add_paragraph_with_style(new_doc, self.paragraphs['three_b'])
        else:  # Agency Assist
            add_paragraph_with_style(new_doc, self.paragraphs['three_a'])

        # Add Paragraph Four
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['four'])

        # Add Paragraph Five - based on extraction type
        if self.extraction_type == "Cellebrite":
            if hasattr(self, 'multiple_extractions') and self.multiple_extractions:
                add_paragraph_with_style(new_doc, self.paragraphs['five_cellebrite_multiple'])
            else:
                add_paragraph_with_style(new_doc, self.paragraphs['five_cellebrite'])
        else: 
            add_paragraph_with_style(new_doc, self.paragraphs['five_graykey'])

        # Add Paragraph Six
        add_paragraph_with_style(new_doc, self.paragraphs['six'])

        # Add Paragraph Seven - Forensic Processing Header
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['seven'])

        # Add Paragraph Eight - Updated logic for multiple software combinations
        selected_software = self.get_selected_forensic_software()
        
        if len(selected_software) == 1:
            if "Axiom" in selected_software:
                add_paragraph_with_style(new_doc, self.paragraphs['eight_axiom'])
            elif "Cellebrite" in selected_software:
                add_paragraph_with_style(new_doc, self.paragraphs['eight_cellebrite'])
            elif "Griffeye" in selected_software:
                # You may want to add a specific paragraph for Griffeye-only processing
                add_paragraph_with_style(new_doc, self.paragraphs.get('eight_griffeye', self.paragraphs['griffeye_para']))
        elif len(selected_software) == 2:
            if "Axiom" in selected_software and "Cellebrite" in selected_software:
                add_paragraph_with_style(new_doc, self.paragraphs['eight_both'])
            elif "Griffeye" in selected_software:
                # Handle combinations with Griffeye
                other_software = [s for s in selected_software if s != "Griffeye"][0]
                if other_software == "Axiom":
                    add_paragraph_with_style(new_doc, self.paragraphs['eight_axiom'])
                elif other_software == "Cellebrite":
                    add_paragraph_with_style(new_doc, self.paragraphs['eight_cellebrite'])
        elif len(selected_software) >= 3:
            # All three software packages
            add_paragraph_with_style(new_doc, self.paragraphs['eight_both'])

        # Add Griffeye paragraph if Griffeye is selected
        if self.griffeye_var.get() == 1:
            add_paragraph_with_style(new_doc, self.paragraphs['griffeye_para'])

        # Check if "No Evidence Found" is checked
        if self.no_evidence_var.get() == 1:
            # Add Paragraph Nine - Findings of Exam Header
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['nine'])
            
            # Add No Evidence paragraph instead of paragraphs 10 and on
            add_paragraph_with_style(new_doc, self.paragraphs['Paragraph_NoEv'])
        else:
            # Normal flow - Add Paragraph Nine - Findings of Exam Header
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['nine'])
            
            # Add Paragraph Ten
            add_paragraph_with_style(new_doc, self.paragraphs['ten'])

            # Add Paragraph Eleventeen (Yes, eleventeen)
            add_paragraph_with_style(new_doc, self.paragraphs['eleventeen'])
            
            # Add Paragraph Twelveteen (They're totally words)
            add_paragraph_with_style(new_doc, self.paragraphs['twelveteen'])

            # Add Paragraph Thirteen
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['thirteen'])
            
            # Add Paragraph Fourteen
            add_paragraph_with_style(new_doc, self.paragraphs['fourteen'])

            # Add Paragraph Fifteen
            add_paragraph_with_style(new_doc, self.paragraphs['fifteen'])

            # Add Paragraph Sixteen
            add_paragraph_with_style(new_doc, self.paragraphs['sixteen'])

            # Add Paragraph Seventeen - Artifacts Header
            self.add_bold_underline_paragraph(new_doc, self.paragraphs['seventeen'])

            # BEGIN ARTIFACTS
            if hasattr(self, 'selected_artifacts') and self.selected_artifacts:
                # Use the previously selected artifacts
                self.generate_artifact_paragraphs(new_doc, self.selected_artifacts)
            elif self.axiom_var.get() == 1:
                # If no artifacts were selected but Axiom is used, add a message
                p = new_doc.add_paragraph()
                p.style = new_doc.styles['Normal']
                run = p.add_run("No artifacts were selected. Use the 'Select Artifacts' button to include artifacts in the report.")
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run.font.italic = True

    def save_document(self, doc):
        output_filename = self.output_filename.get().strip()
        
        if not output_filename:
            model = ""
            if hasattr(self, "extraction_file") and self.extraction_file:
                try:
                    model = self.parse_extraction_file().get("device_model", "")
                except Exception:
                    model = ""
            output_filename = suggested_report_filename(
                self.DFR_Num.get().strip(),
                "Mobile",
                self.device_owner.get().strip(),
                model,
            )
        
        save_location = self.save_location.get().strip()
        if not save_location or not os.path.exists(save_location):
            save_location = default_export_dir()
            self.save_location.delete(0, tk.END)
            self.save_location.insert(0, save_location)
        remember_folder("export", save_location)
        output_path = unique_output_path(save_location, output_filename)
        try:
            doc.save(output_path)
            remember_titles_from_form(self)
            remember_agencies_from_form(self)
            messagebox.showinfo("Success", f"Report generated and saved as:\n{output_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save report: {str(e)}")


    def get_authority_paragraph(self):
        authority_map = {
            'Search Warrant': 'auth_sw',
            'Consent': 'auth_c',
            'Implied Consent': 'auth_i',
            'Parole': 'auth_p'
        }
        
        # Get the appropriate authority field based on role
        if self.role_type.get() == "Case Agent":
            selected_authority = self.legal_self.get()
            
            # For Case Agent with Search Warrant, don't include the auth_sw paragraph
            # because there's no requesting officer information to reference
            if selected_authority == 'Search Warrant':
                return ""  # Return empty string to skip the authority paragraph
        else:
            selected_authority = self.legal_authority.get()
        
        if selected_authority not in authority_map:
            raise ValueError(f"Invalid legal authority selected: {selected_authority}")
        return self.paragraphs[authority_map[selected_authority]]

    def format_agency(self, agency, return_abbreviation=False):
        if not agency:
            return "" if not return_abbreviation else ("", "")
                
        # Dictionary mapping normalized lowercase inputs to (full_name, abbreviation)
        agencies = {
            "federal bureau of investigation": ("Federal Bureau of Investigation", "FBI"),
            "fbi": ("Federal Bureau of Investigation", "FBI"),
            "drug enforcement administration": ("Drug Enforcement Administration", "DEA"),
            "dea": ("Drug Enforcement Administration", "DEA"),
            "bureau of alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "atf": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "department of homeland security": ("Department of Homeland Security", "DHS"),
            "dhs": ("Department of Homeland Security", "DHS"),
            "united states marshals service": ("United States Marshals Service", "USMS"),
            "usms": ("United States Marshals Service", "USMS"),
            "south dakota division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sddci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sd dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "bureau of indian affairs": ("Bureau of Indian Affairs", "BIA"),
            "bia": ("Bureau of Indian Affairs", "BIA"),
            "south dakota highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "sdhp": ("South Dakota Highway Patrol", "SDHP"),
            "national security agency": ("National Security Agency", "NSA"),
            "nsa": ("National Security Agency", "NSA"),
            "federal emergency management agency": ("Federal Emergency Management Agency", "FEMA"),
            "fema": ("Federal Emergency Management Agency", "FEMA"),
            "internal revenue service": ("Internal Revenue Service", "IRS"),
            "irs": ("Internal Revenue Service", "IRS"),
            "department of defense": ("Department of Defense", "DOD"),
            "dod": ("Department of Defense", "DOD"),
            "department of justice": ("Department of Justice", "DOJ"),
            "doj": ("Department of Justice", "DOJ")
        }
        
        agency_lower = agency.lower().strip()
        
        # Check for direct match in the agencies dictionary
        if agency_lower in agencies:
            full_name, abbr = agencies[agency_lower]
            if return_abbreviation:
                return full_name, abbr
            else:
                return full_name
        
        # Preserve original for PD/SO handling
        original_agency = agency
        original_lower = original_agency.lower().strip()
        
        # Handle PD and SO expansion and abbreviation creation
        if original_lower.endswith(" pd"):
            # Extract the location part (everything before " pd")
            location = original_agency[:-3].strip()
            agency = f"{location} Police Department"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + PD
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "PD"
        elif original_lower.endswith(" so"):
            # Extract the location part (everything before " so")
            location = original_agency[:-3].strip()
            agency = f"{location} Sheriff's Office"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + SO
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "SO"
        else:
            # Default to None if not PD/SO
            abbreviation = None
        
        # Format the agency name with title case
        formatted_name = title_agency_words(agency)
        
        # For short agencies (4 or fewer chars), use uppercase
        if len(agency) <= 4:
            formatted_name = agency.upper()
        
        # If abbreviation wasn't set (non-PD/SO), calculate it
        if abbreviation is None:
            # Check for substring matches in agencies dict (fallback for partial matches)
            found_match = False
            for full_lower, (full_name, abbr) in agencies.items():
                if full_lower in agency_lower:
                    formatted_name = full_name
                    abbreviation = abbr
                    found_match = True
                    break
            
            # If still no match, generate abbreviation dynamically
            if not found_match:
                if len(agency) <= 4:
                    abbreviation = agency.upper()
                else:
                    words = agency.split()
                    skip_words = ['of', 'the', 'and', 'in', 'on', 'at', 'by', 'for', 'with', 'a', 'an']
                    abbr_chars = [word[0].upper() for word in words if word.lower() not in skip_words and word]
                    abbreviation = ''.join(abbr_chars)
        
        if return_abbreviation:
            return formatted_name, abbreviation
        else:
            return formatted_name

    def format_title(self, title):
        if len(title) <= 2:
            return title.upper()
        return ' '.join(word.capitalize() for word in title.split())

    def parse_request_date(self, date_string):
        return parse_mdy_date(date_string)

    def back_to_start(self):
        close_and_return(self)
        
    def on_closing(self):
        close_and_return(self)


# Ω Digital Forensics Report Writer Ω (ver. 1.0.1) © 2026 #