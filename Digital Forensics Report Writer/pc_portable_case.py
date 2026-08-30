import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt
import os
import sys
import time
import ctypes
import re
import PyPDF2
from tkinter.scrolledtext import ScrolledText
from tkinterdnd2 import DND_FILES, TkinterDnD
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from settings_manager import SettingsManager
from ui_theme import apply_theme, add_header_bar, add_action_bar, pack_right_actions, size_window, build_extracted_info_pane, parse_mdy_date, add_date_entry, COLORS, FormTabs, close_and_return
from app_menu import attach_app_menu
from paragraphs_manager import fill_paragraph, load_paragraphs
from report_common import (
    ask_open_file,
    ask_directory,
    default_export_dir,
    remember_folder,
    suggested_report_filename,
    unique_output_path,
    apply_suggested_filename,
    show_placeholder_preview,
    pc_preview_rows,
    looks_like_digital_collector_log,
    parse_digital_collector_log,
    apply_log_device_fields_to_form,
    merge_log_device_into_report_data,
    load_request_titles,
    setup_title_combobox,
    remember_titles_from_form,
    setup_agency_combobox,
    remember_agencies_from_form,
    title_agency_words,
    saved_examiner_agency,
    bind_prefix_typeahead,
    refresh_request_title_values,
    current_dfr_prefix,
    is_complete_dfr_number,
    add_template_picker,
    sync_template_choice,
)
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from lxml import etree

#ctypes.windll.user32.ShowWindow(ctypes.windll.kernel32.GetConsoleWindow(), 0)

class PCPortableCase(TkinterDnD.Tk):
    def __init__(self, master=None):
        super().__init__()

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.master = master
        apply_theme(self)
        attach_app_menu(self)
        add_header_bar(self, "Computer Portable Case", "TX1, FTK, X-Ways, and Digital Collector image reports")
        self.action_bar = add_action_bar(self)
        self.preview_button = ttk.Button(self.action_bar, text="Preview", command=self.preview_placeholders)
        self.generate_button = ttk.Button(self.action_bar, text="Generate Report", command=self.generate_report)
        self.exit_button = ttk.Button(self.action_bar, text="Exit", command=self.on_closing)
        pack_right_actions(self.preview_button, self.generate_button, self.exit_button)
        
        # Initialize settings manager
        self.settings_manager = SettingsManager()
        self.current_settings = self.settings_manager.load_settings()
        self._save_timer = None 
        
        # Initialize PC exam specific paragraphs
        self.paragraphs = {}
        self._paragraph_kind = "pc_portable"
        self.initialize_pc_exam_paragraphs()
        
        self.title("Ω Digital Forensics Report Writer - Computer Portable Case Ω")
        size_window(self, 1260, 700)
        
        self.main_frame = ttk.Frame(self)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.paned_window = ttk.PanedWindow(self.main_frame, orient=tk.HORIZONTAL)
        self.paned_window.pack(fill=tk.BOTH, expand=True)
        
        self.left_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.left_frame, weight=3)

        self.form_tabs = FormTabs(self.left_frame)
        self.tab_request = self.form_tabs.add_tab("request", "Request Info")
        self.tab_examiner = self.form_tabs.add_tab("examiner", "Examiner Info")
        self.tab_device = self.form_tabs.add_tab("device", "Device Info")
        self.tab_output = self.form_tabs.add_tab("output", "Output Info")
        self.scrollable_frame = self.tab_request
        self.middle_frame = self.tab_device

        self.right_frame = ttk.Frame(self.paned_window)
        self.paned_window.add(self.right_frame, weight=2)
        
        self.create_widgets()
        self.create_info_display()
        
        self.bind_all("<MouseWheel>", self._on_mousewheel)
        self.bind_tab_status_events()
        self.refresh_tab_status()

    def _on_mousewheel(self, event):
        if hasattr(self, "form_tabs"):
            self.form_tabs.on_mousewheel(event)

    def create_widgets(self):
        """Main method to create all UI widgets - calls individual frame creation methods"""
        style = ttk.Style()
        style.configure('Header.TLabel', font=('Arial', 12, 'bold'))
        style.configure('Section.TFrame', padding='10')

        # Create each section in separate methods
        self.create_left_column_widgets()
        self.create_middle_column_widgets() 
        self.create_right_column_widgets()

    def _field_filled(self, widget):
        if widget is None:
            return False
        try:
            return bool(widget.get().strip())
        except Exception:
            return False

    def bind_tab_status_events(self):
        for name in (
            "request_date", "request_agency", "request_officer", "case_type",
            "examiner_name", "DFR_Num", "device_owner", "save_location",
        ):
            widget = getattr(self, name, None)
            if widget is None:
                continue
            try:
                widget.bind("<KeyRelease>", lambda e: self.refresh_tab_status(), add="+")
                widget.bind("<<ComboboxSelected>>", lambda e: self.refresh_tab_status(), add="+")
            except Exception:
                pass

    def refresh_tab_status(self):
        if not hasattr(self, "form_tabs"):
            return
        request_ok = all(self._field_filled(getattr(self, name, None)) for name in (
            "request_date", "request_agency", "request_officer", "case_type"
        ))
        examiner_ok = self._field_filled(getattr(self, "examiner_name", None)) and is_complete_dfr_number(self.DFR_Num.get() if hasattr(self, "DFR_Num") else "")
        device_ok = self._field_filled(getattr(self, "device_owner", None))
        try:
            device_ok = device_ok and bool(self.get_selected_forensic_software())
        except Exception:
            pass
        output_ok = self._field_filled(getattr(self, "save_location", None))
        if output_ok:
            try:
                output_ok = os.path.isdir(self.save_location.get().strip())
            except Exception:
                output_ok = False
        self.form_tabs.set_complete("request", request_ok)
        self.form_tabs.set_complete("examiner", examiner_ok)
        self.form_tabs.set_complete("device", device_ok)
        self.form_tabs.set_complete("output", output_ok)

    def create_left_column_widgets(self):
        """Create all widgets for the left column"""
        self.create_request_information_frame()
        self.create_transfer_information_frame()
        self.create_examiner_information_frame()

    def create_request_information_frame(self):
        """Create the request information frame (Agency Assist fields only)"""
        request_frame = ttk.LabelFrame(self.scrollable_frame, text="Request Information", padding="10")
        request_frame.pack(fill=tk.X, padx=5, pady=5)

        # Only create Agency Assist fields (no Case Agent frame)
        self.create_agency_assist_fields_in_frame(request_frame)

    def create_agency_assist_fields_in_frame(self, parent_frame):
        """Create agency assist specific fields directly in parent frame"""
        # Request Date
        ttk.Label(parent_frame, text="Exam Request Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.request_date = add_date_entry(parent_frame, row=0, column=1)

        # Request Agency
        ttk.Label(parent_frame, text="Requesting Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.request_agency = ttk.Combobox(parent_frame)
        self.request_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.request_agency)

        # Request Title
        ttk.Label(parent_frame, text="Requesting Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        request_title_frame = ttk.Frame(parent_frame)
        request_title_frame.grid(row=2, column=1, sticky="ew", pady=2)

        self.request_title_type = ttk.Combobox(request_title_frame, values=load_request_titles())
        self.request_title_type.pack(fill=tk.X, expand=True)
        self.request_title_type.set("")
        bind_prefix_typeahead(self.request_title_type)

        self.request_title_entry = ttk.Entry(request_title_frame)
        self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        self.request_title_entry.pack_forget()

        # Request Officer
        ttk.Label(parent_frame, text="Requesting Officer:").grid(row=3, column=0, sticky="w", pady=2)
        self.request_officer = ttk.Entry(parent_frame)
        self.request_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Offense Type
        ttk.Label(parent_frame, text="Primary Case Offense:").grid(row=4, column=0, sticky="w", pady=2)
        self.case_type = ttk.Entry(parent_frame)
        self.case_type.grid(row=4, column=1, sticky="ew", pady=2)

        # Legal Authority
        ttk.Label(parent_frame, text="Legal Authority:").grid(row=5, column=0, sticky="w", pady=2)
        self.legal_authority = ttk.Combobox(parent_frame, 
            values=['Search Warrant', 'Consent', 'Parole', 'Implied Consent'], state="readonly")
        self.legal_authority.grid(row=5, column=1, sticky="ew", pady=2)
        self.legal_authority.set('Search Warrant')
        self.legal_authority.bind('<<ComboboxSelected>>', self.toggle_time_frame_section)
        self.setup_auto_complete_dropdown(self.legal_authority)

        # Time Frame Section
        self.time_frame_label = ttk.Label(parent_frame, text="Time Frame Limited?:")
        self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)

        self.time_frame_checkbox_frame = ttk.Frame(parent_frame)
        self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)

        self.time_frame_var = tk.IntVar(parent_frame)
        self.time_frame_checkbox = ttk.Checkbutton(
            self.time_frame_checkbox_frame,
            text="Yes",
            variable=self.time_frame_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_time_frame_fields
        )
        self.time_frame_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        self.time_frame_var.set(0)

        # Time Frame Date Fields
        self.time_frame_fields = ttk.Frame(parent_frame)
        self.time_frame_fields.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        self.time_frame_fields.grid_remove()

        ttk.Label(self.time_frame_fields, text="Time Frame Start Date (M/D/Y):").grid(row=0, column=0, sticky="w", pady=2)
        self.time_frame_start = add_date_entry(self.time_frame_fields, row=0, column=1)

        ttk.Label(self.time_frame_fields, text="Time Frame End Date (M/D/Y):").grid(row=1, column=0, sticky="w", pady=2)
        self.time_frame_end = add_date_entry(self.time_frame_fields, row=1, column=1)

        # Configure grid weights
        self.time_frame_fields.columnconfigure(1, weight=1)
        parent_frame.columnconfigure(1, weight=1)

    def toggle_time_frame_section(self, event=None):
        """Show/hide time frame section based on legal authority selection for Agency Assist"""
        if self.legal_authority.get() == 'Search Warrant':
            # Show time frame section
            self.time_frame_label.grid(row=6, column=0, sticky="w", pady=2)
            self.time_frame_checkbox_frame.grid(row=6, column=1, sticky="ew", pady=2)
            # If checkbox is checked, also show the fields
            self.toggle_time_frame_fields()
        else:
            # Hide time frame section and clear values
            self.time_frame_label.grid_remove()
            self.time_frame_checkbox_frame.grid_remove()
            self.time_frame_fields.grid_remove()
            
            # Clear values
            self.time_frame_var.set(0)
            if hasattr(self, 'time_frame_start'):
                self.time_frame_start.delete(0, tk.END)
            if hasattr(self, 'time_frame_end'):
                self.time_frame_end.delete(0, tk.END)

    def toggle_time_frame_fields(self, event=None):
        """Show/hide time frame date fields based on checkbox state for Agency Assist"""
        if self.time_frame_var.get() == 1:
            self.time_frame_fields.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            self.time_frame_fields.grid_remove()
            # Clear values when hiding
            if hasattr(self, 'time_frame_start'):
                self.time_frame_start.delete(0, tk.END)
            if hasattr(self, 'time_frame_end'):
                self.time_frame_end.delete(0, tk.END)

    def toggle_case_agent_time_frame_fields(self, event=None):
        """Show/hide time frame date fields based on checkbox state for Case Agent"""
        if self.case_agent_time_frame_var.get() == 1:
            self.case_agent_time_frame_fields.grid(row=4, column=0, columnspan=2, sticky="ew", pady=2)
        else:
            self.case_agent_time_frame_fields.grid_remove()
            # Clear values when hiding
            if hasattr(self, 'case_agent_time_frame_start'):
                self.case_agent_time_frame_start.delete(0, tk.END)
            if hasattr(self, 'case_agent_time_frame_end'):
                self.case_agent_time_frame_end.delete(0, tk.END)

    def create_transfer_information_frame(self):
        """Create the transfer information frame"""
        transfer_frame = ttk.LabelFrame(self.scrollable_frame, text="Transfer Information", padding="10")
        transfer_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Transfer - Changed from Combobox to Single Yes Checkbox
        ttk.Label(transfer_frame, text="Device Transferred by Another Officer:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(transfer_frame)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variable
        self.device_transfer_var = tk.IntVar(transfer_frame)
        
        # Create single Yes checkbox
        self.transfer_yes_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Yes",
            variable=self.device_transfer_var,
            onvalue=1,
            offvalue=0,
            command=self.toggle_transfer_fields
        )
        self.transfer_yes_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        # Set default to unchecked (0 = No)
        self.device_transfer_var.set(0)

        # Transfer Fields (initially hidden)
        self.transfer_fields = ttk.Frame(transfer_frame)
        self.transfer_fields.grid(row=1, column=0, columnspan=2, sticky="ew", pady=2)
        self.transfer_fields.grid_remove()

        ttk.Label(self.transfer_fields, text="Transfer Date:").grid(row=0, column=0, sticky="w", pady=2)
        self.transfer_date = add_date_entry(self.transfer_fields, row=0, column=1)

        ttk.Label(self.transfer_fields, text="Transfer Officer Agency:").grid(row=1, column=0, sticky="w", pady=2)
        self.transfer_agency = ttk.Combobox(self.transfer_fields)
        self.transfer_agency.grid(row=1, column=1, sticky="ew", pady=2)
        setup_agency_combobox(self.transfer_agency)

        ttk.Label(self.transfer_fields, text="Transfer Officer Title:").grid(row=2, column=0, sticky="w", pady=2)
        self.transfer_title = ttk.Combobox(self.transfer_fields, values=load_request_titles())
        self.transfer_title.grid(row=2, column=1, sticky="ew", pady=2)
        setup_title_combobox(self.transfer_title)

        ttk.Label(self.transfer_fields, text="Transfer Officer Name:").grid(row=3, column=0, sticky="w", pady=2)
        self.transfer_officer = ttk.Entry(self.transfer_fields)
        self.transfer_officer.grid(row=3, column=1, sticky="ew", pady=2)

        # Configure grid weights
        transfer_frame.columnconfigure(1, weight=1)
        self.transfer_fields.columnconfigure(1, weight=1)

    def create_examiner_information_frame(self):
        examiner_frame = ttk.LabelFrame(self.tab_examiner, text="Examiner Information", padding="10")
        examiner_frame.pack(fill=tk.X, padx=5, pady=5)

        # Examiner Agency
        ttk.Label(examiner_frame, text="Examiner Agency:").grid(row=0, column=0, sticky="w", pady=2)
        examiner_agency_frame = ttk.Frame(examiner_frame)
        examiner_agency_frame.grid(row=0, column=1, sticky="ew", pady=2)

        self.examiner_agency_type = ttk.Combobox(examiner_agency_frame)
        self.examiner_agency_type.pack(fill=tk.X, expand=True)
        setup_agency_combobox(
            self.examiner_agency_type,
            saved=saved_examiner_agency(self.current_settings),
            on_change=self.auto_save_settings,
        )

        self.examiner_agency_entry = ttk.Entry(examiner_agency_frame)
        self.examiner_agency_entry.pack_forget()

        # Examiner Title
        ttk.Label(examiner_frame, text="Examiner Title:").grid(row=1, column=0, sticky="w", pady=2)
        examiner_title_frame = ttk.Frame(examiner_frame)
        examiner_title_frame.grid(row=1, column=1, sticky="ew", pady=2)

        self.examiner_title_type = ttk.Combobox(examiner_title_frame, values=load_request_titles())
        self.examiner_title_type.pack(fill=tk.X, expand=True)
        setup_title_combobox(
            self.examiner_title_type,
            saved=self.current_settings.get("examiner_title", ""),
            on_change=self.auto_save_settings,
        )

        self.examiner_title_entry = ttk.Entry(examiner_title_frame)
        self.examiner_title_entry.pack_forget()

        # Examiner Name
        ttk.Label(examiner_frame, text="Examiner Name:").grid(row=2, column=0, sticky="w", pady=2)
        self.examiner_name = ttk.Entry(examiner_frame)
        self.examiner_name.grid(row=2, column=1, sticky="ew", pady=2)
        # Load saved setting
        if self.current_settings.get("examiner_name"):
            self.examiner_name.insert(0, self.current_settings["examiner_name"])
        
        # Bind auto-save events
        self.examiner_name.bind('<KeyRelease>', self.auto_save_settings)
        self.examiner_name.bind('<FocusOut>', self.auto_save_settings)

        # Case Number
        ttk.Label(examiner_frame, text="Case Number:").grid(row=3, column=0, sticky="w", pady=2)
        self.case_number = ttk.Entry(examiner_frame)
        self.case_number.grid(row=3, column=1, sticky="ew", pady=2)

        # Evidence Number
        ttk.Label(examiner_frame, text="Evidence Number:").grid(row=4, column=0, sticky="w", pady=2)
        self.evidence_number = ttk.Entry(examiner_frame)
        self.evidence_number.grid(row=4, column=1, sticky="ew", pady=2)

        # DFR Report Number
        ttk.Label(examiner_frame, text="DFR Report #:").grid(row=5, column=0, sticky="w", pady=2)
        self.DFR_Num = ttk.Entry(examiner_frame)
        self.DFR_Num.grid(row=5, column=1, sticky="ew", pady=2)
        
        # Load saved prefix or use default
        self.DFR_Num.insert(0, current_dfr_prefix())

        examiner_frame.columnconfigure(1, weight=1)


    def toggle_examiner_title_entry(self, event=None):
        if self.examiner_title_type.get() == "Other (specify)":
            self.examiner_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.examiner_title_entry.pack_forget()
            self.examiner_title_entry.delete(0, tk.END)
        
        # Auto-save when selection changes
        self.auto_save_settings()

    def on_agency_type_changed(self, event=None):
        # Call the original toggle method
        self.toggle_agency_entry(event)
        # Auto-save the change
        self.auto_save_settings()

    def auto_save_settings(self, event=None):
        # Use after_idle to prevent excessive saving during rapid typing
        if hasattr(self, '_save_timer') and self._save_timer:
            self.after_cancel(self._save_timer)
        
        # Schedule save for after user stops typing (500ms delay)
        self._save_timer = self.after(500, self._perform_auto_save)

    def _perform_auto_save(self):
        try:
            # Get examiner title value based on dropdown selection
            examiner_title_value = self.examiner_title_type.get()
                
            settings_to_save = {
                "examiner_agency_type": self.examiner_agency_type.get(),
                "examiner_agency_custom": "",
                "examiner_title": examiner_title_value,
                "examiner_name": self.examiner_name.get(),
                "dfr_number_prefix": self.get_dfr_prefix(),
                "version": "1.0.1"
            }
            
            self.settings_manager.save_settings(settings_to_save)
            # Update current_settings to reflect the save
            self.current_settings.update(settings_to_save)
            remember_agencies_from_form(self)
            
        except Exception as e:
            print(f"Error auto-saving examiner settings: {e}")

    def get_dfr_prefix(self):
        dfr_value = self.DFR_Num.get()
        # Find the last occurrence of "DFR" and include everything up to and including any year/dash pattern
        import re
        match = re.match(r'(DFR\d{4}-)', dfr_value)
        if match:
            return match.group(1)
        else:
            # Fallback to current year
            from datetime import datetime
            current_year = datetime.now().year
            return f"DFR{current_year}-"

    def create_middle_column_widgets(self):
        self.create_forensic_software_frame()
        self.create_device_info_frame()
        self.create_output_file_frame()

    def create_forensic_software_frame(self):
        forensic_software_frame = ttk.LabelFrame(self.tab_device, text="Forensic Software", padding="10")
        forensic_software_frame.pack(fill=tk.X, padx=5, pady=5)

        forensic_software_content = ttk.Frame(forensic_software_frame)
        forensic_software_content.pack(fill=tk.X, expand=True)

        # Forensic Processing Software - Using checkboxes
        ttk.Label(forensic_software_content, text="Forensic Processing Software:").grid(row=0, column=0, sticky="w", pady=2, padx=(0, 10))
        
        # Create checkbox frame
        checkbox_frame = ttk.Frame(forensic_software_content)
        checkbox_frame.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Create checkbox variables (removed Griffeye)
        self.axiom_var = tk.IntVar(forensic_software_frame)
        self.xways_var = tk.IntVar(forensic_software_frame)
        
        # Create checkboxes (removed Griffeye)
        self.axiom_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="Axiom",
            variable=self.axiom_var,
            onvalue=1,
            offvalue=0
        )
        self.axiom_checkbox.grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        self.xways_checkbox = ttk.Checkbutton(
            checkbox_frame,
            text="X-Ways",
            variable=self.xways_var,
            onvalue=1,
            offvalue=0
        )
        self.xways_checkbox.grid(row=0, column=1, sticky="w", padx=(0, 10))
        
        forensic_software_content.columnconfigure(1, weight=1)

    def on_forensic_software_change(self):
        # Get selected software
        selected_software = []
        if self.axiom_var.get():
            selected_software.append("Axiom")
        if self.xways_var.get():
            selected_software.append("X-Ways")
        
        # Set the main forensic software variable (for backwards compatibility)
        if selected_software:
            self.forensic_software = type('obj', (object,), {'get': lambda: ', '.join(selected_software)})()
        else:
            self.forensic_software = type('obj', (object,), {'get': lambda: ''})()

    def create_device_info_frame(self):
        device_frame = ttk.LabelFrame(self.tab_device, text="Device Information", padding="10")
        device_frame.pack(fill=tk.X, padx=5, pady=5)

        # Device Owner
        ttk.Label(device_frame, text="Device Owner:").grid(row=0, column=0, sticky="w", pady=2)
        self.device_owner = ttk.Entry(device_frame)
        self.device_owner.grid(row=0, column=1, sticky="ew", pady=2)

        # Device Type with dropdown
        ttk.Label(device_frame, text="Device Type:").grid(row=1, column=0, sticky="w", pady=2)
        
        # Create a frame to contain the combobox and optional entry field
        device_type_frame = ttk.Frame(device_frame)
        device_type_frame.grid(row=1, column=1, sticky="ew", pady=2)
        device_type_frame.columnconfigure(0, weight=1)  # Make sure this frame expands properly
        
        self.device_type = ttk.Combobox(device_type_frame, 
            values=["Computer", "Loose Hard Drive", "USB Drive", "Memory Card", "Other Storage Device"], 
            state="readonly")
        self.device_type.pack(fill=tk.X, expand=True)
        self.device_type.set("Computer")
        self.device_type.bind("<<ComboboxSelected>>", self.toggle_device_type_fields)
        self.setup_auto_complete_dropdown(self.device_type)
        
        # Custom device type entry field (initially hidden)
        self.device_type_entry = ttk.Entry(device_type_frame)
        self.device_type_entry.pack(fill=tk.X, expand=True, pady=2)
        self.device_type_entry.pack_forget()  # Hide initially

        # Device Make
        ttk.Label(device_frame, text="Device Make:").grid(row=2, column=0, sticky="w", pady=2)
        self.device_PCMan = ttk.Entry(device_frame)
        self.device_PCMan.grid(row=2, column=1, sticky="ew", pady=2)

        # Device Model
        ttk.Label(device_frame, text="Device Model:").grid(row=3, column=0, sticky="w", pady=2)
        self.device_PCMod = ttk.Entry(device_frame)
        self.device_PCMod.grid(row=3, column=1, sticky="ew", pady=2)

        # Serial Number
        ttk.Label(device_frame, text="Device Serial Number:").grid(row=4, column=0, sticky="w", pady=2)
        self.device_PCSerial = ttk.Entry(device_frame)
        self.device_PCSerial.grid(row=4, column=1, sticky="ew", pady=2)

        # Device Color
        ttk.Label(device_frame, text="Device Color:").grid(row=5, column=0, sticky="w", pady=2)
        self.device_color = ttk.Entry(device_frame)
        self.device_color.grid(row=5, column=1, sticky="ew", pady=2)

        # Device Password
        ttk.Label(device_frame, text="Device Password:").grid(row=6, column=0, sticky="w", pady=2)
        self.device_password = ttk.Entry(device_frame)
        self.device_password.grid(row=6, column=1, sticky="ew", pady=2)

        # Computer-specific hard drive fields
        self.computer_fields_frame = ttk.Frame(device_frame)
        self.computer_fields_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)

        # Hard Drive Make
        ttk.Label(self.computer_fields_frame, text="Hard Drive Make:").grid(row=0, column=0, sticky="w", pady=2)
        self.hd_make = ttk.Entry(self.computer_fields_frame)
        self.hd_make.grid(row=0, column=1, sticky="ew", pady=2)

        # Hard Drive Model
        ttk.Label(self.computer_fields_frame, text="Hard Drive Model:").grid(row=1, column=0, sticky="w", pady=2)
        self.hd_model = ttk.Entry(self.computer_fields_frame)
        self.hd_model.grid(row=1, column=1, sticky="ew", pady=2)

        # Hard Drive Serial Number
        ttk.Label(self.computer_fields_frame, text="Hard Drive Serial:").grid(row=2, column=0, sticky="w", pady=2)
        self.hd_serial = ttk.Entry(self.computer_fields_frame)
        self.hd_serial.grid(row=2, column=1, sticky="ew", pady=2)

        # Configure grid weights for computer fields frame - THIS IS THE KEY FIX
        # Make the label column the same width as the main device_frame labels
        self.computer_fields_frame.columnconfigure(0, weight=0, minsize=120)  # Fixed width for labels
        self.computer_fields_frame.columnconfigure(1, weight=1)  # Entry fields expand

        # Storage Capacity - SINGLE FIELD that moves position based on device type
        self.capacity_label = ttk.Label(device_frame, text="Storage Capacity:")
        self.device_capacity = ttk.Entry(device_frame)

        # Configure grid weights for main device frame
        device_frame.columnconfigure(0, weight=0, minsize=120)  # Fixed width for labels to match
        device_frame.columnconfigure(1, weight=1)  # Entry fields expand
        
        # Show appropriate fields initially since "Computer" is the default
        self.toggle_device_type_fields()
        
        return device_frame
    
    def create_output_file_frame(self):
        output_frame = ttk.LabelFrame(self.tab_output, text="Output File", padding="10")
        output_frame.pack(fill=tk.X, padx=5, pady=5)

        output_content = ttk.Frame(output_frame)
        output_content.pack(fill=tk.X, expand=True)
        
        # Output filename
        ttk.Label(output_content, text="Output File Name:").grid(row=0, column=0, sticky="w", pady=2)
        self.output_filename = ttk.Entry(output_content)
        self.output_filename.grid(row=0, column=1, sticky="ew", pady=2)
        
        # Hint for filename (directly under the filename field)
        hint_style = ttk.Style()
        hint_style.configure("Hint.TLabel", font=('Arial', 7))  
        hint_label = ttk.Label(output_content, 
                              text="*Default File Name: \"(DFR #) - (Owner Name) (Device Model)\"", 
                              style="Hint.TLabel",
                              wraplength=350)
        hint_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 5))
        
        # Save location
        ttk.Label(output_content, text="Save Location:").grid(row=2, column=0, sticky="w", pady=2)
        
        # Frame for save location entry and browse button
        location_frame = ttk.Frame(output_content)
        location_frame.grid(row=2, column=1, sticky="ew", pady=2)
        location_frame.columnconfigure(0, weight=1)
        
        # Save location entry (initialize with desktop path)
        self.save_location = ttk.Entry(location_frame)
        self.save_location.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        
        # Set default to desktop
        self.save_location.insert(0, default_export_dir())
        
        # Browse button
        browse_button = ttk.Button(location_frame, text="Browse", command=self.browse_save_location)
        browse_button.grid(row=0, column=1)
        
        output_content.columnconfigure(1, weight=1)

    def browse_save_location(self):
        """Allow user to browse and select a save location"""
        try:
            # Get current path from the entry, default to desktop if empty
            current_path = self.save_location.get().strip()
            if not current_path or not os.path.exists(current_path):
                current_path = os.path.join(os.path.expanduser("~"), "Desktop")
            
            # Open folder dialog
            selected_path = ask_directory("export", current_path, "Select Save Location")
            
            if selected_path:
                # Update the entry with selected path
                self.save_location.delete(0, tk.END)
                self.save_location.insert(0, selected_path)
                
        except Exception as e:
            messagebox.showerror("Error", f"Could not browse for save location: {str(e)}")

    def create_right_column_widgets(self):
        self.create_file_upload_frame()

    def create_file_upload_frame(self):
        file_frame = ttk.LabelFrame(self.right_frame, text="File Upload", padding="10")
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        add_template_picker(self, file_frame, preferred="DFR Computer (2026).docx", keywords=("pc", "storage", "computer"))

        # Extraction File Drag-and-Drop
        self.create_extraction_file_ui(file_frame)    

    def toggle_device_type_fields(self, event=None):
        selected_type = self.device_type.get()
        
        if selected_type == "Other Storage Device":
            # Show custom entry field for other storage devices
            self.device_type_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            # Hide custom entry field
            self.device_type_entry.pack_forget()
        
        if selected_type == "Computer":
            # Show computer-specific hard drive fields
            self.computer_fields_frame.grid(row=7, column=0, columnspan=2, sticky="ew", pady=2)
            # Position capacity field after hard drive fields (row 8)
            self.capacity_label.grid(row=8, column=0, sticky="w", pady=2)
            self.device_capacity.grid(row=8, column=1, sticky="ew", pady=2)
        else:
            # Hide computer-specific fields and clear their values
            self.computer_fields_frame.grid_remove()
            if hasattr(self, 'hd_make'):
                self.hd_make.delete(0, tk.END)
            if hasattr(self, 'hd_model'):
                self.hd_model.delete(0, tk.END)
            if hasattr(self, 'hd_serial'):
                self.hd_serial.delete(0, tk.END)
            
            # Position capacity field after device password (row 7)
            self.capacity_label.grid(row=7, column=0, sticky="w", pady=2)
            self.device_capacity.grid(row=7, column=1, sticky="ew", pady=2)
    
    def create_extraction_file_ui(self, file_frame):
        # Extraction File Drag-and-Drop
        self.extraction_drop_label = ttk.Label(file_frame, text="Drag & Drop Log File or Click to Browse",
                                      relief="solid", borderwidth=2, padding=10)
        self.extraction_drop_label.pack(fill=tk.X, pady=10, padx=5)
        self.extraction_drop_label.bind('<Button-1>', self.browse_extraction_file)
        self.extraction_drop_label.drop_target_register(DND_FILES)
        self.extraction_drop_label.dnd_bind('<<Drop>>', self.drop_extraction_file)

    def create_info_display(self):
        self.info_frame, self.info_display = build_extracted_info_pane(
            self.right_frame,
            "Add a log file (.txt) to see the extracted information.",
        )
        self.no_evidence_var = tk.IntVar(self)
        self.no_evidence_checkbox = ttk.Checkbutton(
            self.info_frame,
            text="No Evidence Found",
            variable=self.no_evidence_var,
            onvalue=1,
            offvalue=0,
        )
        self.no_evidence_checkbox.pack(anchor="w", pady=(6, 0))

    def update_info_display(self, message=None, data=None):
        self.info_display.config(state=tk.NORMAL)
        self.info_display.delete(1.0, tk.END)
        
        if message:
            self.info_display.insert(tk.END, message)
        
        if data:
            # Add file type header based on extraction type
            if hasattr(self, "extraction_type"):
                if self.extraction_type == "TX1":
                    self.info_display.insert(tk.END, "TX1 Log File Detected\n", "header")
                elif self.extraction_type == "FTK":
                    self.info_display.insert(tk.END, "FTK Imager Log File Detected\n", "header")
                elif self.extraction_type == "XWAYS":
                    self.info_display.insert(tk.END, "X-Ways Forensics Log File Detected\n", "header")
                elif self.extraction_type == "DC":
                    self.info_display.insert(tk.END, "Cellebrite Digital Collector Log Detected\n", "header")
            
            self.info_display.insert(tk.END, "\nExtracted Information:\n\n", "section")
            
            display_data = data.copy()
            
            exclude_keys = ['article', 'extraction_date', 'examiner']
            
            for key in exclude_keys:
                if key in display_data:
                    del display_data[key]
            
            # Determine extraction type for display formatting
            is_tx1 = self.extraction_type == "TX1" if hasattr(self, "extraction_type") else False
            is_ftk = self.extraction_type == "FTK" if hasattr(self, "extraction_type") else False
            is_xways = self.extraction_type == "XWAYS" if hasattr(self, "extraction_type") else False
            is_dc = self.extraction_type == "DC" if hasattr(self, "extraction_type") else False
            
            # Define the display order based on extraction type
            if is_tx1:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('TX1_OS', 'TX1 Version'),
                    ('extraction_serial', 'TX1 Serial Number'),
                    ('case_number', 'Case Number'),
                    ('case_id', 'Case ID'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Case Notes'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_ftk:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('FTK_OS', 'FTK Imager Version'),
                    ('case_number', 'Case Number'),
                    ('evidence_number', 'Evidence Number'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Description'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_xways:
                display_order = [   
                    ('formatted_date', 'Extraction Date/Time'),
                    ('xways_OS', 'X-Ways Version'),
                    ('device_model', 'Source Device Model'),
                    ('case_notes', 'Internal Description'),
                    ('md5_hash', 'MD5 Hash'),
                ]
            elif is_dc:
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('DC_OS', 'Digital Collector Version'),
                    ('case_number', 'Case Number'),
                    ('case_name', 'Case Name'),
                    ('evidence_number', 'Exhibit ID / Evidence #'),
                    ('device_model', 'Hard Drive Model'),
                    ('device_serial', 'Hard Drive Serial'),
                    ('device_capacity', 'Capacity'),
                    ('case_notes', 'Description'),
                    ('image_format', 'Image Format'),
                    ('md5_hash', 'MD5 Hash'),
                    ('sha1_hash', 'SHA1 Hash'),
                ]
            else:
                # Generic fallback display order
                display_order = [
                    ('formatted_date', 'Extraction Date/Time'),
                    ('device_model', 'Source Device Model'),
                ]
            
            displayed_keys = set()
            
            for key_pair in display_order:
                key, display_name = key_pair
                if key in display_data and display_data[key] and key not in displayed_keys:
                    value = display_data[key]
                    self.info_display.insert(tk.END, f"{display_name}: ", "key")
                    self.info_display.insert(tk.END, f"{value}\n", "value")
                    displayed_keys.add(key)
             
            # Configure text tags for styling
            self.info_display.tag_configure("header", font=("Arial", 12, "bold"), foreground="green", justify="center")
            self.info_display.tag_configure("key", font=("Arial", 10, "bold"))
            self.info_display.tag_configure("value", font=("Arial", 10))
            self.info_display.tag_configure("section", font=("Arial", 11, "bold"))
        
        self.info_display.config(state=tk.DISABLED)

    def parse_extraction_file(self):
        if hasattr(self, 'extraction_type'):
            return self.parse_log_file()
        else:
            messagebox.showerror("Error", "No extraction file selected")
            return {}

    def toggle_request_content(self, event=None):
        # Clear all request information fields when role is changed
        if hasattr(self, 'request_date'):
            self.request_date.delete(0, tk.END)
        if hasattr(self, 'request_agency'):
            self.request_agency.delete(0, tk.END)
        if hasattr(self, 'request_title_type'):
            self.request_title_type.set("")  # Reset to default
        if hasattr(self, 'request_title_entry'):
            self.request_title_entry.delete(0, tk.END)
            self.request_title_entry.pack_forget()  # Hide custom entry
        if hasattr(self, 'request_officer'):
            self.request_officer.delete(0, tk.END)
        if hasattr(self, 'case_type'):
            self.case_type.delete(0, tk.END)
        if hasattr(self, 'legal_authority'):
            self.legal_authority.set('Search Warrant')  # Reset to default
        if hasattr(self, 'offense_type'):
            self.offense_type.delete(0, tk.END)
        if hasattr(self, 'legal_self'):
            self.legal_self.set('Search Warrant')  # Reset to default
        if hasattr(self, 'sw_service_date'):
            self.sw_service_date.delete(0, tk.END)
        
        # Clear time frame fields for both roles
        if hasattr(self, 'time_frame_var'):
            self.time_frame_var.set(0)
        if hasattr(self, 'time_frame_start'):
            self.time_frame_start.delete(0, tk.END)
        if hasattr(self, 'time_frame_end'):
            self.time_frame_end.delete(0, tk.END)
        if hasattr(self, 'case_agent_time_frame_var'):
            self.case_agent_time_frame_var.set(0)
        if hasattr(self, 'case_agent_time_frame_start'):
            self.case_agent_time_frame_start.delete(0, tk.END)
        if hasattr(self, 'case_agent_time_frame_end'):
            self.case_agent_time_frame_end.delete(0, tk.END)
        
        # Hide search warrant date fields when switching roles
        if hasattr(self, 'sw_date_frame'):
            for widget in self.sw_date_frame.winfo_children():
                widget.grid_remove()
        
        # Portable case is Agency Assist only. Guard leftover Case Agent UI hooks.
        if not hasattr(self, "role_type") or self.role_type is None:
            if hasattr(self, "toggle_time_frame_section"):
                self.toggle_time_frame_section()
            return

        if self.role_type.get() == "Agency Assist":
            if hasattr(self, "agency_assist_frame"):
                self.agency_assist_frame.pack(fill=tk.X, padx=5, pady=5)
            if hasattr(self, "case_agent_frame"):
                self.case_agent_frame.pack_forget()
            if hasattr(self, "toggle_time_frame_section"):
                self.toggle_time_frame_section()
        else:
            if hasattr(self, "agency_assist_frame"):
                self.agency_assist_frame.pack_forget()
            if hasattr(self, "case_agent_frame"):
                self.case_agent_frame.pack(fill=tk.X, padx=5, pady=5)
            if hasattr(self, "toggle_sw_date_and_time_frame"):
                self.toggle_sw_date_and_time_frame()

    def toggle_title_entry(self, event=None):
        if self.request_title_type.get() == "Other (specify)":
            self.request_title_entry.pack(fill=tk.X, expand=True, pady=2)
        else:
            self.request_title_entry.pack_forget()

    def toggle_agency_entry(self, event=None):
        return

    def toggle_transfer_fields(self, event=None):
        if self.device_transfer_var.get() == 1:
            self.transfer_fields.grid()
        else:
            self.transfer_fields.grid_remove()
            if hasattr(self, 'transfer_title'):
                self.transfer_title.delete(0, tk.END)
            if hasattr(self, 'transfer_officer'):
                self.transfer_officer.delete(0, tk.END)
            if hasattr(self, 'transfer_agency'):
                self.transfer_agency.delete(0, tk.END)
            if hasattr(self, 'transfer_date'):
                self.transfer_date.delete(0, tk.END)

    def setup_auto_complete_dropdown(self, combobox):
        combobox.bind('<KeyRelease>', lambda event, cb=combobox: self.auto_complete(event, cb))

    def auto_complete(self, event, combobox):
        if event.keysym in ('Up', 'Down', 'Left', 'Right', 'Return', 'Tab', 'BackSpace', 'Delete'):
            return
        
        value = combobox.get()
        
        if not value:
            return
        
        all_values = combobox['values']
        
        previous_value = combobox.get()
        
        for option in all_values:
            if option.lower().startswith(value.lower()):
                combobox.set(option)
                combobox.icursor(len(option))
                combobox.selection_range(len(value), len(option))
                           
                break

    def drop_extraction_file(self, event):
        file_path = self.tk.splitlist(event.data)[0]
        
        if file_path.lower().endswith('.txt'):
            # Clear previously populated fields from log files
            self.clear_log_populated_fields()
            
            # Set the new extraction file
            self.extraction_file = file_path
            self.extraction_drop_label.configure(text=f"Selected: {os.path.basename(file_path)} (Forensic Log)")
            
            # Parse the file and update the display
            extraction_data = self.parse_log_file()
            apply_log_device_fields_to_form(self, extraction_data)
            self.update_info_display(data=extraction_data)
            
            # No popup messages - the file type will be shown in the display window header
        else:
            messagebox.showerror("Error", "Please select a TXT log file (TX1, FTK, X-Ways, or Digital Collector)")
            self.update_info_display("Error: Please select a valid TXT log file.")

    def browse_extraction_file(self, event):
        file_path = ask_open_file([("Log files", "*.txt")], folder_kind="extraction", title="Select image log")
        
        if file_path:
            if file_path.lower().endswith('.txt'):
                # Clear previously populated fields from log files
                self.clear_log_populated_fields()
                
                # Set the new extraction file
                self.extraction_file = file_path
                self.extraction_drop_label.configure(text=f"Selected: {os.path.basename(file_path)} (Forensic Log)")
                
                # Parse the file and update the display
                extraction_data = self.parse_log_file()
                apply_log_device_fields_to_form(self, extraction_data)
                self.update_info_display(data=extraction_data)
                
                # No popup messages - the file type will be shown in the display window header
            else:
                messagebox.showerror("Error", "Please select a TXT log file")

    def clear_log_populated_fields(self):        
        # Clear any previously stored extraction data
        if hasattr(self, 'extraction_data'):
            delattr(self, 'extraction_data')
        
        # Clear extraction type
        if hasattr(self, 'extraction_type'):
            delattr(self, 'extraction_type')
        
        # Clear the set of populated fields
        if hasattr(self, 'log_populated_fields'):
            self.log_populated_fields.clear()
        
        # Reset extraction file reference if it exists
        if hasattr(self, 'extraction_file'):
            # Don't delete the file path, just clear any cached data
            pass
     
    def drop_template_file(self, event):
        file_path = self.tk.splitlist(event.data)[0]
        if file_path.endswith('.docx'):
            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        else:
            messagebox.showerror("Error", "Please select a Word document template")
     
    def browse_template_file(self, event):
        file_path = ask_open_file([("Word documents", "*.docx")], folder_kind="template", title="Select DFR template")
        if file_path:
            self.template_file = file_path
            sync_template_choice(self, file_path)
            self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")

    def add_bold_underline_paragraph(self, doc, text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.underline = True
        run.font.name = 'Arial'  # This line already exists
        run.font.size = Pt(11)   # This line already exists
        return p

    def parse_log_file(self):
        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            messagebox.showerror("Error", "No log file selected")
            return {}
            
        try:
            with open(self.extraction_file, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Initialize/reset the set of fields that were populated from the log
                if not hasattr(self, 'log_populated_fields'):
                    self.log_populated_fields = set()
                else:
                    self.log_populated_fields.clear()
                
                # Determine the log type
                if "TX1 Log Entry" in content:
                    self.extraction_type = "TX1"
                    return self.parse_tx1_log(content)
                elif "FTK® Imager" in content or "FTK Imager" in content:
                    self.extraction_type = "FTK"
                    return self.parse_ftk_log(content)
                elif "X-Ways Forensics" in content:
                    self.extraction_type = "XWAYS"
                    return self.parse_xways_log(content)
                elif looks_like_digital_collector_log(content):
                    self.extraction_type = "DC"
                    data, populated = parse_digital_collector_log(content)
                    self.log_populated_fields.update(populated)
                    return data
                else:
                    messagebox.showwarning("Unknown Log Format", 
                                         "The selected file doesn't appear to be a TX1, FTK Imager, X-Ways, or Digital Collector log.")
                    return {}
        except Exception as e:
            messagebox.showerror("Error", f"Failed to parse log file: {str(e)}")
            return {}

    def parse_xways_log(self, content):
        extraction_data = {
            'extraction_tool': 'X-Ways Forensics',
            'extraction_type': 'Disk Imaging',  # Default value
        }
        
        lines = content.split('\n')
        
        # Extract acquisition start date from the first line
        if lines:
            first_line = lines[0].strip()
            # Remove BOM character if present
            first_line = first_line.lstrip('\ufeff')
            
            # Format: 12/13/2024, 10:19:06
            extraction_data['extraction_date'] = first_line  # Store original (cleaned)
            extraction_data['formatted_date'] = self.parse_xways_date(first_line)
            self.log_populated_fields.add('extraction_date')
            self.log_populated_fields.add('formatted_date')
        
        # Extract X-Ways version from second line
        if len(lines) > 1:
            second_line = lines[1].strip()
            # Format: X-Ways Forensics 20.9 SR-5 x64
            version_match = re.search(r'X-Ways Forensics\s+(\d+\.\d+)', second_line)
            if version_match:
                extraction_data['xways_OS'] = version_match.group(1)
                self.log_populated_fields.add('xways_OS')
        
        # Extract other fields line by line
        for line in lines:
            line = line.strip()
            
            # Extract Model
            if line.startswith('Model:'):
                model = line[len('Model:'):].strip()
                extraction_data['device_model'] = model
                self.log_populated_fields.add('device_model')
            
            # Extract Serial Number
            elif line.startswith('Serial No.:'):
                serial = line[len('Serial No.:'):].strip()
                # X-Ways shows serial like: 0000000000000000 / 0xecc4800670043
                # Take the part before the slash if there is one
                if '/' in serial:
                    serial = serial.split('/')[0].strip()
                extraction_data['device_serial'] = serial
                self.log_populated_fields.add('device_serial')
            
            # Extract Internal description (this could be used as case notes or additional device info)
            elif line.startswith('Internal description:'):
                description = line[len('Internal description:'):].strip()
                extraction_data['case_notes'] = description
                self.log_populated_fields.add('case_notes')
            
            # Extract MD5 hash
            elif line.startswith('Hash of source data:') and 'MD5' in line:
                # Format: Hash of source data: 640605156B461F053675C093173C7648 (MD5)
                hash_match = re.search(r'Hash of source data:\s*([a-fA-F0-9]+)\s*\(MD5\)', line)
                if hash_match:
                    extraction_data['md5_hash'] = hash_match.group(1)
                    self.log_populated_fields.add('md5_hash')
            
            # Extract Examiner(s) field
            elif line.startswith('Examiner(s):'):
                examiner = line[len('Examiner(s):'):].strip()
                extraction_data['examiner'] = examiner
                self.log_populated_fields.add('examiner')
        
        # If no formatted date was set, set a default
        if 'formatted_date' not in extraction_data:
            extraction_data['formatted_date'] = "Unknown Date"
            self.log_populated_fields.add('formatted_date')
        
        return extraction_data
    
    def parse_xways_date(self, date_str):
        try:
            # Remove BOM character if present
            date_str = date_str.lstrip('\ufeff')
            
            # Split by comma to get date and time parts
            if ', ' in date_str:
                date_part, time_part = date_str.split(', ')
            else:
                # If no comma, assume the whole string is just the date
                date_part = date_str.strip()
                time_part = "00:00:00"  # Default time
            
            # Parse the date and time using datetime
            full_datetime_str = f"{date_part.strip()} {time_part.strip()}"
            date_obj = datetime.strptime(full_datetime_str, '%m/%d/%Y %H:%M:%S')
            
            # Format the date for output 
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing X-Ways date: {e}")
            return date_str
    
    def parse_tx1_log(self, content):
        extraction_data = {
            'extraction_tool': 'TX1',
            'extraction_type': 'Disk Duplication',  # Default value
        }
        
        # Handle special parsing for Case ID and Case Notes (which may be indented on next line)
        # First, split the content into sections by looking for lines that start with hyphens
        sections = []
        current_section = []
        
        for line in content.split('\n'):
            if line.startswith('----'):  # Section separator
                if current_section:
                    sections.append('\n'.join(current_section))
                    current_section = []
            current_section.append(line)
        
        if current_section:
            sections.append('\n'.join(current_section))
        
        # Look for the section with Case ID and Case Notes
        for section in sections:
            if 'Case ID:' in section:
                # Extract Case ID - store it as case_id, not case_number
                case_id_match = re.search(r'Case ID:\s*\n?\s*(.*?)(?:\n\s*[A-Za-z]|\n\n)', section, re.DOTALL)
                if case_id_match:
                    case_id = case_id_match.group(1).strip()
                    extraction_data['case_id'] = case_id
                    # Track that this field was populated from the log
                    self.log_populated_fields.add('case_id')
                
                # Extract Case Notes (which typically follow Case ID and may be indented)
                case_notes_match = re.search(r'Case Notes:\s*\n?\s*(.*?)(?:\n\s*[A-Za-z]|\n\n|\n-----)', section, re.DOTALL)
                if case_notes_match:
                    # Clean up newlines and excessive whitespace in the notes
                    notes = case_notes_match.group(1).strip()
                    notes = re.sub(r'\n\s+', ' ', notes)  # Replace newline+whitespace with a single space
                    extraction_data['case_notes'] = notes
                    # Track that this field was populated from the log
                    self.log_populated_fields.add('case_notes')
                break
        
        # Find the Source Disk section specifically for device information
        source_disk_section = None
        for section in sections:
            if 'Source Disk' in section and '-----' in section:
                source_disk_section = section
                break
        
        # Process general fields (not device-specific) from the entire content
        general_patterns = {
            'Task:': ('extraction_type', lambda x: x.strip()),
            'Imager Ver:': ('TX1_OS', lambda x: x.strip()),
            'TX1 S/N:': ('extraction_serial', lambda x: x.strip()),
            'Created:': ('extraction_date', self.parse_tx1_date),
            'Examiner:': ('examiner', lambda x: x.strip()),
            'Case Number:': ('case_number', lambda x: x.strip()),
            'Acquisition Md5:': ('md5_hash', lambda x: x.replace(' ', '').strip()), 
        }
        
        # Parse general fields from entire content - DO NOT UPDATE ANY UI FIELDS
        for line in content.split('\n'):
            line = line.strip()
            for key, (field_name, processor) in general_patterns.items():
                if line.startswith(key):
                    value = line[len(key):].strip()
                    extraction_data[field_name] = processor(value)
                    
                    # Track that this field was populated from the log
                    self.log_populated_fields.add(field_name)
                    
                    # DO NOT UPDATE ANY UI FIELDS - just store in extraction_data
                    break
        
        # Process device-specific fields ONLY from Source Disk section
        if source_disk_section:
            device_patterns = {
                'Model:': ('device_model', lambda x: x.strip()),
                'Serial number:': ('device_serial', lambda x: x.strip()),
            }
            
            for line in source_disk_section.split('\n'):
                line = line.strip()
                for key, (field_name, processor) in device_patterns.items():
                    if line.startswith(key):
                        value = line[len(key):].strip()
                        extraction_data[field_name] = processor(value)
                        
                        # Track that this field was populated from the log
                        self.log_populated_fields.add(field_name)
                        
                        # DO NOT update device model or serial fields in the UI - only store in extraction_data
                        break
        
        # Generate formatted date in a readable format
        if 'extraction_date' in extraction_data:
            extraction_data['formatted_date'] = extraction_data['extraction_date']
        else:
            extraction_data['formatted_date'] = "Unknown Date"
                
        return extraction_data
    
    def parse_ftk_log(self, content):
        if re.search(r'Exterro', content, re.IGNORECASE):
            tool_name = 'Exterro FTK Imager'
        elif re.search(r'AccessData', content, re.IGNORECASE):
            tool_name = 'AccessData FTK Imager'
        else:
            tool_name = 'FTK Imager'

        extraction_data = {
            'extraction_tool': tool_name,
            'extraction_type': 'Disk Imaging',
        }

        version_match = re.search(r'FTK®?\s+Imager\s+(\d+\.\d+\.\d+\.\d+)', content)
        if version_match:
            extraction_data['FTK_OS'] = version_match.group(1)
            self.log_populated_fields.add('FTK_OS')

        def field(pattern):
            match = re.search(pattern, content)
            if not match:
                return ''
            return match.group(1).strip()

        case_number = field(r'Case Number:[ \t]*([^\r\n]*)')
        if case_number:
            extraction_data['case_number'] = case_number
            self.log_populated_fields.add('case_number')

        evidence_number = field(r'Evidence Number:[ \t]*([^\r\n]*)')
        if evidence_number:
            extraction_data['evidence_number'] = evidence_number
            self.log_populated_fields.add('evidence_number')

        description = field(r'Unique [Dd]escription:[ \t]*([^\r\n]*)')
        if description:
            extraction_data['case_notes'] = description
            self.log_populated_fields.add('case_notes')

        examiner = field(r'Examiner:[ \t]*([^\r\n]*)')
        if examiner:
            extraction_data['examiner'] = examiner
            self.log_populated_fields.add('examiner')

        date_str = field(r'Acquisition started:[ \t]*([^\r\n]*)')
        if date_str:
            extraction_data['extraction_date'] = self.parse_ftk_date(date_str)
            extraction_data['formatted_date'] = extraction_data['extraction_date']
            self.log_populated_fields.add('extraction_date')
            self.log_populated_fields.add('formatted_date')

        model = field(r'Drive Model:[ \t]*([^\r\n]*)')
        if model:
            extraction_data['device_model'] = model
            self.log_populated_fields.add('device_model')

        serial = field(r'Drive Serial Number:[ \t]*([^\r\n]*)')
        if serial:
            extraction_data['device_serial'] = serial
            self.log_populated_fields.add('device_serial')

        md5_match = re.search(r'MD5 checksum:\s*([a-fA-F0-9]+)', content)
        if md5_match:
            extraction_data['md5_hash'] = md5_match.group(1).strip()
            self.log_populated_fields.add('md5_hash')

        if 'formatted_date' not in extraction_data:
            extraction_data['formatted_date'] = "Unknown Date"
            self.log_populated_fields.add('formatted_date')

        return extraction_data
   
    def parse_tx1_date(self, date_str):
        try:
            # Format: Mon Apr  7 10:05:07 2025 (UTC-0500)
            # Remove timezone info in parentheses for parsing
            cleaned_date = re.sub(r'\s*\([^)]*\)', '', date_str).strip()
            
            # Parse the date using datetime
            date_obj = datetime.strptime(cleaned_date, '%a %b %d %H:%M:%S %Y')
            
            # Format the date for output
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing TX1 date: {e}")
            return date_str

    def parse_ftk_date(self, date_str):
        try:
            # AccessData/Exterro: "Thu Mar 27 10:10:41 2025" or "Fri Jan  9 12:48:09 2026"
            date_str = ' '.join(date_str.split())
            date_obj = datetime.strptime(date_str, '%a %b %d %H:%M:%S %Y')
            
            # Format the date for output
            formatted_date = date_obj.strftime("%A, %B %d, %Y at %H:%M")
            formatted_date = formatted_date.replace(' 0', ' ')  # Remove leading zeros from day
            
            # Get local timezone name
            import time
            is_dst = time.localtime().tm_isdst > 0
            tz_name = self.get_local_timezone_name(0, is_dst)  # Using 0 as offset since we're working with local time
            
            return f"{formatted_date} {tz_name}"
        except Exception as e:
            print(f"Error parsing FTK date: {e}")
            return date_str

    def get_local_timezone_name(self, utc_offset, is_dst):
        import time
        
        # Get the local timezone name directly from the system
        timezone_name = time.tzname[1] if is_dst else time.tzname[0]
        
        # Dictionary of timezone abbreviations to full names
        timezone_names = {
            "PST": "Pacific Standard Time",
            "PDT": "Pacific Daylight Time",
            "MST": "Mountain Standard Time",
            "MDT": "Mountain Daylight Time", 
            "CST": "Central Standard Time",
            "CDT": "Central Daylight Time",
            "EST": "Eastern Standard Time",
            "EDT": "Eastern Daylight Time",
            "HST": "Hawaii Standard Time",
            "AKST": "Alaska Standard Time",
            "AKDT": "Alaska Daylight Time",
            "AST": "Atlantic Standard Time",
            "ADT": "Atlantic Daylight Time"
        }
        
        return timezone_names.get(timezone_name, timezone_name)
        
    def browse_template_file(self, event):
        try:
            file_path = ask_open_file([("Word documents", "*.docx")], folder_kind="template", title="Select DFR template")
            if file_path:
                with open(file_path, 'rb') as test_file:
                    pass
                self.template_file = file_path
                sync_template_choice(self, file_path)
                self.template_drop_label.configure(text=f"Selected: {os.path.basename(file_path)}")
        except IOError as e:
            messagebox.showerror("File Error", f"Could not read the selected file: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

### PARAGRAPHS ###      
    def initialize_pc_exam_paragraphs(self):
        self.paragraphs = load_paragraphs("pc_portable")

    def validate_fields(self):
        missing_fields = []

        examiner_agency = (self.examiner_agency_type.get() or "").strip()
        
        examiner_title = self.examiner_title_type.get()
        if not examiner_title.strip():
            missing_fields.append("Examiner Title")

        device_type = self.device_type_entry.get() if self.device_type.get() == "Other Storage Device" else self.device_type.get()

        # Base fields that are always required
        base_fields = {
            "Device Owner": self.device_owner.get(),
            "Device Type": device_type,
            "Device Make": self.device_PCMan.get(),
            "Examiner Agency": examiner_agency,
            "Examiner Title": examiner_title,
            "Examiner Name": self.examiner_name.get(),
            "Case Number": self.case_number.get(),
            "DFR Number": self.DFR_Num.get() if is_complete_dfr_number(self.DFR_Num.get()) else "",
        }

        # Agency Assist fields (always used now)
        role_fields = {
            "Exam Request Date": self.request_date.get(),
            "Requesting Agency": self.request_agency.get(),
            "Requesting Officer": self.request_officer.get(),
            "Case Type": self.case_type.get(),
        }
        
        if not self.get_request_title():
            missing_fields.append("Requesting Officer Title")
        
        # Check time frame fields
        if (self.legal_authority.get() == 'Search Warrant' and 
            hasattr(self, 'time_frame_var') and self.time_frame_var.get() == 1):
            if not self.time_frame_start.get().strip():
                missing_fields.append("Time Frame Start Date")
            if not self.time_frame_end.get().strip():
                missing_fields.append("Time Frame End Date")

        # Combine all fields
        fields = {**base_fields, **role_fields}

        # Check for custom entry fields
        
        if self.device_type.get() == "Other Storage Device" and not self.device_type_entry.get().strip():
            missing_fields.append("Custom Device Type")

        # Check all fields
        for field_name, field_value in fields.items():
            if not field_value.strip():
                missing_fields.append(field_name)

        # Check transfer fields if transfer is Yes
        if self.device_transfer_var.get() == 1:
            if not self.transfer_title.get().strip():
                missing_fields.append("Transfer Officer Title")
            if not self.transfer_officer.get().strip():
                missing_fields.append("Transfer Officer Name")
            if not self.transfer_agency.get().strip():
                missing_fields.append("Transfer Officer Agency")
            if not self.transfer_date.get().strip():
                missing_fields.append("Transfer Date")

        # Check that at least one forensic software is selected (removed Griffeye)
        selected_software = []
        if hasattr(self, 'axiom_var') and self.axiom_var.get():
            selected_software.append("Axiom")
        if hasattr(self, 'xways_var') and self.xways_var.get():
            selected_software.append("X-Ways")
        
        if not selected_software:
            missing_fields.append("At least one Forensic Software must be selected")

        # Check that files are selected
        if not hasattr(self, 'extraction_file') or not self.extraction_file:
            missing_fields.append("Extraction File (TX1, FTK, X-Ways, or Digital Collector Log)")
        if not hasattr(self, 'template_file') or not self.template_file:
            missing_fields.append("Template File")
        if not hasattr(self, 'extraction_type') or not self.extraction_type:
            missing_fields.append("Extraction Type (TX1, FTK, X-Ways, or Digital Collector)")

        return missing_fields

    def set_document_default_font(self, doc):
        # Set the default character style to Arial
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

    def search_and_replace_content_controls_simple(self, doc, search_docs):
        """
        Enhanced version that handles multi-paragraph replacements for PY_TEXT and PY_ACQUIRE
        with proper newline handling
        """
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        # Get the document XML as bytes
        doc_xml_str = etree.tostring(doc._element, encoding='unicode')
        
        # Perform replacements directly in the XML string
        for search_string, replacement_doc in search_docs.items():
            if search_string in doc_xml_str:
                # Special handling for PY_TEXT and PY_ACQUIRE - multi-paragraph content
                if search_string in ["PY_TEXT", "PY_ACQUIRE"]:
                    # Build the complete replacement XML for all paragraphs
                    replacement_xml_parts = []
                    
                    for para in replacement_doc.paragraphs:
                        # Get the paragraph text
                        para_text = para.text if para.text else ""
                        
                        # For PY_ACQUIRE, don't split by newlines since we already handled that in generate_replacements
                        # Each paragraph already represents a line from the log file
                        text_parts = [para_text] if search_string == "PY_ACQUIRE" else para_text.split('\n')
                        
                        for text_part in text_parts:
                            # Create paragraph XML with proper formatting
                            para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            
                            # Add paragraph properties if needed (for alignment, spacing, etc.)
                            para_xml += '<w:pPr></w:pPr>'
                            
                            # Check if this entire paragraph should be bold/underline
                            # (for headers like "FORENSIC EXTRACTION")
                            is_header = False
                            for run in para.runs:
                                if run.font.bold and run.font.underline:
                                    is_header = True
                                    break
                            
                            if text_part or not text_parts or len(text_parts) == 1:  # Always add run for single empty paragraph
                                para_xml += '<w:r>'
                                para_xml += '<w:rPr>'
                                
                                # Apply formatting based on the original run formatting
                                if is_header or (para.runs and para.runs[0].font.bold):
                                    para_xml += '<w:b/>'
                                if is_header or (para.runs and para.runs[0].font.underline):
                                    para_xml += '<w:u w:val="single"/>'
                                if para.runs and para.runs[0].font.italic:
                                    para_xml += '<w:i/>'
                                
                                # Font settings
                                font_name = 'Arial'
                                font_size = 22  # 11pt * 2 (half-points)
                                if para.runs and para.runs[0].font.name:
                                    font_name = para.runs[0].font.name
                                if para.runs and para.runs[0].font.size:
                                    font_size = para.runs[0].font.size.pt * 2 if hasattr(para.runs[0].font.size, 'pt') else 22
                                
                                para_xml += f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>'
                                para_xml += f'<w:sz w:val="{int(font_size)}"/>'
                                
                                para_xml += '</w:rPr>'
                                
                                if text_part:  # Only add text if there's content
                                    # Add the text content (escape XML special characters)
                                    text_content = text_part
                                    text_content = text_content.replace('&', '&amp;')
                                    text_content = text_content.replace('<', '&lt;')
                                    text_content = text_content.replace('>', '&gt;')
                                    text_content = text_content.replace('"', '&quot;')
                                    text_content = text_content.replace("'", '&apos;')
                                    
                                    para_xml += f'<w:t xml:space="preserve">{text_content}</w:t>'
                                
                                para_xml += '</w:r>'
                            
                            para_xml += '</w:p>'
                            replacement_xml_parts.append(para_xml)
                    
                    # Join all paragraphs
                    replacement_text = ''.join(replacement_xml_parts)
                    
                else:
                    # Handle single paragraph replacements (existing logic)
                    if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                        replacement_text = replacement_doc.paragraphs[0].text
                    else:
                        replacement_text = ""
                    
                    # Escape XML special characters in replacement text
                    replacement_text = replacement_text.replace('&', '&amp;')
                    replacement_text = replacement_text.replace('<', '&lt;')
                    replacement_text = replacement_text.replace('>', '&gt;')
                    replacement_text = replacement_text.replace('"', '&quot;')
                    replacement_text = replacement_text.replace("'", '&apos;')
                
                # Replace in the XML string
                doc_xml_str = doc_xml_str.replace(search_string, replacement_text)
                replaced_strings[search_string] = True
                
                if search_string in ["PY_TEXT", "PY_ACQUIRE"]:
                    paragraph_count = len(replacement_doc.paragraphs) if replacement_doc.paragraphs else 0
                    print(f"Replaced '{search_string}' with formatted content ({paragraph_count} paragraphs)")
                else:
                    print(f"Replaced '{search_string}' with '{replacement_text[:50]}...'")
        
        # Parse the modified XML back into the document
        new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
        doc._element.clear()
        for child in new_element:
            doc._element.append(child)
        
        return replaced_strings
           
    def search_and_replace_split_placeholders(self, doc, search_docs):
        """
        Handle placeholders that are split across XML elements
        """
        import re  # Move the import to the top
        
        replaced_strings = {}
        
        # Initialize tracking
        for search_string in search_docs.keys():
            replaced_strings[search_string] = False
        
        try:
            # Get the document XML as string
            doc_xml_str = etree.tostring(doc._element, encoding='unicode')
            original_xml_str = doc_xml_str
            
            print("\nHandling split placeholders...")
            
            # Look for placeholders that might be split by XML tags
            target_placeholders = ['PY_DFR', 'PY_OWNER', 'PY_EXAMINER', 'PY_CASENUMBER', 'PY_EVIDENCE', 'PY_REQOFF', 'PY_PCSERIAL', 'PY_DEVMAKE', 'PY_TX1VER', 'PY_XWVER', 'PY_DCVER'
            ]
            
            for search_string in target_placeholders:
                if replaced_strings[search_string]:
                    continue
                    
                if search_string not in search_docs:
                    continue
                    
                replacement_doc = search_docs[search_string]
                if replacement_doc.paragraphs and replacement_doc.paragraphs[0].text:
                    replacement_text = replacement_doc.paragraphs[0].text
                else:
                    replacement_text = ""
                
                # Escape XML special characters
                replacement_text = replacement_text.replace('&', '&amp;')
                replacement_text = replacement_text.replace('<', '&lt;')
                replacement_text = replacement_text.replace('>', '&gt;')
                replacement_text = replacement_text.replace('"', '&quot;')
                replacement_text = replacement_text.replace("'", '&apos;')
                
                # Create pattern allowing XML tags between each character
                pattern_parts = []
                for char in search_string:
                    pattern_parts.append(re.escape(char))
                
                # Join with pattern that allows XML tags between characters
                pattern = r'(?:<[^>]*>)*'.join(pattern_parts)
                
                # Find and replace the split placeholder
                matches = list(re.finditer(pattern, doc_xml_str))
                
                if matches:
                    # Replace from last match to first to avoid position shifts
                    for match in reversed(matches):
                        start_pos = match.start()
                        end_pos = match.end()
                        doc_xml_str = doc_xml_str[:start_pos] + replacement_text + doc_xml_str[end_pos:]
                    
                    replaced_strings[search_string] = True
                    print(f"  ✓ Replaced split placeholder '{search_string}' ({len(matches)} occurrences)")
                else:
                    print(f"  ✗ Could not find split placeholder '{search_string}'")
            
            # Parse back the document if changes were made
            if doc_xml_str != original_xml_str:
                new_element = etree.fromstring(doc_xml_str.encode('utf-8'))
                doc._element.clear()
                for child in new_element:
                    doc._element.append(child)
            
        except Exception as e:
            print(f"Split placeholder replacement failed: {e}")
        
        return replaced_strings

    def preview_placeholders(self):
        self.generate_report(preview_only=True)

    def generate_report(self, preview_only=False):
        try:
            missing_fields = self.validate_fields()
            if missing_fields and not preview_only:
                messagebox.showerror(
                    "Missing Fields",
                    f"The following fields are missing:\n\n{', '.join(missing_fields)}\n\nPlease complete them before proceeding."
                )
                return

            # Always use Agency Assist fields
            officer_name = self.request_officer.get().strip()
            officer_last_name = officer_name.split()[-1] if officer_name and ' ' in officer_name else officer_name

            examiner_agency = self.get_examiner_agency()
            request_date = self.parse_request_date(self.request_date.get())
            
            request_agency_formatted, request_agency_abbr = self.format_agency(self.request_agency.get(), return_abbreviation=True)
            examiner_agency_formatted, examiner_agency_abbr = self.format_agency(examiner_agency, return_abbreviation=True)

            # Always use case_type (Agency Assist field)
            case_type = self.case_type.get()

            device_type = self.device_type_entry.get() if self.device_type.get() == "Other Storage Device" else self.device_type.get()

            if device_type.lower() in ['computer', 'loose hard drive']:
                source_device = "source hard drive"
            else:
                source_device = "source device"

            examiner_title = self.examiner_title_type.get()

            data = {
                'Request_Date': request_date,
                'Request_Agency': request_agency_formatted,
                'Request_Agency_Abbr': request_agency_abbr,
                'Request_Title': self.format_title(self.get_request_title()),
                'Request_Officer': self.request_officer.get().title(),
                'Request_Officer_LastName': officer_last_name.title() if officer_last_name else '',
                'Request_Case': case_type,
                'Device_Owner': self.device_owner.get().title(),
                'Device_Type': device_type,
                'device_type': device_type,
                'Device_Model': self.device_PCMod.get().title(),
                'Device_Serial': self.device_PCSerial.get(),
                'Device_Capacity': self.device_capacity.get(),
                'device_capacity': self.device_capacity.get(),
                'Device_Color': self.device_color.get().title() if hasattr(self, 'device_color') else '',
                'Examiner_Agency': examiner_agency_formatted,
                'Examiner_Agency_Abbr': examiner_agency_abbr,
                'Examiner_Title': self.format_title(examiner_title),
                'Examiner_Name': self.examiner_name.get().title(),
                'Forensic_Software': self.get_selected_forensic_software(),
                'Case_Number': self.case_number.get(),
                'evidence_ID': self.evidence_number.get(),
                'DFR_Num': self.DFR_Num.get(),
                'source_device': source_device,
                'device_PCMan': self.device_PCMan.get(),
                'device_PCMod': self.device_PCMod.get(),
                'device_PCSerial': self.device_PCSerial.get(),
            }

            if self.device_type.get() == "Computer":
                data.update({
                    'device_password': self.device_password.get().strip() if hasattr(self, 'device_password') else '',
                    'hd_make': self.hd_make.get().title(),
                    'hd_model': self.hd_model.get().title(),
                    'hd_serial': self.hd_serial.get(),
                    'article_hd': 'an' if self.hd_make.get().lower().startswith(('a','e','i','o','u')) else 'a',
                })

            if self.device_transfer_var.get() == 1:
                transfer_agency_formatted, transfer_agency_abbr = self.format_agency(self.transfer_agency.get(), return_abbreviation=True)
                data.update({
                    'Transfer_Date': self.parse_request_date(self.transfer_date.get()),
                    'Transfer_Title': self.format_title(self.transfer_title.get()),
                    'Transfer_Officer': self.transfer_officer.get().title(),
                    'Transfer_Agency': transfer_agency_formatted,
                    'Transfer_Agency_Abbr': transfer_agency_abbr,
                })

            try:
                extraction_data = self.parse_extraction_file()
            except Exception:
                extraction_data = {}
            
            extraction_fields_to_copy = [
                'case_id',
                'extraction_type',
                'extraction_tool',
                'TX1_OS',
                'FTK_OS',
                'xways_OS',
                'DC_OS',
                'case_name',
                'sha1_hash',
                'image_format',
                'extraction_software',  
                'extraction_serial',
                'extraction_date',
                'formatted_date',
                'md5_hash',
                'evidence_number',
            ]
            
            for field in extraction_fields_to_copy:
                if field in extraction_data and not data.get(field):
                    data[field] = extraction_data[field]
            data = merge_log_device_into_report_data(data, extraction_data, device_type)
                        
            data['article'] = 'an' if device_type.lower().startswith(('a','e','i','o','u')) else 'a'
            if preview_only:
                model = data.get("device_PCMod") or data.get("hd_model") or ""
                suggested = apply_suggested_filename(self, "PCPortable", model)
                show_placeholder_preview(
                    self,
                    pc_preview_rows(
                        data,
                        officer_text=self.format_case_officer(data) if hasattr(self, "format_case_officer") else data.get("Request_Officer", ""),
                        image_date=data.get("formatted_date", ""),
                    ),
                    suggested,
                )
                return
            
            doc = Document(self.template_file)
            self.set_document_default_font(doc)
            
            py_text_found = False
            for para_index, para in enumerate(doc.paragraphs):
                if 'PY_TEXT' in para.text:
                    py_text_found = True
                    print(f"Found PY_TEXT in paragraph {para_index}")
                    
                    para.clear()
                    
                    temp_doc = Document()
                    self.generate_paragraphs(data, temp_doc)
                    
                    parent = para._element.getparent()
                    insert_index = parent.index(para._element)
                    
                    for new_para in reversed(temp_doc.paragraphs):
                        copied_para = doc.add_paragraph()
                        
                        for run in new_para.runs:
                            new_run = copied_para.add_run(run.text)
                            new_run.font.bold = run.font.bold
                            new_run.font.italic = run.font.italic
                            new_run.font.underline = run.font.underline
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                        
                        parent.insert(insert_index, copied_para._element)
                    
                    parent.remove(para._element)
                    break
            
            search_docs = {
                "PY_DFR": Document(),
                "PY_CASENUMBER": Document(),
                "PY_EVIDENCE": Document(),
                "PY_REQDATE": Document(),
                "PY_OWNER": Document(),
                "PY_REQAGENCY": Document(),
                "PY_REQOFF": Document(),
                "PY_LIMITSTART": Document(),
                "PY_LIMITEND": Document(),
                "PY_EXAMINER": Document(),
                "PY_IMAGEDATE": Document(),
                "PY_DEVMAKE": Document(),
                "PY_DEVMODEL": Document(),
                "PY_PCSERIAL": Document(),
                "PY_COLOR": Document(),
                "PY_PASSCODE": Document(),
                "PY_HDMAKE": Document(),
                "PY_HDMODEL": Document(),
                "PY_HDSERIAL": Document(),
                "PY_CAPACITY": Document(),        
                "PY_FTKVER": Document(),
                "PY_TX1VER": Document(),
                "PY_XWVER": Document(),
                "PY_DCVER": Document(),
                "PY_ACQUIRE": Document(),   
            }
            
            self.generate_replacements(data, search_docs)
            
            if not py_text_found:
                print("PY_TEXT not found in paragraphs, trying XML method...")
                new_doc = Document()
                self.generate_paragraphs(data, new_doc)
                search_docs["PY_TEXT"] = new_doc
            
            try:
                replaced_strings = self.search_and_replace_content_controls_simple(doc, search_docs)
                
                missing_strings = [s for s, replaced in replaced_strings.items() if not replaced]
                
                if missing_strings:
                    print(f"Trying split placeholder method for: {missing_strings}")
                    split_replaced = self.search_and_replace_split_placeholders(doc, search_docs)
                    
                    for search_string, was_replaced in split_replaced.items():
                        if was_replaced:
                            replaced_strings[search_string] = True
                
                final_missing = [s for s, replaced in replaced_strings.items() if not replaced]
                
                if final_missing:
                    if "PY_TEXT" in final_missing:
                        messagebox.showwarning("Warning", "The string 'PY_TEXT' was not found in the document.")
                    else:
                        messagebox.showwarning("Warning", f"The following strings were not found in the document: {', '.join(final_missing)}")
                    
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                print(f"Replacement method failed: {e}")
                print(f"Error details: {error_details}")

            apply_suggested_filename(self, "PCPortable", data.get("device_PCMod") or data.get("hd_model") or "")
            self.save_document(doc)
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            messagebox.showerror("Error", f"Failed to generate report:\n\n{str(e)}\n\nDetails:\n{error_details[:500]}")
 
    def get_selected_forensic_software(self):
        selected = []
        if hasattr(self, 'axiom_var') and self.axiom_var.get():
            selected.append("Axiom")
        if hasattr(self, 'xways_var') and self.xways_var.get():
            selected.append("X-Ways")

        return ', '.join(selected) if selected else ''

    def format_request_agency_full(self, data):
        """Returns the full agency name (not abbreviated) for PY_REQAGENCY"""
        # For Case Agent role, use examiner agency (full name)
        if self.role_type.get() == "Case Agent":
            return data.get('Examiner_Agency', '')
        else:  # Agency Assist role
            return data.get('Request_Agency', '')

    def generate_replacements(self, data, search_docs):
        def add_text_to_doc(doc, text):
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p
        
        def add_formatted_log_to_doc(doc, log_content):
            if not log_content:
                return add_text_to_doc(doc, "")
            
            lines = log_content.split('\n')
            
            for i, line in enumerate(lines):
                if not line.strip():
                    p = doc.add_paragraph()
                else:
                    p = doc.add_paragraph(line)
                    
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
        
        def format_date_for_image(extraction_data):                
            if 'formatted_date' in extraction_data and extraction_data['formatted_date']:
                formatted_date_str = extraction_data['formatted_date']
                
                try:
                    import re
                    date_match = re.search(r'(\w+),\s+(\w+)\s+(\d+),\s+(\d+)', formatted_date_str)
                    if date_match:
                        day_name = date_match.group(1)
                        month_name = date_match.group(2)
                        day = int(date_match.group(3))
                        year = int(date_match.group(4))
                        
                        formatted_date = f"{day_name}, {month_name} {day}, {year}"
                        formatted_date = formatted_date.replace(' 0', ' ')
                        return formatted_date
                                
                except Exception as e:
                    print(f"Error parsing formatted_date: {e}")
            
            if 'extraction_date' not in extraction_data:
                return ""
            
            try:
                date_str = extraction_data['extraction_date']
                
                if hasattr(self, 'extraction_type') and self.extraction_type == "TX1":
                    cleaned_date = re.sub(r'\s*\([^)]*\)', '', date_str).strip()
                    date_obj = datetime.strptime(cleaned_date, '%a %b %d %H:%M:%S %Y')
                    formatted_date = date_obj.strftime("%A, %B %d, %Y")
                    formatted_date = formatted_date.replace(' 0', ' ')
                    return formatted_date
                
                elif hasattr(self, 'extraction_type') and self.extraction_type == "FTK":
                    date_obj = datetime.strptime(date_str, '%a %b %d %H:%M:%S %Y')
                    formatted_date = date_obj.strftime("%A, %B %d, %Y")
                    formatted_date = formatted_date.replace(' 0', ' ')
                    return formatted_date
                
                return ""
                
            except Exception as e:
                print(f"Error formatting raw date for PY_IMAGEDATE: {e}")
                return ""
        
        log_content = ""
        if hasattr(self, 'extraction_file') and self.extraction_file:
            try:
                with open(self.extraction_file, 'r', encoding='utf-8', errors='ignore') as file:
                    log_content = file.read()
            except Exception as e:
                print(f"Error reading log file for PY_ACQUIRE: {e}")
                log_content = ""
        
        extraction_data = self.parse_extraction_file() if hasattr(self, 'parse_extraction_file') else {}
        
        image_date = format_date_for_image(extraction_data)
        
        # Get time frame dates (always from Agency Assist fields now)
        time_frame_start = ""
        time_frame_end = ""
        
        if (self.legal_authority.get() == 'Search Warrant' and 
            hasattr(self, 'time_frame_var') and self.time_frame_var.get() == 1):
            try:
                if hasattr(self, 'time_frame_start') and self.time_frame_start.get().strip():
                    time_frame_start = self.parse_request_date(self.time_frame_start.get())
                if hasattr(self, 'time_frame_end') and self.time_frame_end.get().strip():
                    time_frame_end = self.parse_request_date(self.time_frame_end.get())
            except Exception as e:
                print(f"Error parsing time frame dates: {e}")
        
        # Map search strings to data values
        replacement_map = {
            "PY_DFR": data.get('DFR_Num', ''),
            "PY_CASENUMBER": data.get('Case_Number', ''),
            "PY_EVIDENCE": data.get('evidence_ID', ''),
            "PY_REQDATE": data.get('Request_Date', ''),
            "PY_OWNER": data.get('Device_Owner', ''),
            "PY_REQAGENCY": self.format_request_agency_full(data),
            "PY_REQOFF": self.format_case_officer(data),
            "PY_LIMITSTART": time_frame_start,
            "PY_LIMITEND": time_frame_end,
            "PY_EXAMINER": f"{data.get('Examiner_Title', '')} {data.get('Examiner_Name', '')}".strip(),
            "PY_IMAGEDATE": image_date,
            "PY_DEVMAKE": data.get('device_PCMan', ''),
            "PY_DEVMODEL": data.get('device_PCMod', ''),
            "PY_PCSERIAL": data.get('device_PCSerial', ''),
            "PY_COLOR": data.get('Device_Color', ''),
            "PY_PASSCODE": data.get('device_password', ''),
            "PY_HDMAKE": data.get('hd_make', ''),
            "PY_HDMODEL": data.get('hd_model', ''),
            "PY_HDSERIAL": data.get('hd_serial', ''),
            "PY_CAPACITY": data.get('Device_Capacity', ''),
            "PY_FTKVER": data.get('FTK_OS', ''),
            "PY_TX1VER": data.get('TX1_OS', ''),
            "PY_XWVER": data.get('xways_OS', ''),
            "PY_DCVER": data.get('DC_OS', ''),
        }

        if self.device_type.get() != "Computer":
            replacement_map.update({
                "PY_PCMAN": data.get('device_PCMan', ''),
                "PY_PCMOD": data.get('device_PCMod', ''),
            })

        if self.device_type.get() == "Computer":
            replacement_map.update({
                "PY_HDMAKE": data.get('hd_make', ''),
                "PY_HDMODEL": data.get('hd_model', ''),
                "PY_HDSERIAL": data.get('hd_serial', ''),
            })
        
        print("\nGenerating replacement content:")
        
        for search_string, doc in search_docs.items():
            if search_string == "PY_TEXT":
                para_count = len(doc.paragraphs) if doc.paragraphs else 0
                print(f"  {search_string}: {para_count} paragraphs already generated")
                continue 
            
            if search_string == "PY_ACQUIRE":
                print(f"  {search_string}: preserving log file formatting ({len(log_content)} characters)")
                add_formatted_log_to_doc(doc, log_content)
                continue
            
            value = replacement_map.get(search_string, '')
            
            if not value:
                value = ''
                print(f"  {search_string}: (empty)")
            else:
                print(f"  {search_string}: '{value[:50]}...'")
            
            add_text_to_doc(doc, value)
    
    def format_time_frame_date(self, date_string):
        """
        Format time frame dates using the same format as other dates in the application
        """
        if not date_string or not date_string.strip():
            return ""
        
        try:
            # Use the existing parse_request_date method to format consistently
            return self.parse_request_date(date_string.strip())
        except ValueError as e:
            # If parsing fails, return the original string
            print(f"Warning: Could not format time frame date '{date_string}': {e}")
            return date_string
    
    def format_software_name(self, software_name):
        if not software_name:
            return ''
        
        if 'FTK' in software_name:
            return software_name if software_name else 'FTK Imager'
        
        # If the software name contains "X-Ways", simplify it to just "X-Ways Forensics"
        if 'X-Ways' in software_name:
            return 'X-Ways Forensics'
        
        return software_name

    def format_case_officer(self, data):
        parts = [] 
        # Always use Request information (Agency Assist)
        if data.get('Request_Title'):
            abbreviated_title = self.get_title_abbreviation(data['Request_Title'])
            parts.append(abbreviated_title)
        
        if data.get('Request_Officer'):
            parts.append(data['Request_Officer'])

        return ' '.join(parts) if parts else ''

    def format_request_agency_full(self, data):
        """Returns the full agency name (not abbreviated) for PY_REQAGENCY"""
        return data.get('Request_Agency', '')

    def get_title_abbreviation(self, full_title):
        title_abbreviations = {
            "officer": "Off.",
            "deputy": "Dep.",
            "special agent": "SA",
            "supervisory special agent": "SSA",
            "special assistant attorney general": "SAAG",
            "saag": "SAAG",
            "intel analyst": "IA",
            "detective": "Det.",
            "investigator": "Inv.",
            "trooper": "Trooper",
            "computer forensic examiner": "CFE",
        }
        
        # Convert to lowercase for comparison
        title_lower = full_title.lower().strip()
        
        # Return the abbreviation if found, otherwise return the original title
        return title_abbreviations.get(title_lower, full_title)      

    def get_examiner_agency(self):
        agency = (self.examiner_agency_type.get() or "").strip()
        if agency == "South Dakota DCI":
            return "South Dakota Division of Criminal Investigation"
        return agency

    def get_request_title(self):
        return (self.request_title_type.get() or "").strip()

    def generate_paragraphs(self, data, new_doc):
        def add_paragraph_with_style(doc, text):
            text = fill_paragraph(text, data)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
            return p

        device_type = data.get('Device_Type', '').lower()
        
        # Select paragraph based on device type and transfer status
        if device_type == 'computer':
            if self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_computer'
            else:
                paragraph_key = 'one_c_computer'
        elif device_type == 'loose hard drive':
            if self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_loose'
            else:
                paragraph_key = 'one_c_loose'
        else:
            if self.device_transfer_var.get() == 0:
                paragraph_key = 'one_b_storage'
            else:
                paragraph_key = 'one_c_storage'

        # Paragraph 1
        add_paragraph_with_style(new_doc, self.paragraphs[paragraph_key])

        # Paragraph 2 - Authority
        authority_paragraph = self.get_authority_paragraph()
        if authority_paragraph:
            add_paragraph_with_style(new_doc, authority_paragraph)

        # Paragraph 3
        add_paragraph_with_style(new_doc, self.paragraphs['three_a'])

        # Paragraph 4 - Header
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['four'])

        # Paragraph 5 - Extraction method
        if hasattr(self, 'extraction_type'):
            if self.extraction_type == "TX1":
                add_paragraph_with_style(new_doc, self.paragraphs['five_tx1'])
            elif self.extraction_type == "FTK":
                add_paragraph_with_style(new_doc, self.paragraphs['five_ftk'])
            elif self.extraction_type == "XWAYS":
                add_paragraph_with_style(new_doc, self.paragraphs['five_xways'])
            elif self.extraction_type == "DC":
                add_paragraph_with_style(new_doc, self.paragraphs.get('five_dc', self.paragraphs['five_ftk']))

        # Paragraph 6
        add_paragraph_with_style(new_doc, self.paragraphs['six'])
        
        # Paragraph 7 - Header
        self.add_bold_underline_paragraph(new_doc, self.paragraphs['seven'])
        
        # Paragraph 8
        add_paragraph_with_style(new_doc, self.paragraphs['eight'])

        # Paragraph 9
        add_paragraph_with_style(new_doc, self.paragraphs['nine'])

        # Paragraph 10
        add_paragraph_with_style(new_doc, self.paragraphs['ten'])
    
    def save_document(self, doc):
        output_filename = self.output_filename.get().strip()
        
        if not output_filename:
            model = self.device_PCMod.get().strip() if hasattr(self, "device_PCMod") else ""
            output_filename = suggested_report_filename(
                self.DFR_Num.get().strip(),
                "PCPortable",
                self.device_owner.get().strip(),
                model,
            )
        
        save_location = self.save_location.get().strip()
        if not save_location or not os.path.exists(save_location):
            save_location = default_export_dir()
            self.save_location.delete(0, tk.END)
            self.save_location.insert(0, save_location)
        remember_folder("export", save_location)
        output_path = unique_output_path(save_location, output_filename)
        try:
            doc.save(output_path)
            remember_titles_from_form(self)
            remember_agencies_from_form(self)
            messagebox.showinfo("Success", f"Report generated and saved as:\n{output_path}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save report: {str(e)}")
        
    def get_authority_paragraph(self):
        authority_map = {
            'Search Warrant': 'auth_sw',
            'Consent': 'auth_c',
            'Implied Consent': 'auth_i',
            'Parole': 'auth_p'
        }
        
        # Always use legal_authority (Agency Assist field)
        selected_authority = self.legal_authority.get()
        
        if selected_authority not in authority_map:
            raise ValueError(f"Invalid legal authority selected: {selected_authority}")
        return self.paragraphs[authority_map[selected_authority]]

    def format_agency(self, agency, return_abbreviation=False):
        if not agency:
            return "" if not return_abbreviation else ("", "")
                
        # Dictionary mapping normalized lowercase inputs to (full_name, abbreviation)
        agencies = {
            "federal bureau of investigation": ("Federal Bureau of Investigation", "FBI"),
            "fbi": ("Federal Bureau of Investigation", "FBI"),
            "drug enforcement administration": ("Drug Enforcement Administration", "DEA"),
            "dea": ("Drug Enforcement Administration", "DEA"),
            "bureau of alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "alcohol tobacco firearms and explosives": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "atf": ("Bureau of Alcohol Tobacco Firearms and Explosives", "ATF"),
            "department of homeland security": ("Department of Homeland Security", "DHS"),
            "dhs": ("Department of Homeland Security", "DHS"),
            "united states marshals service": ("United States Marshals Service", "USMS"),
            "usms": ("United States Marshals Service", "USMS"),
            "south dakota division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "division of criminal investigation": ("South Dakota Division of Criminal Investigation", "DCI"),
            "dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sddci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "sd dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "south dakota dci": ("South Dakota Division of Criminal Investigation", "DCI"),
            "bureau of indian affairs": ("Bureau of Indian Affairs", "BIA"),
            "bia": ("Bureau of Indian Affairs", "BIA"),
            "south dakota highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "highway patrol": ("South Dakota Highway Patrol", "SDHP"),
            "sdhp": ("South Dakota Highway Patrol", "SDHP"),
            "national security agency": ("National Security Agency", "NSA"),
            "nsa": ("National Security Agency", "NSA"),
            "federal emergency management agency": ("Federal Emergency Management Agency", "FEMA"),
            "fema": ("Federal Emergency Management Agency", "FEMA"),
            "internal revenue service": ("Internal Revenue Service", "IRS"),
            "irs": ("Internal Revenue Service", "IRS"),
            "department of defense": ("Department of Defense", "DOD"),
            "dod": ("Department of Defense", "DOD"),
            "department of justice": ("Department of Justice", "DOJ"),
            "doj": ("Department of Justice", "DOJ")
        }
        
        agency_lower = agency.lower().strip()
        
        # Check for direct match in the agencies dictionary
        if agency_lower in agencies:
            full_name, abbr = agencies[agency_lower]
            if return_abbreviation:
                return full_name, abbr
            else:
                return full_name
        
        # Preserve original for PD/SO handling
        original_agency = agency
        original_lower = original_agency.lower().strip()
        
        # Handle PD and SO expansion and abbreviation creation
        if original_lower.endswith(" pd"):
            # Extract the location part (everything before " pd")
            location = original_agency[:-3].strip()
            agency = f"{location} Police Department"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + PD
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "PD"
        elif original_lower.endswith(" so"):
            # Extract the location part (everything before " so")
            location = original_agency[:-3].strip()
            agency = f"{location} Sheriff's Office"
            agency_lower = agency.lower()
            
            # Generate abbreviation: first letters of location + SO
            words = location.split()
            abbr_chars = [word[0].upper() for word in words if word]
            abbreviation = ''.join(abbr_chars) + "SO"
        else:
            # Default to None if not PD/SO
            abbreviation = None
        
        # Format the agency name with title case
        formatted_name = title_agency_words(agency)
        
        # For short agencies (4 or fewer chars), use uppercase
        if len(agency) <= 4:
            formatted_name = agency.upper()
        
        # If abbreviation wasn't set (non-PD/SO), calculate it
        if abbreviation is None:
            # Check for substring matches in agencies dict (fallback for partial matches)
            found_match = False
            for full_lower, (full_name, abbr) in agencies.items():
                if full_lower in agency_lower:
                    formatted_name = full_name
                    abbreviation = abbr
                    found_match = True
                    break
            
            # If still no match, generate abbreviation dynamically
            if not found_match:
                if len(agency) <= 4:
                    abbreviation = agency.upper()
                else:
                    words = agency.split()
                    skip_words = ['of', 'the', 'and', 'in', 'on', 'at', 'by', 'for', 'with', 'a', 'an']
                    abbr_chars = [word[0].upper() for word in words if word.lower() not in skip_words and word]
                    abbreviation = ''.join(abbr_chars)
        
        if return_abbreviation:
            return formatted_name, abbreviation
        else:
            return formatted_name
    
    def format_title(self, title):
        if len(title) <= 2:
            return title.upper()
        return ' '.join(word.capitalize() for word in title.split())

    def parse_request_date(self, date_string):
        return parse_mdy_date(date_string)

    def back_to_start(self):
        close_and_return(self)
        
    def on_closing(self):
        close_and_return(self)


# Ω Digital Forensics Report Writer Ω (ver. 1.0.1) © 2026 #